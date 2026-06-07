"""
纲目刷格式工具：将纯文本纲目套用 Word 模板样式，返回 (docx 字节流, 文件名)。

模板（位于 format_utils.py 所在目录的上一级 test_B/translate/）：
    中文简体刷格式template.docx
    繁体纲目格式刷template.docx
    英文刷格式template.docx（英文 / 韩文 / 西班牙文共用）

公共逻辑：
    1. 按 \\n 切割文本为段落列表；
    2. 用 Document(template) 加载模板，删除 body 内所有 w:p 与 w:tbl；
    3. 逐段新增，套样式统一直接写 w:pStyle（用 style_id），不用 doc.styles[name]；
    4. 用 io.BytesIO 保存为 bytes，不写磁盘。

分隔符规范化：
    不同翻译方向下，层级标记（壹/一/1/a 与 I./A./1./a.）与内容之间的分隔符可能是
    Tab、全角空格、普通空格、顿号「、」或分号「；」。本模块在识别层级后，统一把分隔符
    规范化为单个 Tab 输出。
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

_BASE = Path(__file__).resolve().parent.parent
ZH_TEMPLATE = str(_BASE / "中文简体刷格式template.docx")
ZHTW_TEMPLATE = str(_BASE / "繁体纲目格式刷template.docx")
EN_TEMPLATE = str(_BASE / "英文刷格式template.docx")

# 中文句末标点
_ZH_END_PUNCT = set("，。、；：？！…—·.,;:!?）)》〉」』】")

# 标记与内容之间可能出现的分隔符（统一规范化为 Tab）
_SEP_CHARS = "\t\u3000 、；"

# 中文大点字符（同时覆盖简体与繁体写法）
_ZH_BIG_CHARS = "壹贰貳叁參参肆伍陆陸柒捌玖拾"
_ZH_MID_CHARS = "一二三四五六七八九十"

_FILENAME_BAD = set('\\/:*?"<>|')


def _first_index(lines: List[str], pred: Callable[[str], bool]) -> Optional[int]:
    for i, t in enumerate(lines):
        if t and pred(t):
            return i
    return None


def _ends_punct_zh(t: str) -> bool:
    return bool(t) and t[-1] in _ZH_END_PUNCT


def _clean_filename(name: Optional[str]) -> str:
    """去掉非法字符与 Tab/换行（保留普通空格），strip 后为空则回退 '纲目'。"""
    if not name:
        return "纲目"
    cleaned = "".join(c for c in name if c not in _FILENAME_BAD and c not in "\t\r\n")
    cleaned = cleaned.strip()
    return cleaned or "纲目"


def _clean_filename_keep_tab(name: Optional[str]) -> str:
    """去掉非法字符与换行，保留内部 Tab；首尾空白 strip 后为空则回退 '纲目'。"""
    if not name:
        return "纲目"
    cleaned = "".join(c for c in name if c not in _FILENAME_BAD and c not in "\r\n")
    cleaned = cleaned.strip()
    return cleaned or "纲目"


def _set_style_id(p, style_id: str) -> None:
    """直接写 w:pStyle（val=style_id）。"""
    pPr = p.get_or_add_pPr()
    existing = pPr.find(qn("w:pStyle"))
    if existing is not None:
        pPr.remove(existing)
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.insert(0, pStyle)


def _set_page_break_before(p) -> None:
    pPr = p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


def _clear_body(doc) -> None:
    """删除 body 内所有 w:p 与 w:tbl，保留 sectPr 等其它节点。"""
    body = doc.element.body
    for el in list(body):
        if el.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(el)


def _add_paragraph(doc, text: str, style_id: Optional[str], page_break_before: bool) -> None:
    """在 sectPr 之前插入段落，套样式并按需加段前分页符。"""
    body = doc.element.body
    p = OxmlElement("w:p")
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(p)
    else:
        body.append(p)
    if style_id:
        _set_style_id(p, style_id)
    if page_break_before:
        _set_page_break_before(p)
    para = Paragraph(p, doc)
    if text:
        para.add_run(text)


def _render(template_path: str, items: List[Tuple[str, Optional[str], bool]]) -> bytes:
    doc = Document(template_path)
    _clear_body(doc)
    for text, style_id, page_break in items:
        _add_paragraph(doc, text, style_id, page_break)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ---------------------------------------------------------------------------
# 中文（简体 / 繁体）
# ---------------------------------------------------------------------------

_DUJING_PREFIXES = ("读经", "讀經")
_MINISTRY_PREFIXES = ("职事信息摘录", "職事信息摘錄", "职事摘录", "職事摘錄")
_REF_PREFIXES = ("参考与参读资料", "參考與參讀資料", "参读资料", "參讀資料")

# 各级标记的“核心”正则（不含其后的句点/分隔符）
_BIG_CORE = re.compile(r"^[" + _ZH_BIG_CHARS + r"]+")
_MID_CORE = re.compile(r"^[" + _ZH_MID_CHARS + r"]+")
_PAREN_CORE = re.compile(r"^（[一二三四五六七八九十百]+）")
_DIGIT_CORE = re.compile(r"^\d+")
_LOWER_CORE = re.compile(r"^[a-z]+")
_LEADING_NUM_RE = re.compile(r"^(\d+)")

_SEP_SET = set(_SEP_CHARS)
_PERIOD_CHARS = ".．。"


def _split_marker_zh(t: str, core_re, self_delim: bool = False) -> Optional[str]:
    """若 t 以某层级标记开头，返回规范化后的 'marker\\t内容'（分隔符统一为 Tab）；
    否则返回 None。

    self_delim=True 用于（一）这类自带边界的标记，标记后允许直接接内容；
    其余标记要求标记后是句点或分隔符，避免把正文（如 '十二使徒'、'11~13节'）误判。
    保留标记原有的句点（如 '1.'、'a.'），与大写数字（壹/一）惯例并存。
    """
    m = core_re.match(t)
    if not m:
        return None
    marker_core = m.group(0)
    tail = t[len(marker_core):]
    period = ""
    if tail[:1] in _PERIOD_CHARS:
        period = tail[0]
        tail = tail[1:]
    if not self_delim and not period:
        if tail and tail[0] not in _SEP_SET:
            return None
    content = tail.lstrip(_SEP_CHARS)
    marker = marker_core + period
    return marker + ("\t" + content if content else "")


def _is_marker_zh(t: str) -> bool:
    return any(
        _split_marker_zh(t, core, sd) is not None
        for core, sd in (
            (_BIG_CORE, False), (_MID_CORE, False),
            (_PAREN_CORE, True), (_DIGIT_CORE, False),
        )
    )


def _render_zh_like(text: str, template: str, styles: Dict[str, Optional[str]]) -> Tuple[bytes, str]:
    lines = [ln.strip() for ln in (text or "").split("\n")]

    dujing_idx = _first_index(lines, lambda t: t.startswith(_DUJING_PREFIXES))
    ministry_idx = _first_index(lines, lambda t: t.startswith(_MINISTRY_PREFIXES))
    ref_idx = _first_index(lines, lambda t: t.startswith(_REF_PREFIXES))

    # 标题区：读经行之上的所有非空行；若无读经行，则取第一个层级标记之前
    limit = dujing_idx
    if limit is None:
        for k, tt in enumerate(lines):
            if tt and _is_marker_zh(tt):
                limit = k
                break
    if limit is None:
        limit = len(lines)
    title_indices = [k for k in range(limit) if lines[k]]
    title_pos = {idx: pos for pos, idx in enumerate(title_indices)}
    n_titles = len(title_indices)

    # 文件名取篇题行（最后一个标题行），保留内部 Tab
    filename = (
        _clean_filename_keep_tab(lines[title_indices[-1]])
        if title_indices
        else "纲目"
    )

    items: List[Tuple[str, Optional[str], bool]] = []
    for i, t in enumerate(lines):
        if not t:
            items.append((t, None, False))
            continue
        sid: Optional[str] = None
        pb = False
        out = t
        if i in title_pos:
            pos = title_pos[i]
            if pos == n_titles - 1:
                sid = styles["title_last"]
            elif pos == 0:
                sid = styles["title_first"]
            else:
                sid = styles["title_mid"]
        elif dujing_idx is not None and i == dujing_idx:
            sid = styles["dujing"]
        elif t.startswith(_MINISTRY_PREFIXES):
            sid, pb = styles["ministry"], True
        elif t.startswith(_REF_PREFIXES):
            sid = styles["ref"]
        elif (
            ref_idx is not None and i > ref_idx
            and styles.get("ref_small") and _LEADING_NUM_RE.match(t)
        ):
            num = int(_LEADING_NUM_RE.match(t).group(1))
            sid = styles["ref_small"] if num <= 9 else styles["ref_large"]
            split = _split_marker_zh(t, _DIGIT_CORE)
            out = split if split is not None else t
        elif (
            ministry_idx is not None and i > ministry_idx
            and (ref_idx is None or i < ref_idx)
        ):
            sid = styles["excerpt_body"] if _ends_punct_zh(t) else styles["excerpt_title"]
        elif _split_marker_zh(t, _BIG_CORE) is not None:
            sid = styles["big"]
            out = _split_marker_zh(t, _BIG_CORE)
        elif _split_marker_zh(t, _MID_CORE) is not None:
            sid = styles["mid"]
            out = _split_marker_zh(t, _MID_CORE)
        elif _split_marker_zh(t, _PAREN_CORE, self_delim=True) is not None:
            sid = styles["paren"]
            out = _split_marker_zh(t, _PAREN_CORE, self_delim=True)
        elif _split_marker_zh(t, _DIGIT_CORE) is not None:
            sid = styles["digit"]
            out = _split_marker_zh(t, _DIGIT_CORE)
        elif _split_marker_zh(t, _LOWER_CORE) is not None:
            sid = styles["lower"]
            out = _split_marker_zh(t, _LOWER_CORE)
        else:
            sid = styles["default"]
        items.append((out, sid, pb))
    return _render(template, items), filename


_ZH_STYLES: Dict[str, Optional[str]] = {
    "title_first": "aa",   # 第一行 总题
    "title_mid": "ac",     # 第二行 系列题
    "title_last": "ae",    # 第三行 篇题
    "dujing": "af0",       # 正文前 读经
    "ministry": "af2",     # 职事信息摘录（段前分页）
    "ref": "af4",          # 参考与参读资料
    "big": "10",           # 1级 壹大点
    "mid": "20",           # 2级 一中点
    "paren": "50",         # 5级 小小点
    "digit": "31",         # 3级 1小点
    "lower": "4a",         # 4级 a小点
    "excerpt_title": "af6",  # 摘录一级标题
    "excerpt_body": "af8",   # 摘录正文
    "ref_small": "19",     # 参考参读 1~9
    "ref_large": "100",    # 参考参读 10+
    "default": "af8",
}

_ZHTW_STYLES: Dict[str, Optional[str]] = {
    "title_first": "ae",   # 新总题
    "title_mid": "ae",     # （繁体无系列题，沿用新总题）
    "title_last": "af0",   # 新篇题
    "dujing": "af2",       # 读经
    "ministry": "afa",     # 职事摘录（段前分页）
    "ref": "aff2",         # 参读资料
    "big": "af4",          # 新大点
    "mid": "af6",          # 新中点
    "paren": "aff8",       # 小小点样式
    "digit": "af8",        # 新小点
    "lower": "ab0",        # 新ab点
    "excerpt_title": "afc",  # 摘录标题
    "excerpt_body": "afe",   # 新正文
    "ref_small": None,
    "ref_large": None,
    "default": "afe",
}


def format_zh(text: str) -> Tuple[bytes, str]:
    return _render_zh_like(text, ZH_TEMPLATE, _ZH_STYLES)


def format_zhtw(text: str) -> Tuple[bytes, str]:
    return _render_zh_like(text, ZHTW_TEMPLATE, _ZHTW_STYLES)


# ---------------------------------------------------------------------------
# 英文 / 韩文 / 西班牙文（共用英文模板）
# ---------------------------------------------------------------------------

# 大点罗马数字 → 样式 id
_EN_ROMAN_MAP = {
    "I": "2I",
    "II": "2IIV", "V": "2IIV", "X": "2IIV",
    "III": "2III",
    "IV": "2IVVI", "VI": "2IVVI", "IX": "2IVVI", "XI": "2IVVI",
    "VII": "2VII",
    "VIII": "2VIII", "XIII": "2VIII",
}

# 大点期望序列（按层级顺序消歧 I/V/X/C 等罗马字符与中点字母）
_EN_ROMANS = [
    "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
    "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII",
]

# 摘录区起始行（职事信息摘录）——兼容英 / 韩 / 西 多种写法
_EXCERPT_PREFIXES = (
    "ministry excerpts", "excerpts from the ministry", "ministry excerpt",
    "사역 메시지 발췌", "메시지 발췌", "사역 발췌", "발췌",
    "extractos del ministerio", "extracto del ministerio",
)
# 参考与参读资料
_REFS_PREFIXES = (
    "references and further reading", "references for further reading", "further reading",
    "참고 및 추가", "참고자료", "참고 자료", "추가 읽을 자료",
    "referencias y lecturas adicionales", "referencias y lectura adicional", "referencias",
)
# 摘录正文（句子）句末标点：以此结尾视为正文，否则视为小标题
_EXCERPT_BODY_END = "。.!！?？;；…"
_EXCERPT_TRAIL_STRIP = "\"”’）)】》」』 　"

_EN_SEP_CHARS = "\t\u3000 "
_EN_UPPER_RE = re.compile(r"^((?:[IVXLCDM]+|[A-Z])\.)")
_EN_DIGIT_RE = re.compile(r"^(\d+\.)")
_EN_LOWER_RE = re.compile(r"^([a-z]\.)")
# 任一层级标记（要求标记后是分隔符或行尾，用于定位正文起点）
_EN_ANY_MARKER = re.compile(
    r"^((?:[IVXLCDM]+|[A-Z])\.|\d+\.|[a-z]\.)(?=[" + re.escape(_EN_SEP_CHARS) + r"]|$)"
)


def _en_startswith(t: str, prefixes) -> bool:
    tl = t.lower()
    return any(tl.startswith(p) for p in prefixes)


def _retab_en(text: str, marker: str) -> str:
    rest = text[len(marker):].lstrip(_EN_SEP_CHARS)
    return marker + ("\t" + rest if rest else "")


def _classify_en_marker(letter: str, state: Dict[str, int]) -> Tuple[str, Optional[str]]:
    """按层级顺序判断大写标记是大点(big)还是中点(mid)，消除 I/V/X/C 的歧义。"""
    big_count = state["big"]
    mid_count = state["mid"]
    exp_big = _EN_ROMANS[big_count] if big_count < len(_EN_ROMANS) else None
    exp_mid = chr(ord("A") + mid_count) if mid_count < 26 else None

    if letter == exp_big:
        state["big"] += 1
        state["mid"] = 0
        return ("big", letter)
    if letter == exp_mid:
        state["mid"] += 1
        return ("mid", None)
    if re.fullmatch(r"[IVXLCDM]{2,}", letter):  # 多字符罗马数字 → 大点
        state["big"] += 1
        state["mid"] = 0
        return ("big", letter)
    if len(letter) == 1 and letter in "IVXLM":  # 单字符罗马（排除 C/D 等中点字母）→ 大点
        state["big"] += 1
        state["mid"] = 0
        return ("big", letter)
    state["mid"] += 1
    return ("mid", None)


def format_en(text: str) -> Tuple[bytes, str]:
    lines = [ln.strip() for ln in (text or "").split("\n")]

    # 正文起点 = 第一个层级标记行；其上为标题区（按位置识别，兼容英/韩/西）
    first_marker_idx = _first_index(lines, lambda t: bool(_EN_ANY_MARKER.match(t)))
    header_end = first_marker_idx if first_marker_idx is not None else len(lines)
    header_nonempty = [k for k in range(header_end) if lines[k]]
    scripture_idx = header_nonempty[-1] if header_nonempty else None
    title_idx = header_nonempty[-2] if len(header_nonempty) >= 2 else None

    if title_idx is not None:
        filename = _clean_filename(lines[title_idx])
    elif scripture_idx is not None:
        filename = _clean_filename(lines[scripture_idx])
    else:
        filename = "纲目"

    excerpt_idx = _first_index(lines, lambda t: _en_startswith(t, _EXCERPT_PREFIXES))
    ref_idx = _first_index(lines, lambda t: _en_startswith(t, _REFS_PREFIXES))

    state = {"big": 0, "mid": 0}
    items: List[Tuple[str, Optional[str], bool]] = []
    for i, t in enumerate(lines):
        if not t:
            items.append((t, None, False))
            continue
        sid: Optional[str] = None
        pb = False
        out = t
        if scripture_idx is not None and i == scripture_idx:
            sid = "1"    # 读经
        elif title_idx is not None and i == title_idx:
            sid = "11"   # 篇题
        elif i < header_end and lines[i]:
            sid = "111"  # 篇题上方的大标题
        elif _en_startswith(t, _EXCERPT_PREFIXES):
            sid, pb = "B1", True
        elif _en_startswith(t, _REFS_PREFIXES):
            sid = "B3"
        elif (
            excerpt_idx is not None and i > excerpt_idx
            and (ref_idx is None or i < ref_idx)
        ):
            # 摘录区：以句末标点结尾视为摘录正文，否则视为摘录小标题
            stripped = t.rstrip(_EXCERPT_TRAIL_STRIP)
            if stripped and stripped[-1] in _EXCERPT_BODY_END:
                sid = "00"   # 摘录正文
            else:
                sid = "B21"  # B2摘录标题
        elif _EN_DIGIT_RE.match(t):
            marker = _EN_DIGIT_RE.match(t).group(1)
            sid = "4"    # 小点
            out = _retab_en(t, marker)
        elif _EN_LOWER_RE.match(t):
            marker = _EN_LOWER_RE.match(t).group(1)
            sid = "5"    # 小小点
            out = _retab_en(t, marker)
        elif _EN_UPPER_RE.match(t):
            marker = _EN_UPPER_RE.match(t).group(1)  # 含句点，如 "I." / "C."
            letter = marker[:-1]
            role, _ = _classify_en_marker(letter, state)
            sid = _EN_ROMAN_MAP.get(letter, "2I") if role == "big" else "3"
            out = _retab_en(t, marker)
        else:
            sid = "00"
        items.append((out, sid, pb))
    return _render(EN_TEMPLATE, items), filename
