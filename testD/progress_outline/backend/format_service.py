# -*- coding: utf-8 -*-
import cn2an
import io
import re
import zipfile
from copy import deepcopy
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
BACKEND_DIR = Path(__file__).resolve().parent
ZH_TEMPLATE = BACKEND_DIR / "中文纲目模板.docx"

# ── 英文/西班牙文层级识别 ──────────────────────────────────────────────
ROMAN_STYLE_MAP = [
    (r"^(VIII\.|Ⅷ\.)", "2大点VIII"),
    (r"^(XIII\.)", "2大点VIII"),
    (r"^(VII\.|Ⅶ\.)", "2大点VII"),
    (r"^(XII\.|Ⅻ\.)", "2大点VII"),
    (r"^(IV\.|Ⅳ\.)", "2大点IV，VI"),
    (r"^(VI\.|Ⅵ\.)", "2大点IV，VI"),
    (r"^(IX\.|Ⅸ\.)", "2大点IV，VI"),
    (r"^(XI\.|Ⅺ\.)", "2大点IV，VI"),
    (r"^(III\.|Ⅲ\.)", "2大点III"),
    (r"^(II\.|Ⅱ\.)", "2大点II，V"),
    (r"^(V\.|Ⅴ\.)", "2大点II，V"),
    (r"^(X\.|Ⅹ\.)", "2大点II，V"),
    (r"^(I\.|Ⅰ\.)", "2大点I"),
]
LEVEL_PATTERNS = {
    "大点": [p for p, _ in ROMAN_STYLE_MAP],
    "中点": [r"^[A-H]\."],
    "小点": [r"^[1-9]\."],
    "小小点": [r"^[a-k]\."],
}
LEVEL_ORDER = ["大点", "中点", "小点", "小小点"]
TRAILING_PUNCTS = set(".,:;!?")

# 中文格式规则
ZH_PATTERN_STYLES = [
    (r"^([壹贰貳叁参肆伍陆陸柒捌玖拾佰仟萬亿億]+)　", "81级标题"),
    (r"^([一二三四五六七八九十百千万萬亿億]+)　", "82级标题"),
    (r"^(\d+)　", "83级标题"),
    (r"^([a-z])　", "84级标题"),
    (r"^([壹贰貳叁参肆伍陆陸柒捌玖拾佰仟萬亿億]+)\t", "2大点"),
    (r"^([一二三四五六七八九十百千万萬亿億]+)\t", "3中点"),
    (r"^(\d+)\t", "4小点"),
    (r"^([a-z])\t", "5a点"),
]


# ════════════════════════════════════════════════════════════
#  中文文字处理
# ════════════════════════════════════════════════════════════


def fullwidth_to_halfwidth(text: str) -> str:
    result = []
    for char in text:
        code = ord(char)
        if 0xFF10 <= code <= 0xFF19:
            result.append(chr(code - 0xFF10 + ord("0")))
        elif 0xFF21 <= code <= 0xFF3A:
            result.append(chr(code - 0xFF21 + ord("A")))
        elif 0xFF41 <= code <= 0xFF5A:
            result.append(chr(code - 0xFF41 + ord("a")))
        else:
            result.append(char)
    return "".join(result)


def replace_english_punctuation(text: str) -> str:
    mapping = {
        ",": "，",
        ";": "；",
        ":": "：",
        "?": "？",
        "!": "！",
        "(": "（",
        ")": "）",
    }
    for en, zh in mapping.items():
        text = text.replace(en, zh)
    text = re.sub(r"(?<![A-Za-z])\.(?![A-Za-z])", "。", text)
    text = text.replace("-", "—")
    return text


def replace_chinese_single_quotes(text: str) -> str:
    return text.replace("\u2018", "\u201c").replace("\u2019", "\u201d")


def process_zh_text(raw: str) -> str:
    lines = raw.splitlines()
    lines = [l for l in lines if l.strip()]
    result = []
    for line in lines:
        t = fullwidth_to_halfwidth(line)
        t = replace_english_punctuation(t)
        t = replace_chinese_single_quotes(t)
        t = t.lstrip("　\t ")
        if not ("读经：" in t or "讀經：" in t):
            t = t.replace("　", "\t").replace(" ", "\t")
        for kw in ["篇", "章", "课"]:
            t = t.replace(f"{kw}\t", f"{kw}　")
        t = t.replace("叄\t", "叁\t").replace("叄　", "叁　")
        t = t.replace("陸\t", "陆\t").replace("陸　", "陆　")
        t = t.replace("貳\t", "贰\t").replace("貳　", "贰　")
        t = t.replace("贮\t", "贰\t").replace("参\t", "叁\t")
        t = t.replace("\t\t", "\t")
        t = t.replace("～", "~")
        result.append(t)

    scripture_idx_text = None
    for idx, t in enumerate(result):
        if "读经：" in t or "讀經：" in t:
            scripture_idx_text = idx
            break

    final = []
    for i, t in enumerate(result):
        stripped = t.rstrip()
        if scripture_idx_text is not None and i <= scripture_idx_text:
            final.append(stripped)
            continue
        next_t = None
        for j in range(i + 1, len(result)):
            if result[j].strip():
                next_t = result[j]
                break
        has_sub = False
        if next_t:
            sub_markers = ("一\t", "1\t", "a\t")
            if any(next_t.startswith(m) for m in sub_markers):
                has_sub = True
        if stripped and not stripped.endswith(
            ("。", "！", "？", "…", '"', "\u2019", "：", "』", "读经：", "讀經：")
        ):
            stripped = stripped + ("：" if has_sub else "。")
        elif stripped.endswith("；"):
            stripped = stripped[:-1] + "。"
        final.append(stripped)

    return "\n".join(final)


# ════════════════════════════════════════════════════════════
#  英文/西班牙文文字处理
# ════════════════════════════════════════════════════════════


def get_en_level(line: str):
    text = line.strip()
    for level in LEVEL_ORDER:
        for pat in LEVEL_PATTERNS[level]:
            if re.match(pat, text):
                return level
    return None


def is_lower_level(cur: str, nxt: str) -> bool:
    if cur not in LEVEL_ORDER or nxt not in LEVEL_ORDER:
        return False
    return LEVEL_ORDER.index(nxt) > LEVEL_ORDER.index(cur)


def replace_marker_sep(line: str, pattern: str) -> str:
    m = re.match(pattern, line.strip())
    if not m:
        return line
    marker = m.group(0)
    rest = line.strip()[len(marker) :]
    body = re.search(r"[A-Za-z\u4e00-\u9fff]", rest)
    if not body:
        return line
    return marker + "\t" + rest[body.start() :]


def process_en_text(raw: str, lang: str = "en") -> str:
    if lang == "es":
        scripture_marker = "Lectura bíblica:"
        excerpts_marker = "Extractos del ministerio:"
    else:
        scripture_marker = "Scripture Reading:"
        excerpts_marker = "Excerpts from the Ministry:"

    lines = raw.splitlines()
    lines = [l for l in lines if l.strip()]
    lines = [l.replace("\u3000", " ") for l in lines]

    processed = []
    for line in lines:
        text = line.strip()
        if (
            text.lower().startswith(scripture_marker.lower())
            or text.lower().startswith(excerpts_marker.lower())
        ):
            processed.append(line)
            continue
        replaced = False
        for pattern, _ in ROMAN_STYLE_MAP:
            if re.match(pattern, text):
                line = replace_marker_sep(text, pattern)
                replaced = True
                break
        if not replaced:
            for pat in [r"^([A-H]\.)", r"^([1-9]\.)", r"^([a-k]\.)"]:
                if re.match(pat, text):
                    line = replace_marker_sep(text, pat)
                    break
        processed.append(line)

    scripture_idx_en = None
    for idx, l in enumerate(processed):
        if l.strip().lower().startswith(scripture_marker.lower()):
            scripture_idx_en = idx
            break

    final = []
    for i, line in enumerate(processed):
        if scripture_idx_en is not None and i <= scripture_idx_en:
            final.append(line)
            continue
        cur_level = get_en_level(line)
        if cur_level is None:
            final.append(line)
            continue
        next_level = None
        for j in range(i + 1, len(processed)):
            if processed[j].strip():
                next_level = get_en_level(processed[j])
                break
        use_colon = next_level is not None and is_lower_level(cur_level, next_level)
        target = ":" if use_colon else "."
        stripped = line.rstrip()
        if stripped and stripped[-1] in TRAILING_PUNCTS:
            stripped = stripped[:-1] + target
        else:
            stripped = stripped + target
        final.append(stripped)

    return "\n".join(final)


# ════════════════════════════════════════════════════════════
#  docx 生成
# ════════════════════════════════════════════════════════════


def copy_styles(src_doc, dst_doc):
    src_elem = src_doc.styles.element
    dst_elem = dst_doc.styles.element
    existing = {s.get(qn("w:styleId")) for s in dst_elem.findall(qn("w:style"))}
    for style_elem in src_elem.findall(qn("w:style")):
        sid = style_elem.get(qn("w:styleId"))
        if not sid:
            continue
        if sid in existing:
            for old in dst_elem.findall(qn("w:style")):
                if old.get(qn("w:styleId")) == sid:
                    dst_elem.remove(old)
                    break
        dst_elem.append(deepcopy(style_elem))


def apply_style(para, style_name: str, doc):
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass


def generate_filename(lines_list: list, scripture_idx: int | None, lang: str) -> str:
    """根据篇题行生成文件名"""
    import re as _re

    if scripture_idx is None or scripture_idx < 1:
        return f"formatted_{lang}.docx"

    title_line = lines_list[scripture_idx - 1].strip()

    if lang == "zh":
        parts = title_line.split("　")
        if len(parts) >= 2:
            serial_part = fullwidth_to_halfwidth(parts[0].strip())
            content_part = parts[1].strip()
            m = _re.match(
                r"^第\s*([0-9]+|[一二三四五六七八九十百千万亿]+)\s*([章篇课])$",
                serial_part,
            )
            if m:
                try:
                    arabic_num = cn2an.cn2an(m.group(1), "smart")
                    serial_str = f"msg. {arabic_num}"
                except Exception:
                    serial_str = serial_part
            else:
                serial_str = serial_part
            content_part = _re.sub(r'[\\/:*?"<>|]', "", content_part)
            return f"{serial_str} {content_part}.docx"
        else:
            name = _re.sub(r'[\\/:*?"<>|]', "", title_line)
            return f"{name}.docx"

    else:
        WORD_TO_NUM = {
            "One": "1",
            "Two": "2",
            "Three": "3",
            "Four": "4",
            "Five": "5",
            "Six": "6",
            "Seven": "7",
            "Eight": "8",
            "Nine": "9",
            "Ten": "10",
            "Eleven": "11",
            "Twelve": "12",
            "Thirteen": "13",
            "Fourteen": "14",
            "Fifteen": "15",
            "Sixteen": "16",
            "Seventeen": "17",
            "Eighteen": "18",
            "Nineteen": "19",
            "Twenty": "20",
        }
        m = _re.match(r"^(Message|Lesson)\s+([A-Z][a-z]+|\d+)(.*)", title_line)
        if m:
            keyword = m.group(1)
            token = m.group(2)
            rest = m.group(3).strip()
            num = WORD_TO_NUM.get(token, token)
            prefix = "msg." if keyword == "Message" else "lsn."
            rest_clean = _re.sub(r'[\\/:*?"<>|]', "", rest)
            name_parts = [prefix, num]
            if rest_clean:
                name_parts.append(rest_clean)
            return " ".join(name_parts) + ".docx"
        else:
            name = _re.sub(r'[\\/:*?"<>|]', "", title_line)
            return f"{name}.docx"


def make_zh_docx(text: str) -> bytes:
    doc = Document(str(ZH_TEMPLATE))
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    lines_list = [l for l in text.splitlines() if l.strip()]

    scripture_idx = None
    for i, line in enumerate(lines_list):
        if "读经：" in line or "讀經：" in line:
            scripture_idx = i
            break

    for i, line in enumerate(lines_list):
        para = doc.add_paragraph(line)

        if scripture_idx is not None and i < scripture_idx:
            diff = scripture_idx - i
            if diff == 1:
                apply_style(para, "00篇题", doc)
            elif diff == 2:
                apply_style(para, "11111西列", doc)
            elif diff >= 3:
                apply_style(para, "0系列", doc)
            continue

        if "读经：" in line or "讀經：" in line:
            apply_style(para, "11读经", doc)
            continue

        skip = False
        for kw in ["职事信息摘录：", "研读问题：", "出处与参读：", "參考與參讀："]:
            if kw in line:
                apply_style(para, "9职事信息摘录", doc)
                skip = True
                break
        if skip:
            continue

        for pattern, style_name in ZH_PATTERN_STYLES:
            if re.match(pattern, line):
                apply_style(para, style_name, doc)
                break

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_en_docx(text: str, lang: str = "en") -> bytes:
    if lang == "es":
        scripture_marker = "Lectura bíblica:"
        excerpts_marker = "Extractos del ministerio:"
    else:
        scripture_marker = "Scripture Reading:"
        excerpts_marker = "Excerpts from the Ministry:"

    doc = Document(str(EN_TEMPLATE))
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    lines_list = [l for l in text.splitlines() if l.strip()]

    scripture_idx = None
    for i, line in enumerate(lines_list):
        if line.strip().lower().startswith(scripture_marker.lower()):
            scripture_idx = i
            break

    for i, line in enumerate(lines_list):
        para = doc.add_paragraph(line)
        text_s = line.strip()

        if scripture_idx is not None and i < scripture_idx:
            diff = scripture_idx - i
            if diff == 1:
                apply_style(para, "11篇题", doc)
            elif diff >= 2:
                apply_style(para, "11大标题", doc)
            continue

        if text_s.lower().startswith(scripture_marker.lower()):
            apply_style(para, "1读经", doc)
            continue

        if text_s.lower().startswith(excerpts_marker.lower()):
            apply_style(para, "B1职事信息摘录", doc)
            continue

        matched = False
        for pattern, style_name in ROMAN_STYLE_MAP:
            if re.match(pattern, text_s):
                apply_style(para, style_name, doc)
                matched = True
                break
        if not matched:
            if re.match(r"^[A-H]\.", text_s):
                apply_style(para, "3中点", doc)
            elif re.match(r"^[1-9]\.", text_s):
                apply_style(para, "4小点", doc)
            elif re.match(r"^[a-k]\.", text_s):
                apply_style(para, "5小小点", doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════
#  接口
# ════════════════════════════════════════════════════════════


def format_zh_docx(text: str) -> tuple[bytes, str]:
    """处理中文纲目文本并返回 (docx_bytes, filename)。"""
    if not text.strip():
        raise ValueError("文字内容不能为空")
    processed = process_zh_text(text)
    lines_list = [l for l in processed.splitlines() if l.strip()]
    scripture_idx = None
    for i, line in enumerate(lines_list):
        if "读经：" in line or "讀經：" in line:
            scripture_idx = i
            break
    docx_bytes = make_zh_docx(processed)
    filename = generate_filename(lines_list, scripture_idx, "zh")
    return docx_bytes, filename
