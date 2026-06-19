# -*- coding: utf-8 -*-
"""
将「主恢复中神圣启示的进展」Word 纲目批量解析并写入 Elasticsearch。
解析目标：「纲目带出处」版本（文件名含「纲目带出处」，或内容防火墙判定）。

用法：
  cd back_mic/backend
  python scripts/progress_outline/ingest_pano.py --source-dir <Word 根目录>
  python scripts/progress_outline/ingest_pano.py --source-dir <dir> --export [FILE]
  python scripts/progress_outline/ingest_pano.py --import-json progress_pano.json
  python scripts/progress_outline/ingest_pano.py --source-dir <dir> --doc-id progress_pano-1-1-1
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Iterator, TextIO

from docx import Document
from dotenv import load_dotenv
from natsort import natsorted

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parents[1]
sys.path.insert(0, str(BACKEND_DIR))
load_dotenv(BACKEND_DIR / ".env")

from es_config import es  # noqa: E402

ES_INDEX = "progress_pano"
EXPORT_JSON = SCRIPT_DIR / "progress_pano.json"
BULK_CHUNK = 200
LOG_FILE = SCRIPT_DIR / "ingest_pano.log"

_log_fp: TextIO | None = None
_root_dir: Path | None = None

SOURCE_GROUP_MAP: dict[str, int] = {
    "倪柝声弟兄职事": 1,
    "李常受弟兄职事第一阶段": 2,
    "李常受弟兄职事第二阶段": 2,
    "李常受弟兄职事第三阶段": 3,
    "李常受弟兄职事第四阶段": 4,
    "李常受弟兄职事高峰阶段": 5,
}

SOURCE_GROUP_TITLE_CANONICAL: dict[int, str] = {
    1: "倪柝声弟兄职事",
    2: "李常受弟兄职事第一阶段",
    3: "李常受弟兄职事第三阶段",
    4: "李常受弟兄职事第四阶段",
    5: "李常受弟兄职事高峰阶段",
}

SERIES_FOLDER_RE = re.compile(r"^(\d+)\s+(.+?)(?:——|—)\d+篇")
MSG_FILE_RE = re.compile(
    r"^msg\.\s*(\d+)\s+(.+?)"
    r"(?:[\s_]*(?:[（(]\s*)?纲目带出处(?:\s*[）)])?[\s_]*)?\.docx$",
    re.IGNORECASE,
)
SOURCE_FILE_TAG = "纲目带出处"
RED_PARA_THRESHOLD = 4

READING_MARKERS = ("读经：", "讀經：", "读经∶", "讀經∶")
MINISTRY_MARKERS = ("职事信息摘录", "職事信息摘錄")

OUTLINE_TYPE_RULES: list[tuple[str, re.Pattern[str]]] = [
    ("ot1", re.compile(r"^[壹贰貳叁參叄肆伍陆陸柒捌玖拾][\t　]")),
    ("ot2", re.compile(r"^[一二三四五六七八九十百千万萬亿億]+[\t　]")),
    ("ot3", re.compile(r"^\d+[\t　]")),
    ("ot4", re.compile(r"^[a-z][\t　]")),
    ("ot6", re.compile(r"^\([一二三四五六七八九十百千万萬亿億]+\)[\t　]")),
    ("ot5", re.compile(r"^\(\d+\)[\t　]")),
    ("ot7", re.compile(r"^\([a-z]\)[\t　]")),
]

INDEX_MAPPING = {
    "mappings": {
        "properties": {
            "series_no": {"type": "integer"},
            "series_title": {"type": "keyword"},
            "source_group_no": {"type": "integer"},
            "source_group_title": {"type": "keyword"},
            "article_no": {"type": "integer"},
            "title": {"type": "keyword"},
            "metadata": {"type": "text"},
            "outline": {
                "type": "nested",
                "properties": {
                    "type": {"type": "keyword"},
                    "text": {"type": "text"},
                    "source": {"type": "keyword"},
                },
            },
            "ministry_excerpt": {
                "type": "nested",
                "properties": {"text": {"type": "text"}},
            },
        }
    }
}


def arabic_to_chinese(n: int) -> str:
    if not 1 <= n <= 999:
        raise ValueError(f"article_no 超出 1-999 范围: {n}")
    digits = "零一二三四五六七八九"
    if n < 10:
        return digits[n]
    if n < 20:
        return "十" + (digits[n % 10] if n % 10 else "")
    if n < 100:
        tens, ones = divmod(n, 10)
        text = digits[tens] + "十"
        if ones:
            text += digits[ones]
        return text
    hundreds, rem = divmod(n, 100)
    text = digits[hundreds] + "百"
    if rem == 0:
        return text
    if rem < 10:
        return text + "零" + digits[rem]
    if rem < 20:
        return text + "十" + (digits[rem % 10] if rem % 10 else "")
    tens, ones = divmod(rem, 10)
    text += digits[tens] + "十"
    if ones:
        text += digits[ones]
    return text


def parse_series_folder(name: str) -> dict[str, Any] | None:
    m = SERIES_FOLDER_RE.match(name.strip())
    if not m:
        return None
    return {"no": int(m.group(1)), "title": m.group(2).strip()}


def normalize_for_match(name: str) -> str:
    return name.replace("弟兄的职事", "弟兄职事")


def match_source_group(folder_name: str) -> dict[str, Any] | None:
    normalized = normalize_for_match(folder_name)
    best_len = 0
    best: dict[str, Any] | None = None
    for title, no in sorted(SOURCE_GROUP_MAP.items(), key=lambda x: len(x[0]), reverse=True):
        if title in normalized and len(title) > best_len:
            best_len = len(title)
            best = {"no": no, "title": title}
    if best:
        best["title"] = SOURCE_GROUP_TITLE_CANONICAL[best["no"]]
    return best


def parse_msg_file(name: str) -> dict[str, Any] | None:
    m = MSG_FILE_RE.match(name.strip())
    if not m:
        return None
    article_no = int(m.group(1))
    topic = m.group(2).strip()
    return {
        "article_no": article_no,
        "title": f"第{arabic_to_chinese(article_no)}篇　{topic}",
    }


def _has_red_font(path: Path) -> bool:
    try:
        doc = Document(str(path))
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                try:
                    rgb = run.font.color.rgb
                    if rgb is not None and str(rgb) == "FF0000":
                        count += 1
                        break
                except Exception:
                    pass
            if count >= RED_PARA_THRESHOLD:
                return True
        return False
    except Exception:
        return False


def is_source_file(path: Path) -> bool:
    if SOURCE_FILE_TAG in path.name:
        return True
    return _has_red_font(path)


def should_skip_file(path: Path) -> bool:
    return not is_source_file(path)


def _contains_marker(text: str, markers: tuple[str, ...]) -> bool:
    return any(marker in text for marker in markers)


def _outline_type(text: str) -> str:
    stripped = text.strip()
    if _contains_marker(stripped, READING_MARKERS):
        return "bible_reading"
    for type_name, pattern in OUTLINE_TYPE_RULES:
        if pattern.match(stripped):
            return type_name
    return "ot2"


def _split_para_by_color(para) -> tuple[str, str | None]:
    text_parts: list[str] = []
    source_parts: list[str] = []
    for run in para.runs:
        content = run.text
        if not content:
            continue
        try:
            rgb = run.font.color.rgb
            is_red = rgb is not None and str(rgb) == "FF0000"
        except Exception:
            is_red = False
        if is_red:
            source_parts.append(content)
        else:
            text_parts.append(content)
    text = "".join(text_parts).strip()
    source = "".join(source_parts).strip() or None

    # fallback：若无红字出处，尝试从行末括号提取
    if not source:
        m = re.search(r"[（(]([^（(）)]{4,})[）)]$", text.strip())
        if m:
            source = m.group(1).strip()
            text = text[: m.start()].strip()

    return text, source


def parse_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    para_objs = doc.paragraphs
    paragraphs = [p.text.strip() for p in para_objs]

    reading_idx: int | None = None
    ministry_idx: int | None = None
    for idx, text in enumerate(paragraphs):
        if not text:
            continue
        if reading_idx is None and _contains_marker(text, READING_MARKERS):
            reading_idx = idx
        if ministry_idx is None and _contains_marker(text, MINISTRY_MARKERS):
            ministry_idx = idx

    metadata: list[str] = []
    outline: list[dict[str, Any]] = []

    if reading_idx is None:
        metadata = [t for t in paragraphs if t]
    else:
        metadata = [t for t in paragraphs[:reading_idx] if t]
        outline_end = ministry_idx if ministry_idx is not None else len(para_objs)
        for para in para_objs[reading_idx:outline_end]:
            raw = para.text.strip()
            if not raw:
                continue
            text, source = _split_para_by_color(para)
            if not text:
                continue
            outline.append({
                "type": _outline_type(text),
                "text": text,
                "source": source,
            })

    return {
        "metadata": metadata,
        "outline": outline,
        "ministry_excerpt": [],
    }


def iter_docx_leaf_dirs(group_dir: Path) -> list[Path]:
    leaves: list[Path] = []
    if any(group_dir.glob("*.docx")):
        leaves.append(group_dir)
    for child in natsorted([p for p in group_dir.iterdir() if p.is_dir()], key=lambda p: p.name):
        leaves.extend(iter_docx_leaf_dirs(child))
    return leaves


def collect_group_docx(group_dir: Path) -> list[Path]:
    files: list[Path] = []
    for leaf in natsorted(iter_docx_leaf_dirs(group_dir), key=lambda p: str(p)):
        for docx_path in natsorted(leaf.glob("*.docx"), key=lambda p: p.name):
            if should_skip_file(docx_path):
                continue
            if parse_msg_file(docx_path.name) is None:
                continue
            files.append(docx_path)
    return natsorted(files, key=lambda p: str(p.relative_to(group_dir)))


def iter_articles(root: Path) -> Iterator[dict[str, Any]]:
    if not root.is_dir():
        raise FileNotFoundError(f"source-dir 不存在: {root}")

    for series_dir in natsorted([p for p in root.iterdir() if p.is_dir()], key=lambda p: p.name):
        series = parse_series_folder(series_dir.name)
        if not series:
            continue

        group_dirs = natsorted(
            [p for p in series_dir.iterdir() if p.is_dir()],
            key=lambda p: p.name,
        )
        branches: dict[int, tuple[dict[str, Any], list[Path]]] = {}
        for group_dir in group_dirs:
            group = match_source_group(group_dir.name)
            if not group:
                continue
            if group["no"] not in branches:
                branches[group["no"]] = (group, [])
            branches[group["no"]][1].append(group_dir)

        for group_no in natsorted(branches.keys()):
            group, dirs = branches[group_no]
            docx_files: list[Path] = []
            for group_dir in natsorted(dirs, key=lambda p: p.name):
                docx_files.extend(collect_group_docx(group_dir))
            docx_files = natsorted(docx_files, key=lambda p: str(p))

            for seq_no, docx_path in enumerate(docx_files, start=1):
                msg = parse_msg_file(docx_path.name)
                if not msg:
                    continue
                parsed = parse_docx(docx_path)
                doc_id = f"{ES_INDEX}-{series['no']}-{group['no']}-{seq_no}"
                yield {
                    "_id": doc_id,
                    "series_no": series["no"],
                    "series_title": series["title"],
                    "source_group_no": group["no"],
                    "source_group_title": group["title"],
                    "article_no": msg["article_no"],
                    "title": msg["title"],
                    **parsed,
                }


def ensure_index(recreate: bool = False) -> None:
    if recreate and es.indices.exists(index=ES_INDEX):
        es.indices.delete(index=ES_INDEX)
        log(f"已删除索引: {ES_INDEX}")
    if not es.indices.exists(index=ES_INDEX):
        es.indices.create(index=ES_INDEX, body=INDEX_MAPPING)
        log(f"已创建索引: {ES_INDEX}")
    else:
        log(f"索引已存在: {ES_INDEX}")


def bulk_index(docs: list[dict[str, Any]]) -> None:
    if not docs:
        return
    operations: list[dict[str, Any]] = []
    for doc in docs:
        payload = dict(doc)
        doc_id = payload.pop("_id")
        operations.append({"index": {"_index": ES_INDEX, "_id": doc_id}})
        operations.append(payload)
    es.bulk(operations=operations, refresh=False)


def log(msg: str) -> None:
    print(msg)
    if _log_fp is not None:
        _log_fp.write(msg + "\n")
        _log_fp.flush()


def open_log(source_dir: Path | None = None) -> None:
    global _log_fp
    _log_fp = LOG_FILE.open("w", encoding="utf-8")
    log(f"# ingest_pano {datetime.now().isoformat(timespec='seconds')}")
    if source_dir is not None:
        log(f"# source-dir: {source_dir}")
    log(f"# LOG_FILE: {LOG_FILE}")
    log("")


def close_log() -> None:
    global _log_fp
    if _log_fp is not None:
        _log_fp.close()
        _log_fp = None


def print_series_summary(series_counts: dict[int, tuple[str, int]]) -> None:
    log("── 系列篇数统计 ──")
    total = 0
    for series_no in natsorted(series_counts.keys()):
        title, count = series_counts[series_no]
        log(f"  系列 {series_no:>2}　{title}：{count} 篇")
        total += count
    log(f"  合计：{total} 篇")
    log("")


def reindex_doc(source_dir: Path, doc_id: str) -> dict[str, Any]:
    target: dict[str, Any] | None = None
    for doc in iter_articles(source_dir):
        if doc["_id"] == doc_id:
            target = doc
            break
    if not target:
        raise ValueError(f"未找到文档: {doc_id}")
    bulk_index([target])
    es.indices.refresh(index=ES_INDEX)
    return target


def export_to_json(source_dir: Path, out_path: Path) -> int:
    log(f"开始导出 → {out_path}")
    documents: list[dict[str, Any]] = []
    series_counts: dict[int, tuple[str, int]] = {}

    for doc in iter_articles(source_dir):
        log(f"✓ {doc['_id']} {doc['title']}")
        documents.append(doc)
        sno = doc["series_no"]
        if sno not in series_counts:
            series_counts[sno] = (doc["series_title"], 0)
        title, cnt = series_counts[sno]
        series_counts[sno] = (title, cnt + 1)

    payload = {
        "index": ES_INDEX,
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "root_dir": str(source_dir),
        "mapping": INDEX_MAPPING,
        "count": len(documents),
        "documents": documents,
    }
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as fp:
        json.dump(payload, fp, ensure_ascii=False, separators=(",", ":"))

    log("")
    print_series_summary(series_counts)
    log(f"已导出 {len(documents)} 篇 → {out_path}")
    log(f"文件大小: {out_path.stat().st_size / 1024 / 1024:.2f} MB")
    return len(documents)


def import_from_json(json_path: Path, recreate: bool = True) -> int:
    if not json_path.is_file():
        raise FileNotFoundError(f"JSON 不存在: {json_path}")

    log(f"读取 {json_path} …")
    with json_path.open("r", encoding="utf-8") as fp:
        payload = json.load(fp)

    documents = payload.get("documents") or []
    if not documents:
        raise ValueError("JSON 中 documents 为空")

    index_name = payload.get("index") or ES_INDEX
    if index_name != ES_INDEX:
        log(f"警告: JSON index={index_name}，将写入 {ES_INDEX}")

    mapping = payload.get("mapping") or INDEX_MAPPING
    if recreate:
        if es.indices.exists(index=ES_INDEX):
            es.indices.delete(index=ES_INDEX)
            log(f"已删除索引: {ES_INDEX}")
        es.indices.create(index=ES_INDEX, body=mapping)
        log(f"已创建索引: {ES_INDEX}")
    elif not es.indices.exists(index=ES_INDEX):
        es.indices.create(index=ES_INDEX, body=mapping)
        log(f"已创建索引: {ES_INDEX}")
    else:
        log(f"索引已存在: {ES_INDEX}（--no-recreate，追加/覆盖同 _id）")

    total = 0
    batch: list[dict[str, Any]] = []
    for doc in documents:
        if "_id" not in doc:
            raise ValueError(f"文档缺少 _id: {doc}")
        batch.append(doc)
        total += 1
        if len(batch) >= BULK_CHUNK:
            bulk_index(batch)
            batch.clear()
            log(f"  已导入 {total} 篇…")

    if batch:
        bulk_index(batch)

    es.indices.refresh(index=ES_INDEX)
    log(f"导入完成，共 {total} 篇（JSON 记录 count={payload.get('count')}）")
    return total


def run_ingest(source_dir: Path, *, recreate: bool = True) -> None:
    global _root_dir
    _root_dir = source_dir
    open_log(source_dir)
    try:
        ensure_index(recreate=recreate)
        total = 0
        batch: list[dict[str, Any]] = []
        series_counts: dict[int, tuple[str, int]] = {}

        for doc in iter_articles(source_dir):
            log(f"✓ {doc['_id']} {doc['title']}")
            batch.append(doc)
            total += 1
            sno = doc["series_no"]
            if sno not in series_counts:
                series_counts[sno] = (doc["series_title"], 0)
            title, cnt = series_counts[sno]
            series_counts[sno] = (title, cnt + 1)
            if len(batch) >= BULK_CHUNK:
                bulk_index(batch)
                batch.clear()

        if batch:
            bulk_index(batch)

        es.indices.refresh(index=ES_INDEX)
        log("")
        print_series_summary(series_counts)
        log(f"总计入库 {total} 篇")
        log(f"日志已写入: {LOG_FILE}")
    finally:
        close_log()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="进展75系列：纲目带出处 Word 入库 / JSON 导出 / JSON 导入"
    )
    parser.add_argument(
        "--source-dir",
        type=Path,
        help="Word 纲目根目录（含系列子文件夹）",
    )
    parser.add_argument(
        "--export",
        nargs="?",
        const=str(EXPORT_JSON),
        metavar="FILE",
        help=f"导出 JSON（默认 {EXPORT_JSON.name}），不需连接 ES",
    )
    parser.add_argument("--import-json", metavar="FILE", help="从 JSON 导入 ES")
    parser.add_argument(
        "--no-recreate",
        action="store_true",
        help="全量入库/导入时不删除已有索引",
    )
    parser.add_argument(
        "--doc-id",
        help="仅重新入库指定 _id，如 progress_pano-1-1-1（需 --source-dir）",
    )
    args = parser.parse_args()

    if args.export is not None:
        if not args.source_dir:
            parser.error("--export 需要同时指定 --source-dir")
        out = Path(args.export)
        open_log(args.source_dir)
        try:
            export_to_json(args.source_dir, out)
            log(f"日志: {LOG_FILE}")
        finally:
            close_log()
        return

    if args.import_json:
        open_log()
        try:
            import_from_json(Path(args.import_json), recreate=not args.no_recreate)
            log(f"日志: {LOG_FILE}")
        finally:
            close_log()
        return

    if args.doc_id:
        if not args.source_dir:
            parser.error("--doc-id 需要同时指定 --source-dir")
        open_log(args.source_dir)
        try:
            ensure_index(recreate=False)
            doc = reindex_doc(args.source_dir, args.doc_id)
            log(
                f"✓ 已重新入库 {args.doc_id} "
                f"outline={len(doc.get('outline', []))} 条"
            )
        finally:
            close_log()
        return

    if not args.source_dir:
        parser.error("请指定 --source-dir，或使用 --import-json / --help")
    run_ingest(args.source_dir, recreate=not args.no_recreate)


if __name__ == "__main__":
    main()
