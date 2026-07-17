# -*- coding: utf-8 -*-
"""导出主题目录筛选对照 Word：列出对照基准篇题，黄底高亮终稿筛除项。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

from build_topic_catalogs import (  # noqa: E402
    CATALOG,
    INTERMEDIATE_TARGET_TOTAL,
    MIN_PER_STAGE,
    STAGE_COUNT,
    STAGE_LABELS,
    STAGE_ZH,
    TARGET_TOTAL,
    TIGHT_FILTER_TOPICS,
    TWO_STAGE_FILTER_TOPICS,
    TOPICS,
    _group_by_book,
    apply_topic_filters,
    is_match,
    parse_catalog,
    to_zh_num,
)

OUT_DIR = _SCRIPT_DIR / "topic_catalogs"


def _collect_hits(text: str, topic: str) -> dict[int, list[tuple[str, str, str]]]:
    hits: dict[int, list[tuple[str, str, str]]] = {
        s: [] for s in range(1, STAGE_COUNT + 1)
    }
    for stage, volume, book, title in parse_catalog(text):
        if is_match(topic, title, book):
            hits[stage].append((volume, book, title))
    return hits


def _kept_set(
    kept: dict[int, list[tuple[str, str, str]]],
) -> set[tuple[int, str, str, str]]:
    return {
        (stage, volume, book, title)
        for stage, items in kept.items()
        for volume, book, title in items
    }


def _add_run(paragraph, text: str, *, highlight: bool = False, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_center_title(doc: Document, text: str, *, size_pt: int = 16, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)


def write_review_docx(
    topic: str,
    initial_hits: dict[int, list[tuple[str, str, str]]],
    review_hits: dict[int, list[tuple[str, str, str]]],
    kept_hits: dict[int, list[tuple[str, str, str]]],
    *,
    output_path: Path | None = None,
) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = output_path or (OUT_DIR / f"{topic}_筛选对照.docx")

    two_stage = topic in TWO_STAGE_FILTER_TOPICS
    single_stage = topic in TIGHT_FILTER_TOPICS and not two_stage
    has_filter_review = two_stage or single_stage

    total_initial = sum(len(v) for v in initial_hits.values())
    total_review = sum(len(v) for v in review_hits.values())
    total_kept = sum(len(v) for v in kept_hits.values())
    total_removed = total_review - total_kept
    kept_lookup = _kept_set(kept_hits)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    _add_center_title(doc, "主恢复中神圣启示的进展", size_pt=18)
    _add_center_title(doc, f"{topic}——筛选对照", size_pt=15)

    note = doc.add_paragraph()
    if two_stage:
        _add_run(
            note,
            f"初筛 {total_initial} 篇；第一阶段保留 {total_review} 篇（目标 ≤{INTERMEDIATE_TARGET_TOTAL} 篇）；"
            f"终稿保留 {total_kept} 篇（目标约 {TARGET_TOTAL} 篇，每阶段至少 {MIN_PER_STAGE} 篇）；"
            f"第二阶段筛除 {total_removed} 篇。",
        )
        note2 = doc.add_paragraph()
        _add_run(
            note2,
            "本文列出第一阶段结果；黄底高亮为第二阶段筛除篇题，仍保留供核对。",
        )
    elif single_stage:
        _add_run(
            note,
            f"初筛 {total_review} 篇；终稿保留 {total_kept} 篇；筛除 {total_removed} 篇。"
            f"（目标约 {TARGET_TOTAL} 篇，每阶段至少 {MIN_PER_STAGE} 篇）",
        )
        note2 = doc.add_paragraph()
        _add_run(note2, "黄底高亮为筛除篇题，仍保留在目录中供核对。")
    else:
        _add_run(note, f"共 {total_review} 篇；未做相关性筛选，初筛结果即终稿。")

    if topic == "召会生活":
        remark = doc.add_paragraph()
        _add_run(remark, "备注：「召会生活」与「教会生活」「聚会生活」「聚会的生活」等同视之。")

    doc.add_paragraph()

    for sn in range(1, STAGE_COUNT + 1):
        items = review_hits[sn]
        if not items:
            continue

        stage_review = len(items)
        stage_kept = len(kept_hits[sn])
        stage_removed = stage_review - stage_kept
        stage_zh = STAGE_ZH[sn]
        stage_label = STAGE_LABELS[sn]

        if has_filter_review:
            stage_title = (
                f"{stage_zh}　{stage_label}"
                f"——对照 {stage_review} 篇，终稿 {stage_kept} 篇，筛除 {stage_removed} 篇"
            )
        else:
            stage_title = f"{stage_zh}　{stage_label}——{stage_review} 篇"

        _add_heading(doc, stage_title, level=2)

        book_groups = _group_by_book(items)
        for book_idx, (volume, book, titles) in enumerate(book_groups, start=1):
            book_kept = sum(
                1 for t in titles if (sn, volume, book, t) in kept_lookup
            )
            book_removed = len(titles) - book_kept
            book_zh = to_zh_num(book_idx)

            if has_filter_review:
                book_line = (
                    f"{book_zh}　{book}——对照 {len(titles)} 篇"
                    f"（终稿 {book_kept}，筛除 {book_removed}）（{volume}）"
                )
            else:
                book_line = f"{book_zh}　{book}——{len(titles)} 篇（{volume}）"

            p_book = doc.add_paragraph()
            _add_run(p_book, book_line, bold=True)

            for title in titles:
                removed = has_filter_review and (sn, volume, book, title) not in kept_lookup
                p_art = doc.add_paragraph(style="List Number")
                prefix = "【筛除】 " if removed else ""
                _add_run(p_art, f"{prefix}{title}", highlight=removed)

            doc.add_paragraph()

    doc.save(path)
    return path


def main() -> None:
    text = CATALOG.read_text(encoding="utf-8")

    for topic in TOPICS:
        initial_hits = _collect_hits(text, topic)
        kept_hits, review_hits, _ = apply_topic_filters(topic, initial_hits)
        try:
            path = write_review_docx(topic, initial_hits, review_hits, kept_hits)
        except PermissionError:
            alt = OUT_DIR / f"{topic}_筛选对照_新.docx"
            path = write_review_docx(
                topic, initial_hits, review_hits, kept_hits, output_path=alt
            )
        total_initial = sum(len(v) for v in initial_hits.values())
        total_review = sum(len(v) for v in review_hits.values())
        total_kept = sum(len(v) for v in kept_hits.values())
        if topic in TWO_STAGE_FILTER_TOPICS:
            print(
                f"{topic}: {total_initial} -> stage1 {total_review} -> final {total_kept}, "
                f"wrote {path.name}"
            )
        else:
            print(f"{topic}: {total_initial} -> {total_kept}, wrote {path.name}")


if __name__ == "__main__":
    main()
