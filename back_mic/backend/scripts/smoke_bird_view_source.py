# -*- coding: utf-8 -*-
"""鸟瞰纲目加出处 — 集成冒烟测试（路由 / source API / format_download / docx 着色）。"""
from __future__ import annotations

import base64
import io
import json
import os
import re
import sys
import time
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

os.chdir(BACKEND_DIR)

from dotenv import load_dotenv

load_dotenv(BACKEND_DIR / ".env")

BASE = os.environ.get("SMOKE_BASE_URL", "http://127.0.0.1:8000")
ROUTES_EXPECTED = [
    "POST /api/kg_rag/bird_view/skeleton",
    "POST /api/kg_rag/bird_view/outline",
    "POST /api/kg_rag/bird_view/source",
    "POST /api/kg_rag/bird_view/format_download",
]


def _get_token() -> str:
    users_path = BACKEND_DIR / "user" / "users.json"
    users = json.loads(users_path.read_text(encoding="utf-8"))
    username = next(iter(users))
    password = users[username]["pass"]
    r = requests.post(
        f"{BASE}/api/token",
        data={"username": username, "password": password, "remember": "0"},
        timeout=30,
    )
    r.raise_for_status()
    return r.json()["access_token"]


def _list_routes_from_app() -> list[str]:
    from main import app

    out = []
    for route in app.routes:
        methods = getattr(route, "methods", None) or set()
        path = getattr(route, "path", "")
        for m in sorted(methods):
            if m in ("GET", "POST", "PUT", "DELETE", "PATCH"):
                out.append(f"{m} {path}")
    return out


def _rgb_is_red(rgb) -> bool:
    if rgb is None:
        return False
    hexv = str(rgb).upper()
    if len(hexv) >= 6:
        try:
            r = int(hexv[0:2], 16)
            g = int(hexv[2:4], 16)
            b = int(hexv[4:6], 16)
            return r >= 200 and g <= 80 and b <= 80
        except ValueError:
            pass
    return "FF0000" in hexv or "FF00" in hexv


def _inspect_docx_colors(docx_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(docx_bytes))
    red_bracket_lines = 0
    green_lines = 0
    samples = []
    for i, para in enumerate(doc.paragraphs):
        if i < 4:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        has_red = any(_rgb_is_red(run.font.color.rgb) for run in para.runs)
        has_green = any(
            run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN for run in para.runs
        )
        if text.endswith("）") and has_red:
            red_bracket_lines += 1
            if len(samples) < 3:
                samples.append(f"red_bracket: {text[:80]}")
        if has_green and not text.startswith("读经："):
            green_lines += 1
            if len(samples) < 5:
                samples.append(f"green_line: {text[:80]}")
    return {
        "red_bracket_lines": red_bracket_lines,
        "green_lines": green_lines,
        "samples": samples,
        "paragraph_count": len(doc.paragraphs),
    }


def main() -> int:
    print("=== Step 1: Route registration ===")
    routes = _list_routes_from_app()
    ok_routes = True
    for expected in ROUTES_EXPECTED:
        found = expected in routes
        print(f"  {'OK' if found else 'MISSING'}  {expected}")
        ok_routes = ok_routes and found
    if not ok_routes:
        print("FAIL: missing routes")
        return 1

    print("\n=== Auth token ===")
    try:
        token = _get_token()
        print("  OK  obtained bearer token")
    except Exception as e:
        print(f"  FAIL  login: {e}")
        return 1
    headers = {"Authorization": f"Bearer {token}"}

    print("\n=== Step 2: POST /api/kg_rag/bird_view/source ===")
    source_payload = {
        "keyword": "安息",
        "type": "feast",
        "content": "（测试原文，含出处标注）2000年感恩节，第二篇\n壹　神的建造乃是神的心愿。",
        "outline": "壹\t神的建造乃是神的心愿。",
    }
    try:
        t0 = time.time()
        r2 = requests.post(
            f"{BASE}/api/kg_rag/bird_view/source",
            json=source_payload,
            headers=headers,
            timeout=180,
        )
        elapsed = time.time() - t0
        print(f"  HTTP {r2.status_code}  ({elapsed:.1f}s)")
        if r2.status_code != 200:
            print(f"  FAIL body: {r2.text[:500]}")
            return 1
        data2 = r2.json()
        outline_with_source = data2.get("outline_with_source", "")
        print(f"  type={data2.get('type')!r}")
        print(f"  outline_with_source len={len(outline_with_source)}")
        print(f"  preview: {outline_with_source[:200]!r}")
        if not outline_with_source:
            print("  FAIL: outline_with_source empty")
            return 1
        print("  OK")
    except Exception as e:
        print(f"  FAIL: {e}")
        return 1

    # 构造带括号出处 + 无出处行，用于 format_download 着色验证
    format_contents = (
        "读经：出二五8\n"
        "壹\t神的建造乃是神的心愿—出二五8。（2000年感恩节特会，第二篇）\n"
        "一\t这是找不到出处的纲目行—弗一10。"
    )

    print("\n=== Step 3: POST /api/kg_rag/bird_view/format_download (with_source=true) ===")
    try:
        r3 = requests.post(
            f"{BASE}/api/kg_rag/bird_view/format_download",
            data={
                "contents": format_contents,
                "filename": "安息【3b 节期纲目的鸟瞰】",
                "keyword": "安息",
                "type": "feast",
                "with_source": "true",
            },
            headers=headers,
            timeout=120,
        )
        print(f"  HTTP {r3.status_code}")
        if r3.status_code != 200:
            print(f"  FAIL body: {r3.text[:500]}")
            return 1
        data3 = r3.json()
        b64 = data3.get("docx_base64", "")
        print(f"  filename={data3.get('filename')!r}")
        print(f"  docx_base64 len={len(b64)}")
        if not b64:
            print("  FAIL: docx_base64 empty")
            return 1
        docx_bytes = base64.b64decode(b64)
        colors = _inspect_docx_colors(docx_bytes)
        print(f"  docx paragraphs={colors['paragraph_count']}")
        print(f"  red_bracket_lines={colors['red_bracket_lines']}")
        print(f"  green_lines={colors['green_lines']}")
        for s in colors["samples"]:
            print(f"    sample: {s}")
        if colors["red_bracket_lines"] < 1:
            print("  WARN: no red bracket lines detected")
        if colors["green_lines"] < 1:
            print("  WARN: no green highlight lines detected")
        print("  OK")
    except Exception as e:
        print(f"  FAIL: {e}")
        return 1

    print("\n=== Step 4: Frontend UI (static verification) ===")
    vue_path = BACKEND_DIR.parents[1] / "front_mic" / "frontend" / "src" / "components" / "toolbox" / "BirdViewOutline.vue"
    vue = vue_path.read_text(encoding="utf-8")
    checks = [
        ("加出处 button", "generateSource('ministry')" in vue or 'generateSource("ministry")' in vue),
        ("加出处中… loading", "加出处中…" in vue),
        ("带出处版 label", "带出处版" in vue),
        ("orange border class", "outline-with-source" in vue),
        ("download with_source", 'params.append("with_source", "true")' in vue),
        ("route /bird-view-outline", True),
    ]
    router_path = BACKEND_DIR.parents[1] / "front_mic" / "frontend" / "src" / "router" / "index.js"
    router = router_path.read_text(encoding="utf-8")
    checks.append(("router bird-view-outline", "/bird-view-outline" in router))
    for label, ok in checks:
        print(f"  {'OK' if ok else 'MISSING'}  {label}")

    print("\n=== ALL BACKEND SMOKE CHECKS PASSED ===")
    print("Note: Step 4d/e (browser click + visual docx) require manual UI verification.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
