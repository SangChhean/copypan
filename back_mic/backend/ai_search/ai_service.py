"""
AI搜索服务 - 核心业务逻辑
负责Elasticsearch检索、Claude API调用、结果处理
"""
import os
import json
import hashlib
import logging
import time
import warnings
from typing import Any, Callable, Dict, List, Optional, Tuple
from datetime import datetime
from io import BytesIO
import re
import threading
from concurrent.futures import ThreadPoolExecutor

# 抑制「Elasticsearch built-in security features are not enabled」的警告（本地开发常见）
try:
    from elasticsearch.exceptions import ElasticsearchWarning
    warnings.filterwarnings("ignore", category=ElasticsearchWarning)
except Exception:
    pass
warnings.filterwarnings("ignore", message=".*security features are not enabled.*")

from es_config import es
import anthropic
from dotenv import load_dotenv
from pathlib import Path

from .monitoring import get_monitoring


def _call_claude_messages_sync(client, system_prompt: str, user_prompt: str, max_tokens: int = 1024) -> Tuple[str, Any]:
    """同步调用 Claude messages API，返回 (首条 content 的 text, usage 或 None)。"""
    if not client:
        return ("", None)
    msg = client.messages.create(
        model=os.getenv("CLAUDE_MODEL", "claude-sonnet-4-6"),
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )
    text = ""
    if msg.content and getattr(msg.content[0], "text", None):
        text = msg.content[0].text
    return (text, getattr(msg, "usage", None))


# 配置日志（必须在导入格式刷之前，因为导入失败时会使用 logger）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("ai_search")

# 加载环境变量（确保从 backend 目录加载 .env）
env_path = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(dotenv_path=env_path)

# 导入格式刷函数（从 backend 目录导入）
try:
    import sys
    backend_dir = Path(__file__).resolve().parent.parent
    logger.debug(f"尝试导入格式刷模块，backend_dir: {backend_dir}")
    if str(backend_dir) not in sys.path:
        sys.path.insert(0, str(backend_dir))
    from format_chinese_outline import format_chinese_outline_docx
    from format_english_outline import format_english_outline_docx, format_plain_docx
    logger.info("格式刷模块导入成功")
except ImportError as e:
    format_chinese_outline_docx = None
    format_english_outline_docx = None
    format_plain_docx = None
    logger.warning(f"格式刷模块未找到，格式化功能将不可用: {e}", exc_info=True)
except Exception as e:
    format_chinese_outline_docx = None
    format_english_outline_docx = None
    format_plain_docx = None
    logger.error(f"格式刷模块导入时发生错误: {e}", exc_info=True)

# 环境变量配置
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
REDIS_HOST = os.getenv("REDIS_HOST", "localhost")
REDIS_PORT = int(os.getenv("REDIS_PORT", "6379"))

# Redis 可选：未安装或连接失败时仅禁用缓存
redis_client = None
try:
    import redis
    redis_client = redis.Redis(host=REDIS_HOST, port=REDIS_PORT, decode_responses=True)
    redis_client.ping()
    logger.info("Redis 连接成功，缓存已启用")
except Exception as e:
    logger.warning(f"Redis 未启用，将跳过缓存: {e}")

# Claude 客户端（需配置 CLAUDE_API_KEY）
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-6")
try:
    claude_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY) if CLAUDE_API_KEY else None
    if claude_client:
        logger.info("Claude 客户端初始化成功")
except Exception as e:
    logger.error(f"Claude 客户端初始化失败: {e}")
    claude_client = None

# Gemini 客户端（用于中文纲目→英文纲目翻译，需配置 GEMINI_API_KEY；使用新 SDK google.genai）
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# 纲目翻译：默认 gemini-2.5-pro 保证可用；若 3.1 不可用(404) 会自动用 GEMINI_TRANSLATION_FALLBACK_MODEL
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-pro")
# 与主模型不同才能在主模型 404/空响应/过载时真正切换到备用（若与 GEMINI_MODEL 相同则代码会跳过备用分支）
GEMINI_TRANSLATION_FALLBACK_MODEL = os.getenv("GEMINI_TRANSLATION_FALLBACK_MODEL", "gemini-2.5-flash")
# 毛胚纲目：首次 gemini-3.1-pro-preview，失败重试 gemini-2.5-pro（官方预览 ID，见 ai.google.dev/models/gemini-3.1-pro-preview）
ROUGH_OUTLINE_GEMINI_MODEL = os.getenv("ROUGH_OUTLINE_GEMINI_MODEL", "gemini-3.1-pro-preview")
ROUGH_OUTLINE_GEMINI_FALLBACK_MODEL = os.getenv("ROUGH_OUTLINE_GEMINI_FALLBACK_MODEL", "gemini-2.5-pro")
# 纲目翻译等其它场景：主模型重试全失败后尝试的备用模型
GEMINI_FALLBACK_MODEL = os.getenv("GEMINI_FALLBACK_MODEL", "gemini-2.5-pro")
gemini_client = None
_gemini_system_instruction_en2zh = None
if GEMINI_API_KEY:
    try:
        from google import genai
        from google.genai import types
        from .gemini_translation_instruction import GEMINI_TRANSLATION_SYSTEM_INSTRUCTION
        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        _gemini_system_instruction = GEMINI_TRANSLATION_SYSTEM_INSTRUCTION
        try:
            from .gemini_translation_instruction_en2zh import GEMINI_TRANSLATION_SYSTEM_INSTRUCTION_EN2ZH
            _gemini_system_instruction_en2zh = GEMINI_TRANSLATION_SYSTEM_INSTRUCTION_EN2ZH
        except Exception:
            _gemini_system_instruction_en2zh = None
        logger.info("Gemini 翻译模型初始化成功")
    except Exception as e:
        logger.error(f"Gemini 翻译模型初始化失败: {e}")
        gemini_client = None
        _gemini_system_instruction = None
        _gemini_system_instruction_en2zh = None
else:
    _gemini_system_instruction = None
    logger.info("Gemini 未配置: GEMINI_API_KEY 未设置（.env 路径: %s）", env_path)

try:
    from .gemini_response_utils import (
        extract_translatable_text,
        gemini_translation_generate_config,
        translation_max_output_tokens,
    )
except ImportError:
    gemini_translation_generate_config = None
    extract_translatable_text = None
    translation_max_output_tokens = None  # type: ignore[misc, assignment]

if gemini_client and translation_max_output_tokens:
    try:
        logger.info(
            "纲目翻译输出 token 上限（GEMINI_TRANSLATION_MAX_OUTPUT_TOKENS）: %s",
            translation_max_output_tokens(),
        )
    except Exception:
        pass

# 毛胚纲目 - 其他 API（.env 中填写对应 API Key 后启用）
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-v3.2")
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
PERPLEXITY_MODEL = os.getenv("PERPLEXITY_MODEL", "sonar-pro")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4")
XAI_API_KEY = os.getenv("XAI_API_KEY")  # Grok
XAI_MODEL = os.getenv("XAI_MODEL", "grok-4-1-fast-reasoning")

# 毛胚纲目 OpenAI 兼容接口费用（美元/百万 token），按 provider 分别计算
ROUGH_OUTLINE_PRICES = {
    "chatgpt": {"in": 2.50, "out": 15.00},
    "grok": {"in": 0.20, "out": 0.50},
    "deepseek": {"in": 0.27, "out": 1.10},
    "perplexity": {"in": 3.00, "out": 15.00},  # sonar-pro
}

# 并发限制：同时进行中的 Claude / Gemini 请求数，避免人多时触发 API 限流（429）
def _parse_concurrent_limit(env_key: str, default: int) -> int:
    try:
        v = int(os.getenv(env_key, str(default)))
        return max(1, v)  # 至少为 1，避免 Semaphore(0) 导致永久阻塞
    except (ValueError, TypeError):
        return default

# Claude Tier 1 约 50 RPM，Gemini 免费/低阶约 5–15 RPM；20/10 在 Tier 2 与付费 Tier 1 下安全，可通过 .env 覆盖
CLAUDE_CONCURRENT_LIMIT = _parse_concurrent_limit("CLAUDE_CONCURRENT_LIMIT", 20)
GEMINI_CONCURRENT_LIMIT = _parse_concurrent_limit("GEMINI_CONCURRENT_LIMIT", 10)
CLAUDE_SEMAPHORE = threading.Semaphore(CLAUDE_CONCURRENT_LIMIT)
GEMINI_SEMAPHORE = threading.Semaphore(GEMINI_CONCURRENT_LIMIT)
logger.info("API 并发限制: Claude=%s, Gemini=%s", CLAUDE_CONCURRENT_LIMIT, GEMINI_CONCURRENT_LIMIT)


def _claude_error_is_retryable(exc: BaseException) -> bool:
    """429/529/5xx 临时错误与连接/超时，可重试。"""
    if isinstance(exc, anthropic.APIConnectionError):
        return True
    if isinstance(exc, anthropic.APIStatusError):
        return exc.status_code in (429, 500, 502, 503, 504, 529)
    return False


def _call_claude_with_retry(
    prompt: str,
    max_tokens: int = 8192,
    max_retries: int = 2,
    backoff_seconds: Tuple[int, ...] = (5, 10),
) -> Any:
    """
    节期纲目等：调用 Claude messages.create；遇可重试错误最多再试 max_retries 次（共 1+max_retries 次）。
    CLAUDE_SEMAPHORE 仅在单次 HTTP 请求期间持有，退避等待期间释放。
    """
    if not claude_client:
        raise RuntimeError("Claude 客户端未初始化")
    for attempt in range(max_retries + 1):
        try:
            with CLAUDE_SEMAPHORE:
                return claude_client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=max_tokens,
                    messages=[{"role": "user", "content": prompt}],
                )
        except Exception as e:
            if not _claude_error_is_retryable(e) or attempt >= max_retries:
                raise
            wait = backoff_seconds[min(attempt, len(backoff_seconds) - 1)] if backoff_seconds else 5
            status = getattr(e, "status_code", None) if isinstance(e, anthropic.APIStatusError) else None
            logger.warning(
                "Claude 调用失败（可重试）attempt=%s/%s status=%s: %s，%s 秒后重试",
                attempt + 1,
                max_retries + 1,
                status,
                e,
                wait,
            )
            time.sleep(wait)


def _claude_message_text(message: Any) -> str:
    """
    从 Messages API 返回的 Message 中拼接所有 text 类型块。
    避免仅用 content[0]：若存在 thinking / tool_use 等块时首块可能无 .text，会 AttributeError 或取错内容。
    """
    if not message or not getattr(message, "content", None):
        return ""
    parts: List[str] = []
    block_types: List[str] = []
    for block in message.content:
        block_types.append(str(getattr(block, "type", type(block).__name__)))
        btype = getattr(block, "type", None)
        if btype == "text":
            t = getattr(block, "text", None)
            if isinstance(t, str) and t.strip():
                parts.append(t)
        # 其它类型（thinking、tool_use 等）跳过
    out = "\n".join(parts).strip()
    if not out and block_types:
        logger.warning(
            "Claude 响应无可用 text 块，块类型=%s，stop_reason=%s",
            block_types,
            getattr(message, "stop_reason", None),
        )
    return out


# 纲目翻译时与原文一起发送的 prompt（【需要翻译的文章】+ 以下说明）
OUTLINE_TRANSLATE_PROMPT_ZH2EN = (
    "请将文章翻译为英文，严格使用System instructions中的专用术语表进行翻译。"
    "格式要求：①中文序号为壹，翻译为英文I.，一翻译为A.，二翻译为B.，1翻译为1.，a翻译为a.，(一)翻译为1)，以此类推；②不要缩进，直接输出。"
)
OUTLINE_TRANSLATE_PROMPT_EN2ZH = (
    "请将文章翻译为中文，严格使用System instructions中的专用术语表进行翻译。"
    "格式要求：①读经格式为缩写，例如：罗一1；②英文序号为I.，翻译为中文壹，A.翻译为一，B.翻译为二，1.翻译为1，a.翻译为a，1)翻译为(一)，以此类推；注意，纲目层级之后只加空格，不加其他符号，如：壹 神爱世人，为世人舍了自己的独生子—约三16：；③不要缩进，直接输出。"
)


def _user_facing_translate_error(api_errors: List[str], empty_body: bool) -> str:
    """将 Gemini 异常摘要为前端可展示的中文（不泄露密钥）。"""
    if empty_body and not api_errors:
        return "模型未返回正文，可尝试缩短内容或稍后重试"
    blob = " ".join(api_errors)
    bl = blob.lower()
    if any(
        x in bl
        for x in (
            "api key",
            "api_key",
            "unauthorized",
            "401",
            "invalid authentication",
            "permission_denied",
            "permission denied",
        )
    ):
        return "Gemini API 密钥无效或未授权，请检查 .env 中的 GEMINI_API_KEY"
    if any(
        x in bl
        for x in (
            "429",
            "resource_exhausted",
            "quota",
            "rate limit",
            "too many requests",
            "exhausted",
        )
    ):
        return "调用频率或配额已达上限，请稍后再试或提高 Gemini 配额"
    if any(x in bl for x in ("404", "not_found", "is not found", "not supported for generatecontent")):
        return "当前 GEMINI_MODEL 不可用，请在 .env 中更换模型或调整 GEMINI_TRANSLATION_FALLBACK_MODEL"
    if any(
        x in bl
        for x in (
            "503",
            "unavailable",
            "500",
            "internal",
            "deadline_exceeded",
            "timeout",
            "temporary",
            "connection",
            "reset",
        )
    ):
        return "翻译服务暂时不稳定，请稍后重试"
    if any(x in bl for x in ("safety", "blocked", "recitation", "prohibited", "harmful")):
        return "内容触发安全策略未返回译文，请删减或改写后重试"
    if "max_tokens" in bl or "max token" in bl or "output_token" in bl or "length limit" in bl:
        return "译文长度达到模型输出上限，请缩短原文或分段翻译，或提高 GEMINI_TRANSLATION_MAX_OUTPUT_TOKENS"
    if api_errors:
        tail = (api_errors[-1] or "")[:280]
        return f"翻译失败：{tail}"
    return "纲目内容翻译失败，请稍后重试"


def _gemini_error_is_retryable(error_msg: str) -> bool:
    em = error_msg.lower()
    return (
        "503" in error_msg
        or "unavailable" in em
        or "429" in error_msg
        or "timeout" in em
        or "temporary" in em
        or "deadline_exceeded" in em
        or "500" in error_msg
        or "internal" in em
        or "resource_exhausted" in em
    )


# 索引配置：索引名 -> 权重（用于每索引取数及排序加权）
# 按纲目性质（special_needs）选择不同权重
INDEXES_CONFIG_BY_NATURE = {
    "一般性": {
        "map_note": {"weight": 1.0},
        "map_dictionary": {"weight": 1.0},
        "map_7feasts": {"weight": 1.0},
        "map_pano": {"weight": 1.0},
        "filewall": {"weight": 1.2},
        "cwwl": {"weight": 1.0},
        "cwwn": {"weight": 1.0},
        "life": {"weight": 1.0},
        "bib": {"weight": 1.0},
        "others": {"weight": 1.0},
    },
    "高真理浓度": {
        "map_note": {"weight": 1.0},
        "map_dictionary": {"weight": 1.0},
        "map_7feasts": {"weight": 1.0},
        "map_pano": {"weight": 1.0},
        "filewall": {"weight": 1.2},
        "cwwl": {"weight": 1.0},  # 94-97 额外 1.5
        "cwwn": {"weight": 1.0},
        "life": {"weight": 1.0},
        "bib": {"weight": 1.0},
        "others": {"weight": 1.0},
    },
    "高生命浓度": {
        "map_note": {"weight": 1.0},
        "map_dictionary": {"weight": 1.0},
        "map_7feasts": {"weight": 1.0},
        "map_pano": {"weight": 1.0},
        "filewall": {"weight": 1.2},
        "cwwl": {"weight": 1.0},
        "cwwn": {"weight": 1.5},
        "life": {"weight": 1.5},
        "bib": {"weight": 1.0},
        "others": {"weight": 1.0},
    },
    "重实行应用": {
        "map_note": {"weight": 1.0},
        "map_dictionary": {"weight": 1.0},
        "map_7feasts": {"weight": 1.0},
        "map_pano": {"weight": 1.0},
        "filewall": {"weight": 1.2},
        "cwwl": {"weight": 1.0},  # 85-93 额外 1.5
        "cwwn": {"weight": 1.0},
        "life": {"weight": 1.0},
        "bib": {"weight": 1.0},
        "others": {"weight": 1.0},
    },
}

# 索引中文名，供后台统计页展示权重
INDEX_LABELS = {
    "map_note": "注解",
    "map_dictionary": "词典",
    "map_7feasts": "节期",
    "map_pano": "上河图",
    "filewall": "防火墙",
    "cwwl": "李常受文集",
    "cwwn": "倪柝声文集",
    "life": "生命读经",
    "bib": "圣经",
    "others": "其他",
}

# 信息检索使用的索引（仅此 8 个，不含 bib）
INFO_RETRIEVAL_INDEXES = (
    "map_note",
    "map_dictionary",
    "map_7feasts",
    "map_pano",
    "cwwl",
    "cwwn",
    "life",
    "others",
)


def get_index_weights_for_display():
    """供后台统计页展示：各纲目性质对应的 AI 检索索引权重。"""
    out = {}
    for nature, config in INDEXES_CONFIG_BY_NATURE.items():
        out[nature] = {idx: config[idx]["weight"] for idx in config}
    out["_notes"] = {
        "一般性": "cwwl 1994-1997 文集 ×1.1",
        "高真理浓度": "cwwl 1994-1997 文集 ×1.5",
        "重实行应用": "cwwl 1985-1993 文集 ×1.5",
    }
    out["_labels"] = INDEX_LABELS
    return out


def _strip_code_fence_for_outline(text: Optional[str]) -> Optional[str]:
    """
    提取文本中最后一个完整的 fenced code block 内容。
    若存在多个代码块（如 Claude 先输出思考过程再输出纲目），取最后一个。
    若文本中没有任何代码块，原样返回。
    """
    if not text or not isinstance(text, str):
        return text
    s = text.strip()
    last_start = -1
    idx = 0
    lines = s.split("\n")
    fence_starts = []
    fence_ends = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("```"):
            if not fence_starts or len(fence_starts) == len(fence_ends):
                fence_starts.append(i)
            elif len(fence_starts) > len(fence_ends):
                fence_ends.append(i)
    if not fence_starts or not fence_ends:
        return text
    # 取最后一对完整的 fence
    last_end = fence_ends[-1]
    last_start = fence_starts[len(fence_ends) - 1]
    inner = lines[last_start + 1 : last_end]
    return "\n".join(inner).strip() if inner else text


class AISearchService:
    """AI智能搜索服务"""

    def __init__(self):
        self.es = es
        self.redis = redis_client
        self.claude = claude_client
        self.cache_ttl = 3600  # 缓存1小时

        logger.info("AISearchService初始化完成")

    _MAP_LIKE_INDICES = frozenset({"map_note", "map_7feasts", "map_dictionary", "map_pano", "filewall"})

    def _parse_doc_id(self, doc_id: str) -> tuple:
        """解析文档 id，提取 message 前缀和段号。如 others_1_1-4 -> (others_1_1-, 4)"""
        if not doc_id or "-" not in doc_id:
            return ("", 0)
        last_dash = doc_id.rfind("-")
        prefix = doc_id[: last_dash + 1]
        try:
            seg = int(doc_id[last_dash + 1 :])
        except ValueError:
            seg = 0
        return (prefix, seg)

    def _fetch_message_docs(self, index_name: str, message_prefix: str) -> List[Dict]:
        """从 ES 获取同一篇（message）内的所有文档，按段号排序"""
        try:
            resp = self.es.search(
                index=index_name,
                body={
                    "query": {"prefix": {"id": message_prefix}},
                    "size": 500,
                    "_source": ["id", "type", "text", "title", "book", "chapter", "verse"],
                },
                request_timeout=10,
            )
            hits = resp.get("hits", {}).get("hits", [])
            docs = []
            for h in hits:
                src = h.get("_source", {})
                pid, seg = self._parse_doc_id(src.get("id", ""))
                docs.append((seg, src))
            docs.sort(key=lambda x: x[0])
            return [d[1] for d in docs]
        except Exception as e:
            logger.warning(f"获取 message 文档失败: {e}")
            return []

    def _get_map_note_reference_from_hit(self, source: Dict, hit: Dict, index_name: str = "") -> str:
        """
        map 类索引的引用：去掉括号，只保留文本。
        - map_note：msg 项内有 source，优先从命中的 msg 项取
        - map_7feasts：用文档外层 source
        - map_dictionary：用文档外层 text
        """
        def _strip_parens(t) -> str:
            t = str(t or "").strip()
            for left, right in [("（", "）"), ("(", ")")]:
                if t.startswith(left) and t.endswith(right):
                    t = t[len(left):-len(right)].strip()
            return t

        # map_dictionary：引用 = 第一个bookname + ", " + title + ", " + 第二个bookname（从 msg 中取）
        if index_name == "map_dictionary":
            msg_list = source.get("msg") or []
            booknames = [str(m.get("text") or "").strip() for m in msg_list if (m.get("type") or "") == "bookname"]
            titles = [str(m.get("text") or "").strip() for m in msg_list if (m.get("type") or "") == "title"]
            b1 = booknames[0] if len(booknames) >= 1 else ""
            t = titles[0] if titles else ""
            b2 = booknames[1] if len(booknames) >= 2 else ""
            parts = [p for p in [b1, t, b2] if p]
            if parts:
                return "，".join(parts)
            s = _strip_parens(source.get("text") or "")
            if s:
                return s
            return _strip_parens(source.get("id") or source.get("_id") or "") or "未知来源"

        # map_pano：清明上河图，+ 外层 text
        if index_name == "map_pano":
            t = str(source.get("text") or "").strip()
            if t:
                return f"清明上河图，{t}"
            s = _strip_parens(source.get("source") or "")
            if s:
                return f"清明上河图，{s}"
            return _strip_parens(source.get("id") or source.get("_id") or "") or "清明上河图"

        # map_7feasts：msg 内无 source，用文档级 source
        if index_name == "map_7feasts":
            s = _strip_parens(source.get("source") or "")
            if s:
                return s
            return _strip_parens(source.get("id") or source.get("_id") or "") or "未知来源"

        # map_note：圣经真理题库，+ 外层 text
        if index_name == "map_note":
            t = str(source.get("text") or "").strip()
            if t:
                return f"圣经真理题库，{t}"
            return "圣经真理题库"

        # 其他 map 类回退
        inner = hit.get("inner_hits", {}).get("matched_msg", {})
        msg_list = source.get("msg") or []
        for ih in inner.get("hits", {}).get("hits", []):
            s = _strip_parens(ih.get("_source", {}).get("source") or "")
            if s:
                return s
            offset = (ih.get("_nested") or {}).get("offset")
            if isinstance(offset, int) and 0 <= offset < len(msg_list):
                s = _strip_parens(msg_list[offset].get("source") or "")
                if s:
                    return s
        s = _strip_parens(source.get("source") or "")
        if s:
            return s
        return _strip_parens(source.get("id") or source.get("_id") or "") or "未知来源"

    def _format_reference(self, source: Dict) -> str:
        """格式化经文引用"""
        # 尝试多种可能的字段名
        book = source.get('book') or source.get('title') or source.get('bookname') or ''
        chapter = source.get('chapter') or source.get('chap') or ''
        verse = source.get('verse') or source.get('vs') or ''
        
        # 尝试获取其他标识信息
        volume = source.get('volume') or source.get('vol') or ''
        page = source.get('page') or source.get('pg') or ''
        section = source.get('section') or source.get('sec') or ''
        
        # 优先使用书卷+章节+节
        if book and chapter and verse:
            return f"{book} {chapter}:{verse}"
        elif book and chapter:
            return f"{book} {chapter}"
        elif book:
            return book
        
        # 尝试使用卷+页
        if volume and page:
            return f"卷{volume} 第{page}页"
        elif volume:
            return f"卷{volume}"
        
        # 尝试使用章节信息
        if section:
            return f"第{section}节"
        
        # 最后尝试ID或其他标识
        doc_id = source.get('id') or source.get('_id') or ''
        if doc_id:
            return f"文档 {doc_id}"
        
        return "未知来源"

    def translate_outline(
        self,
        chinese_outline: str,
        outline_topic: Optional[str] = None,
        use_cache: bool = True,
    ) -> Dict:
        """
        将中文纲目翻译为英文纲目（调用 Gemini）。
        若未配置 GEMINI_API_KEY 返回 error；失败时重试 1 次。
        use_cache=True（如 AI 纲目流程）时按中文内容 hash 缓存 24 小时；use_cache=False（如工具箱 - 纲目翻译）时不缓存。
        同时翻译纲目主题作为英文标题（若提供 outline_topic）。
        """
        MAX_OUTLINE_LENGTH = 100_000
        TRANSLATE_CACHE_TTL = 86400  # 24 小时

        outline = (chinese_outline or "").strip()
        if not outline:
            return {"answer_en": None, "title_en": None, "error": "中文纲目为空"}
        if len(outline) > MAX_OUTLINE_LENGTH:
            return {"answer_en": None, "title_en": None, "error": f"中文纲目过长（最多 {MAX_OUTLINE_LENGTH} 字）"}

        cache_key = f"ai_search:translate:{hashlib.sha256(outline.encode()).hexdigest()[:32]}"
        if use_cache and self.redis:
            try:
                cached = self.redis.get(cache_key)
                if cached:
                    data = json.loads(cached)
                    cached_answer = data.get("answer_en")
                    cached_title = data.get("title_en")
                    if cached_answer:
                        if outline_topic and outline_topic.strip() and not cached_title:
                            topic = outline_topic.strip()
                            _cache_title_404 = [False]
                            _cache_title_retryable = [False]
                            _cache_title_empty = [False]

                            def _title_cfg_cache():
                                if gemini_translation_generate_config:
                                    return gemini_translation_generate_config(_gemini_system_instruction)
                                return types.GenerateContentConfig(system_instruction=_gemini_system_instruction)

                            def _translate_title_for_cache(
                                retry_count: int = 0, model: Optional[str] = None
                            ) -> Optional[str]:
                                use_model = model or GEMINI_MODEL
                                with GEMINI_SEMAPHORE:
                                    try:
                                        title_response = gemini_client.models.generate_content(
                                            model=use_model,
                                            contents=topic,
                                            config=_title_cfg_cache(),
                                        )
                                        log_tc = f"[Gemini翻译-缓存标题] model={use_model}"
                                        if extract_translatable_text:
                                            raw_title = extract_translatable_text(title_response, log_tc)
                                        else:
                                            rt = getattr(title_response, "text", None) if title_response else None
                                            raw_title = rt.strip() if isinstance(rt, str) and rt.strip() else None
                                        if raw_title:
                                            title_en_clean = raw_title
                                            prefixes_to_remove = [
                                                "Translation:", "English:", "翻译：", "英文：",
                                                "The translation is:", "Here is the translation:",
                                                "Title:", "标题："
                                            ]
                                            for prefix in prefixes_to_remove:
                                                if title_en_clean.lower().startswith(prefix.lower()):
                                                    title_en_clean = title_en_clean[len(prefix):].strip()
                                            title_en_clean = title_en_clean.strip('"\'')
                                            try:
                                                usage_meta = getattr(title_response, "usage_metadata", None)
                                                in_tok = int(getattr(usage_meta, "prompt_token_count", 0) or 0)
                                                out_tok = int(getattr(usage_meta, "candidates_token_count", 0) or 0)
                                                cost = (in_tok * 1.25 + out_tok * 10) / 1_000_000
                                                logger.info(
                                                    "[Gemini翻译-缓存标题] model=%s | 输入=%d tokens | 输出=%d tokens | 费用=$%.6f",
                                                    use_model, in_tok, out_tok, cost,
                                                )
                                                if self.redis:
                                                    get_monitoring(self.redis).record_tool_usage(
                                                        "translation_cache_title", use_model, in_tok, out_tok, cost
                                                    )
                                            except Exception:
                                                pass
                                            return title_en_clean
                                        _cache_title_empty[0] = True
                                        logger.warning("缓存标题翻译返回空响应（重试次数: %s）", retry_count)
                                    except Exception as e:
                                        error_msg = str(e)
                                        if "404" in error_msg or "NOT_FOUND" in error_msg or "is not found" in error_msg.lower():
                                            _cache_title_404[0] = True
                                        is_retryable = _gemini_error_is_retryable(error_msg)
                                        if is_retryable:
                                            _cache_title_retryable[0] = True
                                        if is_retryable and retry_count == 0:
                                            logger.warning("缓存标题翻译调用失败（可重试）: %s，等待2秒后重试...", e)
                                            time.sleep(2)
                                        else:
                                            logger.warning("缓存标题翻译调用失败（重试次数: %s）: %s", retry_count, e)
                                return None

                            cached_title = _translate_title_for_cache(retry_count=0)
                            if cached_title is None:
                                cached_title = _translate_title_for_cache(retry_count=1)
                            if cached_title is None and (
                                _cache_title_404[0] or _cache_title_retryable[0] or _cache_title_empty[0]
                            ) and GEMINI_TRANSLATION_FALLBACK_MODEL != GEMINI_MODEL:
                                cached_title = _translate_title_for_cache(
                                    retry_count=0, model=GEMINI_TRANSLATION_FALLBACK_MODEL
                                )
                                if cached_title is None:
                                    cached_title = _translate_title_for_cache(
                                        retry_count=1, model=GEMINI_TRANSLATION_FALLBACK_MODEL
                                    )
                            if cached_title:
                                logger.info("缓存标题翻译成功: '%s' -> '%s'", topic, cached_title)
                            else:
                                logger.warning("缓存标题翻译失败（已重试1次）: '%s'", topic)
                        return {"answer_en": cached_answer, "title_en": cached_title}
            except Exception as e:
                logger.debug("翻译缓存读取失败: %s", e)

        if not gemini_client:
            return {"answer_en": None, "error": "英文翻译服务未配置（请设置 GEMINI_API_KEY）"}

        # 发送内容 = 需要翻译的文章 + 格式与术语说明（中翻英）
        contents_zh2en = outline + "\n\n" + OUTLINE_TRANSLATE_PROMPT_ZH2EN
        _last_error_model_not_found = [False]  # 404 / model not found
        _last_error_retryable = [False]         # 503 / 429 / timeout 等临时性失败
        _last_error_empty = [False]             # 调用成功但无可翻译正文
        zh2en_api_errors: List[str] = []

        def _is_model_not_found(err: str) -> bool:
            return "404" in err or "NOT_FOUND" in err or "is not found" in err.lower()

        def _zh2en_config():
            if gemini_translation_generate_config:
                return gemini_translation_generate_config(_gemini_system_instruction)
            return types.GenerateContentConfig(system_instruction=_gemini_system_instruction)

        def _call_gemini(retry_count: int = 0, model: Optional[str] = None) -> Optional[tuple]:
            use_model = model or GEMINI_MODEL
            with GEMINI_SEMAPHORE:
                try:
                    response = gemini_client.models.generate_content(
                        model=use_model,
                        contents=contents_zh2en,
                        config=_zh2en_config(),
                    )
                    log_p = f"[Gemini翻译] model={use_model}"
                    if extract_translatable_text:
                        text = extract_translatable_text(response, log_p)
                    else:
                        rt = getattr(response, "text", None) if response else None
                        text = rt.strip() if isinstance(rt, str) and rt.strip() else None
                    if text:
                        tokens_zh2en = None
                        try:
                            usage_meta = response.usage_metadata
                            in_tok = int(getattr(usage_meta, "prompt_token_count", 0) or 0)
                            out_tok = int(getattr(usage_meta, "candidates_token_count", 0) or 0)
                            cost = (in_tok * 1.25 + out_tok * 10) / 1_000_000
                            logger.info(f"[Gemini翻译] model={use_model} | 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                            if self.redis:
                                get_monitoring(self.redis).record_tool_usage("translation_zh2en", use_model, in_tok, out_tok, cost)
                            tokens_zh2en = {"input": in_tok, "output": out_tok, "cost": cost}
                        except Exception:
                            pass
                        return (text, tokens_zh2en)
                    _last_error_empty[0] = True
                    logger.warning("Gemini 翻译返回空响应（重试次数: %s）", retry_count)
                except Exception as e:
                    error_msg = str(e)
                    zh2en_api_errors.append(error_msg)
                    if _is_model_not_found(error_msg):
                        _last_error_model_not_found[0] = True
                        logger.warning(f"Gemini 模型不可用(404): {e}，将尝试备用模型 {GEMINI_TRANSLATION_FALLBACK_MODEL}")
                    is_retryable = _gemini_error_is_retryable(error_msg)
                    if is_retryable:
                        _last_error_retryable[0] = True
                    if is_retryable and retry_count == 0:
                        logger.warning(f"Gemini 翻译调用失败（可重试）: {e}，等待2秒后重试...")
                        time.sleep(2)
                    else:
                        logger.warning(f"Gemini 翻译调用失败（重试次数: {retry_count}）: {e}")
            return None

        result = _call_gemini(retry_count=0)
        if result is not None:
            answer_en, tokens_zh2en = result[0], result[1]
        else:
            answer_en, tokens_zh2en = None, None
        if answer_en is None:
            result = _call_gemini(retry_count=1)
            if result is not None:
                answer_en, tokens_zh2en = result[0], result[1]
        if answer_en is None and (
            _last_error_model_not_found[0] or _last_error_retryable[0] or _last_error_empty[0]
        ) and GEMINI_TRANSLATION_FALLBACK_MODEL != GEMINI_MODEL:
            _fb_reason = (
                "模型不存在"
                if _last_error_model_not_found[0]
                else "主模型返回空正文"
                if _last_error_empty[0]
                else "主模型负载过高/暂不可用"
            )
            logger.info("使用备用模型进行中翻英: %s（原因: %s）", GEMINI_TRANSLATION_FALLBACK_MODEL, _fb_reason)
            result = _call_gemini(retry_count=0, model=GEMINI_TRANSLATION_FALLBACK_MODEL)
            if result is not None:
                answer_en, tokens_zh2en = result[0], result[1]
            if answer_en is None:
                result = _call_gemini(retry_count=1, model=GEMINI_TRANSLATION_FALLBACK_MODEL)
                if result is not None:
                    answer_en, tokens_zh2en = result[0], result[1]
        
        # 翻译标题（如果提供）- 独立执行，即使纲目内容翻译失败也尝试翻译标题
        title_en = None
        if outline_topic and outline_topic.strip():
            topic = outline_topic.strip()
            _title_model_not_found = [False]
            _title_retryable = [False]
            _title_empty = [False]

            def _title_cfg_main():
                if gemini_translation_generate_config:
                    return gemini_translation_generate_config(_gemini_system_instruction)
                return types.GenerateContentConfig(system_instruction=_gemini_system_instruction)

            def _translate_title(retry_count: int = 0, model: Optional[str] = None) -> Optional[str]:
                """翻译标题，带重试逻辑；model 为空则用 GEMINI_MODEL"""
                use_model = model or GEMINI_MODEL
                with GEMINI_SEMAPHORE:
                    try:
                        title_response = gemini_client.models.generate_content(
                            model=use_model,
                            contents=topic,
                            config=_title_cfg_main(),
                        )
                        log_tm = f"[Gemini标题] model={use_model}"
                        if extract_translatable_text:
                            raw_title = extract_translatable_text(title_response, log_tm)
                        else:
                            rt = getattr(title_response, "text", None) if title_response else None
                            raw_title = rt.strip() if isinstance(rt, str) and rt.strip() else None
                        if raw_title:
                            title_en_clean = raw_title
                            prefixes_to_remove = [
                                "Translation:", "English:", "翻译：", "英文：",
                                "The translation is:", "Here is the translation:",
                                "Title:", "标题："
                            ]
                            for prefix in prefixes_to_remove:
                                if title_en_clean.lower().startswith(prefix.lower()):
                                    title_en_clean = title_en_clean[len(prefix):].strip()
                            title_en_clean = title_en_clean.strip('"\'')
                            return title_en_clean
                        _title_empty[0] = True
                        logger.warning("标题翻译返回空响应（重试次数: %s）", retry_count)
                    except Exception as e:
                        error_msg = str(e)
                        if _is_model_not_found(error_msg):
                            _title_model_not_found[0] = True
                        is_retryable = _gemini_error_is_retryable(error_msg)
                        if is_retryable:
                            _title_retryable[0] = True
                        if is_retryable and retry_count == 0:
                            logger.warning(f"标题翻译调用失败（可重试）: {e}，等待2秒后重试...")
                            time.sleep(2)
                        else:
                            logger.warning(f"标题翻译调用失败（重试次数: {retry_count}）: {e}")
                return None

            title_en = _translate_title(retry_count=0)
            if title_en is None:
                title_en = _translate_title(retry_count=1)
            if title_en is None and (
                _title_model_not_found[0] or _title_retryable[0] or _title_empty[0]
            ) and GEMINI_TRANSLATION_FALLBACK_MODEL != GEMINI_MODEL:
                title_en = _translate_title(retry_count=0, model=GEMINI_TRANSLATION_FALLBACK_MODEL)
                if title_en is None:
                    title_en = _translate_title(retry_count=1, model=GEMINI_TRANSLATION_FALLBACK_MODEL)

            if title_en:
                logger.info(f"标题翻译成功: '{topic}' -> '{title_en}'")
            else:
                logger.warning(f"标题翻译失败（已重试1次）: '{topic}'")
        
        # 如果纲目内容翻译失败，返回错误（但标题翻译结果仍会返回）
        if answer_en is None:
            return {
                "answer_en": None,
                "title_en": title_en,
                "error": _user_facing_translate_error(zh2en_api_errors, _last_error_empty[0]),
            }

        if use_cache and self.redis:
            try:
                cache_data = {"answer_en": answer_en}
                if title_en:
                    cache_data["title_en"] = title_en
                self.redis.setex(
                    cache_key,
                    TRANSLATE_CACHE_TTL,
                    json.dumps(cache_data, ensure_ascii=False),
                )
            except Exception as e:
                logger.debug("翻译缓存写入失败: %s", e)

        return {"answer_en": answer_en, "title_en": title_en, "tokens": tokens_zh2en or {"input": 0, "output": 0, "cost": 0}}

    def translate_outline_en2zh(self, english_outline: str) -> Dict:
        """
        将英文纲目翻译为中文纲目（调用 Gemini，使用英翻中 instruction）。
        用于工具箱「纲目翻译」- 英翻中。失败时重试 1 次。
        """
        MAX_OUTLINE_LENGTH = 100_000
        outline = (english_outline or "").strip()
        if not outline:
            return {"answer_zh": None, "error": "英文纲目为空"}
        if len(outline) > MAX_OUTLINE_LENGTH:
            return {"answer_zh": None, "error": f"英文纲目过长（最多 {MAX_OUTLINE_LENGTH} 字）"}

        if not gemini_client:
            return {"answer_zh": None, "error": "英文翻译服务未配置（请设置 GEMINI_API_KEY）"}
        if not _gemini_system_instruction_en2zh:
            return {"answer_zh": None, "error": "英翻中 instruction 未配置"}

        # 发送内容 = 需要翻译的文章 + 格式与术语说明（英翻中）
        contents_en2zh = outline + "\n\n" + OUTLINE_TRANSLATE_PROMPT_EN2ZH
        _last_error_model_not_found = [False]  # 404 / model not found
        _last_error_retryable = [False]         # 503 / 429 / timeout 等临时性失败
        _last_error_empty = [False]             # 调用成功但无可翻译正文
        en2zh_api_errors: List[str] = []

        def _is_model_not_found(err: str) -> bool:
            return "404" in err or "NOT_FOUND" in err or "is not found" in err.lower()

        def _en2zh_config():
            if gemini_translation_generate_config:
                return gemini_translation_generate_config(_gemini_system_instruction_en2zh)
            return types.GenerateContentConfig(system_instruction=_gemini_system_instruction_en2zh)

        def _call_gemini(retry_count: int = 0, model: Optional[str] = None) -> Optional[tuple]:
            use_model = model or GEMINI_MODEL
            with GEMINI_SEMAPHORE:
                try:
                    response = gemini_client.models.generate_content(
                        model=use_model,
                        contents=contents_en2zh,
                        config=_en2zh_config(),
                    )
                    log_p = f"[Gemini英翻中] model={use_model}"
                    if extract_translatable_text:
                        text = extract_translatable_text(response, log_p)
                    else:
                        rt = getattr(response, "text", None) if response else None
                        text = rt.strip() if isinstance(rt, str) and rt.strip() else None
                    if text:
                        tokens_en2zh = None
                        try:
                            usage_meta = response.usage_metadata
                            in_tok = int(getattr(usage_meta, "prompt_token_count", 0) or 0)
                            out_tok = int(getattr(usage_meta, "candidates_token_count", 0) or 0)
                            cost = (in_tok * 1.25 + out_tok * 10) / 1_000_000
                            logger.info(f"[Gemini英翻中] model={use_model} | 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                            if self.redis:
                                get_monitoring(self.redis).record_tool_usage("translation_en2zh", use_model, in_tok, out_tok, cost)
                            tokens_en2zh = {"input": in_tok, "output": out_tok, "cost": cost}
                        except Exception:
                            pass
                        return (text, tokens_en2zh)
                    _last_error_empty[0] = True
                    logger.warning("Gemini 英翻中返回空响应（重试次数: %s）", retry_count)
                except Exception as e:
                    error_msg = str(e)
                    en2zh_api_errors.append(error_msg)
                    if _is_model_not_found(error_msg):
                        _last_error_model_not_found[0] = True
                        logger.warning(
                            "Gemini 模型不可用(404): %s，将尝试备用模型 %s",
                            e,
                            GEMINI_TRANSLATION_FALLBACK_MODEL,
                        )
                    is_retryable = _gemini_error_is_retryable(error_msg)
                    if is_retryable:
                        _last_error_retryable[0] = True
                    if is_retryable and retry_count == 0:
                        logger.warning("Gemini 英翻中调用失败（可重试）: %s，等待2秒后重试...", e)
                        time.sleep(2)
                    else:
                        logger.warning("Gemini 英翻中调用失败（重试次数: %s）: %s", retry_count, e)
            return None

        result = _call_gemini(retry_count=0)
        if result is not None:
            answer_zh, tokens_en2zh = result[0], result[1]
        else:
            answer_zh, tokens_en2zh = None, None
        if answer_zh is None:
            result = _call_gemini(retry_count=1)
            if result is not None:
                answer_zh, tokens_en2zh = result[0], result[1]
        if answer_zh is None and (
            _last_error_model_not_found[0] or _last_error_retryable[0] or _last_error_empty[0]
        ) and GEMINI_TRANSLATION_FALLBACK_MODEL != GEMINI_MODEL:
            _fb_reason_e2z = (
                "模型不存在"
                if _last_error_model_not_found[0]
                else "主模型返回空正文"
                if _last_error_empty[0]
                else "主模型负载过高/暂不可用"
            )
            logger.info("使用备用模型进行英翻中: %s（原因: %s）", GEMINI_TRANSLATION_FALLBACK_MODEL, _fb_reason_e2z)
            result = _call_gemini(retry_count=0, model=GEMINI_TRANSLATION_FALLBACK_MODEL)
            if result is not None:
                answer_zh, tokens_en2zh = result[0], result[1]
            if answer_zh is None:
                result = _call_gemini(retry_count=1, model=GEMINI_TRANSLATION_FALLBACK_MODEL)
                if result is not None:
                    answer_zh, tokens_en2zh = result[0], result[1]
        if answer_zh is None:
            return {
                "answer_zh": None,
                "error": _user_facing_translate_error(en2zh_api_errors, _last_error_empty[0]),
            }
        return {"answer_zh": answer_zh, "tokens": tokens_en2zh or {"input": 0, "output": 0, "cost": 0}}

    def outline_to_traditional(self, content: str) -> Dict[str, Optional[str]]:
        """
        将简体纲目转为繁体：先按术语表替换，再 OpenCC s2t（失败回退 zhconv zh-hant）。
        用于 AI 纲目制作「同时生成繁体纲目」。
        
        Args:
            content: 简体中文纲目全文
        
        Returns:
            {"answer_zh_tw": str} 成功时；{"answer_zh_tw": None, "error": str} 失败时
        """
        if not (content or "").strip():
            return {"answer_zh_tw": None, "error": "内容为空"}
        text = (content or "").strip()
        try:
            # 1. 加载台湾繁简术语表（简体 -> 繁体）
            terms_path = Path(__file__).resolve().parents[3] / "shared" / "zh_tw_terms.json"
            placeholders: List[tuple] = []  # (placeholder_str, target_value)
            if terms_path.exists():
                terms = json.loads(terms_path.read_text(encoding="utf-8"))
                # 按键长降序，先替换长词避免短词截断
                sorted_keys = sorted(terms.keys(), key=len, reverse=True)
                for idx, simp in enumerate(sorted_keys):
                    trad = terms[simp]
                    if simp and trad is not None:
                        # 用占位符替换，避免 zhconv 把术语表结果再改掉（如「了解」→「瞭解」）
                        ph = f"__TW_{idx}__"
                        placeholders.append((ph, trad))
                        text = text.replace(simp, ph)
            else:
                logger.warning("繁简术语表不存在: %s，仅做通用简繁转换", terms_path)
            # 2. 通用简→繁体（优先 OpenCC s2t，占位符为 ASCII 不会被改动）
            try:
                from opencc import OpenCC
                cc = OpenCC("s2t")
                text = cc.convert(text)
            except Exception:
                try:
                    import zhconv
                    text = zhconv.convert(text, "zh-hant")
                except ImportError:
                    logger.warning("OpenCC/zhconv 未安装，无法做通用简繁转换")
                    return {"answer_zh_tw": None, "error": "繁简转换依赖未安装（opencc 或 zhconv）"}
            # 3. 把占位符还原为术语表里的目标用字
            for ph, trad in placeholders:
                text = text.replace(ph, trad)
            return {"answer_zh_tw": text}
        except Exception as e:
            logger.error("简转繁失败: %s", e, exc_info=True)
            return {"answer_zh_tw": None, "error": str(e)}

    def traditional_to_simplified(self, content: str) -> Dict[str, Optional[str]]:
        """
        将台湾繁体纲目转为简体：直接使用 zhconv 转换（不经过术语表）。
        用于工具箱「简繁互转」功能。
        
        Args:
            content: 台湾繁体中文纲目全文
        
        Returns:
            {"answer_zh_cn": str} 成功时；{"answer_zh_cn": None, "error": str} 失败时
        """
        if not (content or "").strip():
            return {"answer_zh_cn": None, "error": "内容为空"}
        text = (content or "").strip()
        try:
            # 繁→简：优先 OpenCC t2s（繁体→简体），否则 zhconv
            try:
                from opencc import OpenCC
                cc = OpenCC("t2s")
                text = cc.convert(text)
            except Exception:
                try:
                    import zhconv
                    text = zhconv.convert(text, "zh-cn")
                except ImportError:
                    logger.warning("OpenCC/zhconv 未安装，无法做繁简转换")
                    return {"answer_zh_cn": None, "error": "繁简转换依赖未安装（opencc 或 zhconv）"}
            return {"answer_zh_cn": text}
        except Exception as e:
            logger.error("繁转简失败: %s", e, exc_info=True)
            return {"answer_zh_cn": None, "error": str(e)}

    def _convert_docx_to_pdf(self, docx_path: str) -> Optional[bytes]:
        """
        将 DOCX 文件转换为 PDF bytes。
        优先使用 LibreOffice（Linux），回退到 docx2pdf（Windows/Mac）。
        
        Args:
            docx_path: DOCX 文件路径
        
        Returns:
            PDF bytes，失败时返回 None
        """
        import tempfile
        import os
        import subprocess
        import platform
        
        # 检查 DOCX 文件是否存在
        if not os.path.exists(docx_path):
            logger.error(f"DOCX 文件不存在: {docx_path}")
            return None
        
        # 创建临时 PDF 文件
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            pdf_path = tmp_pdf.name
        
        logger.info(f"开始转换 DOCX 到 PDF: {docx_path} -> {pdf_path}")
        
        # 方法1: 尝试使用 LibreOffice（Linux）
        try:
            # 检查 LibreOffice 是否可用
            result = subprocess.run(
                ["which", "libreoffice"],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                # 使用 LibreOffice 转换
                # 获取输出目录（PDF 文件所在目录）
                output_dir = os.path.dirname(pdf_path)
                # LibreOffice 会生成文件名，我们需要指定完整路径
                # 先转换，然后重命名
                temp_output_dir = output_dir
                temp_output_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
                
                # 使用 PDF/A-2b 导出，会嵌入全部字体，避免移动端打开乱码
                pdf_export_opts = 'pdf:writer_pdf_Export:{"SelectPdfVersion":{"type":"long","value":"2"}}'
                convert_result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to", pdf_export_opts,
                        "--outdir", temp_output_dir,
                        docx_path
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                
                if convert_result.returncode == 0:
                    # LibreOffice 会在输出目录生成 PDF，文件名基于输入文件名
                    generated_pdf = os.path.join(temp_output_dir, temp_output_name)
                    if os.path.exists(generated_pdf):
                        # 移动到目标位置
                        if generated_pdf != pdf_path:
                            os.rename(generated_pdf, pdf_path)
                        
                        # 检查 PDF 文件是否生成成功
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            # 读取 PDF bytes
                            with open(pdf_path, "rb") as f:
                                pdf_bytes = f.read()
                            
                            logger.info(f"PDF 转换成功（LibreOffice）: 大小 {len(pdf_bytes)} bytes")
                            
                            # 清理临时文件
                            try:
                                os.unlink(pdf_path)
                            except Exception:
                                pass
                            
                            return pdf_bytes
        except FileNotFoundError:
            logger.debug("LibreOffice 未找到，尝试其他方法")
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice 转换超时")
        except Exception as e:
            logger.warning(f"LibreOffice 转换失败: {e}")
        
        # 方法2: 回退到 docx2pdf（Windows/Mac，或 Linux 上如果安装了 MS Word）
        # Windows 下 win32com 要求当前线程已调用 CoInitialize，否则报「尚未调用 CoInitialize」
        import sys
        if sys.platform == "win32":
            try:
                import pythoncom
                pythoncom.CoInitialize()
            except Exception:
                pass
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            
            # 检查 PDF 文件是否生成成功
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                # 读取 PDF bytes
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                
                logger.info(f"PDF 转换成功（docx2pdf）: 大小 {len(pdf_bytes)} bytes")
                
                # 清理临时文件
                try:
                    os.unlink(pdf_path)
                except Exception:
                    pass
                
                return pdf_bytes
        except NotImplementedError as e:
            logger.error(f"DOCX 转 PDF 失败: {e}")
        except Exception as e:
            logger.error(f"DOCX 转 PDF 失败: {e}", exc_info=True)
        
        # 清理临时文件
        try:
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
        except Exception:
            pass
        
        return None

    def format_outline_only(
        self,
        direction: str,
        translated_text: str,
        output_format: str = "docx",
        is_outline: bool = True,
    ) -> Dict:
        """
        仅格式化已翻译的纲目文本（不调用翻译 API）。
        
        Args:
            direction: "zh2en" 或 "en2zh"（用于确定使用哪个模板和格式刷函数）
            translated_text: 已翻译的纲目文本
            output_format: "docx" 或 "pdf"，默认 "docx"
            is_outline: True=纲目格式刷，False=通用平铺格式刷（末尾无标点→居中加粗，其余→paragraph 样式）
        
        Returns:
            {
                "docx_bytes": bytes | None,  # DOCX bytes（output_format="docx" 时）
                "pdf_bytes": bytes | None,  # PDF bytes（output_format="pdf" 时）
                "filename": str | None,  # 建议的文件名
                "error": str | None,  # 错误信息
            }
        """
        import shutil
        import tempfile
        from docx import Document

        # 根据方向确定模板和格式刷函数
        if direction == "zh2en":
            template_name = "英文纲目模板.docx"
            format_func = format_english_outline_docx if is_outline else format_plain_docx
            default_filename = "outline_en.docx"
        elif direction == "en2zh":
            template_name = "中文纲目模板.docx"
            format_func = format_chinese_outline_docx if is_outline else format_plain_docx
            default_filename = "outline_zh.docx"
        elif direction in ("zh_cn2tw", "zh_tw2cn"):
            # 简繁转换：都使用中文模板和中文格式刷
            template_name = "中文纲目模板.docx"
            format_func = format_chinese_outline_docx if is_outline else format_plain_docx
            default_filename = "outline_zh.docx"
        else:
            return {
                "docx_bytes": None,
                "pdf_bytes": None,
                "filename": None,
                "error": f"无效的方向: {direction}",
            }

        # 检查格式刷函数是否可用
        if format_func is None:
            logger.warning("格式刷函数未导入，无法格式化")
            return {
                "docx_bytes": None,
                "pdf_bytes": None,
                "filename": default_filename,
                "error": "格式刷函数未导入",
            }

        # 复制模板并写入翻译结果
        backend_dir = Path(__file__).resolve().parent.parent
        template_path = backend_dir / template_name
        if not template_path.exists():
            logger.error(f"模板文件不存在: {template_path}")
            return {
                "docx_bytes": None,
                "pdf_bytes": None,
                "filename": default_filename,
                "error": f"模板文件不存在: {template_name}",
            }

        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                temp_docx_path = tmp_file.name

            # 复制模板
            shutil.copy2(template_path, temp_docx_path)

            # 打开文档并写入翻译结果
            doc = Document(temp_docx_path)
            
            # 策略：清空所有段落，将翻译结果按行添加为新段落
            # 这样格式刷函数可以正确识别并应用样式
            for para in list(doc.paragraphs):
                p_element = para._element
                p_element.getparent().remove(p_element)
            
            # 将翻译结果按行分割并添加为段落
            lines = translated_text.split("\n")
            for line in lines:
                if line.strip() or len(doc.paragraphs) == 0:  # 保留空行或至少有一个段落
                    doc.add_paragraph(line)

            # 保存
            doc.save(temp_docx_path)

            # 调用格式刷（中文纲目且 zh_cn2tw 时使用繁体引号 『』「」）
            try:
                if format_func is format_chinese_outline_docx:
                    format_func(temp_docx_path, traditional_quotes=(direction == "zh_cn2tw"))
                else:
                    format_func(temp_docx_path)
            except Exception as e:
                logger.error(f"格式刷失败: {e}", exc_info=True)
                # 格式刷失败时仍返回未格式化的文档
                pass

            # 读取为 bytes
            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()

            # 生成文件名（基于题目：中文纲目取读经前最后一段，英文纲目取 Scripture reading 前最后一段）
            # 回退逻辑：如果没有读经/Scripture reading，则以第一个大点位置判断题目
            filename = default_filename
            if lines:
                title_line = None
                if direction in ("en2zh", "zh_cn2tw", "zh_tw2cn"):
                    # 中文纲目：优先找"读经："所在行，取它前面一行作为题目
                    scripture_idx = None
                    for idx, line in enumerate(lines):
                        if '读经：' in line or '讀經：' in line:
                            scripture_idx = idx
                            break
                    
                    if scripture_idx is not None:
                        # 找到读经，取读经前一行
                        if scripture_idx > 0:
                            title_line = lines[scripture_idx - 1].strip()
                    else:
                        # 没找到读经，找第一个以"壹"开头的行（第一个大点），取它前面一行
                        for idx, line in enumerate(lines):
                            if line.strip().startswith('壹'):
                                if idx > 0:
                                    title_line = lines[idx - 1].strip()
                                break
                
                elif direction == "zh2en":
                    # 英文纲目：优先找"Scripture reading:"所在行，取它前面一行作为题目
                    scripture_idx = None
                    for idx, line in enumerate(lines):
                        if line.strip().lower().startswith("scripture reading:"):
                            scripture_idx = idx
                            break
                    
                    if scripture_idx is not None:
                        # 找到 Scripture reading，取它前面一行
                        if scripture_idx > 0:
                            title_line = lines[scripture_idx - 1].strip()
                    else:
                        # 没找到 Scripture reading，找第一个罗马数字开头的行（第一个大点，如 "I. "），取它前面一行
                        re_roman = re.compile(r'^([IVXL]+)\.\s')
                        for idx, line in enumerate(lines):
                            if re_roman.match(line.strip()):
                                if idx > 0:
                                    title_line = lines[idx - 1].strip()
                                break
                
                # 如果找到题目行，使用题目；否则回退到第一行
                if title_line:
                    title_text = title_line[:50]
                else:
                    title_text = lines[0].strip()[:50]
                
                if title_text:
                    safe_name = re.sub(r'[\/:*?"<>|]', '_', title_text)
                    if output_format == "pdf":
                        filename = f"{safe_name}.pdf"
                    else:
                        filename = f"{safe_name}.docx"

            # 根据输出格式处理
            if output_format == "pdf":
                # 转换为 PDF
                pdf_bytes = self._convert_docx_to_pdf(temp_docx_path)
                # 清理临时 DOCX 文件
                try:
                    os.unlink(temp_docx_path)
                except Exception:
                    pass
                if pdf_bytes:
                    return {
                        "docx_bytes": None,
                        "pdf_bytes": pdf_bytes,
                        "filename": filename,
                        "error": None,
                    }
                else:
                    # PDF 转换失败，返回 DOCX
                    return {
                        "docx_bytes": docx_bytes,
                        "pdf_bytes": None,
                        "filename": filename.replace(".pdf", ".docx"),
                        "error": "PDF 转换失败，已返回 DOCX 文件",
                    }
            else:
                # 返回 DOCX
                # 清理临时文件
                try:
                    os.unlink(temp_docx_path)
                except Exception:
                    pass
                return {
                    "docx_bytes": docx_bytes,
                    "pdf_bytes": None,
                    "filename": filename,
                    "error": None,
                }

        except Exception as e:
            logger.error(f"格式化失败: {e}", exc_info=True)
            # 清理临时文件
            try:
                if "temp_docx_path" in locals():
                    os.unlink(temp_docx_path)
            except Exception:
                pass
            return {
                "docx_bytes": None,
                "pdf_bytes": None,
                "filename": default_filename,
                "error": f"格式化失败: {str(e)}",
            }

    @staticmethod
    def _colorize_ministerialize_source_suffixes(doc, header_count: int) -> None:
        """格式刷后：将段落末尾「（出处）」设为红色。"""
        import re
        from docx.shared import RGBColor

        for idx, para in enumerate(doc.paragraphs):
            if idx < header_count:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            # 去掉末尾格式刷可能加上的。或：
            clean_text = re.sub(r"[。：]\s*$", "", text).strip()
            # 必须以）结尾才处理
            if not clean_text.endswith("）"):
                continue
            # 从右往左找与末尾）配对的（
            depth = 0
            split_pos = -1
            for i in range(len(clean_text) - 1, -1, -1):
                ch = clean_text[i]
                if ch == "）":
                    depth += 1
                elif ch == "（":
                    depth -= 1
                    if depth == 0:
                        split_pos = i
                        break
            if split_pos == -1:
                continue
            main = clean_text[:split_pos]
            suffix = clean_text[split_pos:]
            para.clear()
            para.add_run(main)
            run_src = para.add_run(suffix)
            run_src.font.color.rgb = RGBColor(0xFF, 0, 0)

    def format_rough_outline_docx(
        self,
        outline_type: str,
        contents: List[str],
        header_lines: Optional[List[str]] = None,
        content_lines: Optional[List[dict]] = None,
        with_source: bool = False,
    ) -> Dict:
        """
        毛胚纲目刷格式并下载：将润色版 4 篇或三分钟分享 6 篇合并为一个 DOCX，使用中文模板与中文刷格式。
        
        Args:
            outline_type: "polish"（润色版）或 "sharing"（三分钟分享）
            contents: 多篇纲目正文，按顺序合并（润色版 4 篇，三分钟分享 6 篇）
            header_lines: 可选前三段（系列/总题/篇题），写入正文之前
        
        Returns:
            {"docx_bytes": bytes | None, "filename": str, "error": str | None}
        """
        import shutil
        import tempfile
        import os
        from docx import Document

        _allowed = ("polish", "sharing", "beginner", "youth", "truth")
        if outline_type not in _allowed:
            return {"docx_bytes": None, "filename": "毛胚纲目.docx", "error": f"不支持的纲目类型: {outline_type}"}
        if content_lines:
            text_parts = [(item.get("text") or "").strip() for item in content_lines if (item.get("text") or "").strip()]
            if not text_parts:
                return {"docx_bytes": None, "filename": "毛胚纲目.docx", "error": "内容不能为空"}
        elif not contents:
            return {"docx_bytes": None, "filename": "毛胚纲目.docx", "error": "内容不能为空"}

        if format_chinese_outline_docx is None:
            return {"docx_bytes": None, "filename": "毛胚纲目.docx", "error": "中文格式刷未导入"}

        backend_dir = Path(__file__).resolve().parent.parent
        template_name = "中文纲目模板.docx"
        template_path = backend_dir / template_name
        if not template_path.exists():
            return {"docx_bytes": None, "filename": "毛胚纲目.docx", "error": f"模板文件不存在: {template_name}"}

        # 合并多篇：三分钟分享各 AI 版本之间多加一行空行以作区分，其余用双换行
        sep = "\n\n\n" if outline_type in ("sharing", "polish") else "\n\n"
        if content_lines:
            combined_text = sep.join(
                (item.get("text") or "").strip() for item in content_lines if (item.get("text") or "").strip()
            )
        else:
            combined_text = sep.join((c or "").strip() for c in contents if (c or "").strip())

        _filename_map = {
            "polish": "毛胚纲目_润色版.docx",
            "sharing": "毛胚纲目_三分钟分享.docx",
            "beginner": "毛胚纲目_初信版.docx",
            "youth": "毛胚纲目_青少年版.docx",
            "truth": "毛胚纲目_真理加强版.docx",
        }
        filename = _filename_map.get(outline_type, "毛胚纲目.docx")
        temp_docx_path = None

        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                temp_docx_path = tmp_file.name
            shutil.copy2(template_path, temp_docx_path)

            doc = Document(temp_docx_path)
            for para in list(doc.paragraphs):
                p_element = para._element
                p_element.getparent().remove(p_element)

            header_count = 0
            # 写入前三段（若用户有填）
            for line in header_lines or []:
                if line.strip():
                    doc.add_paragraph(line.strip())
                    header_count += 1

            if content_lines:
                for item in content_lines:
                    text = (item.get("text") or "").strip()
                    if not text:
                        continue
                    source = (item.get("source") or "").strip()
                    if with_source and source:
                        doc.add_paragraph(f"{text}（{source}）")
                    else:
                        for line in text.split("\n"):
                            if line.strip():
                                doc.add_paragraph(line)
                            elif len(doc.paragraphs) > header_count:
                                doc.add_paragraph("")
            else:
                for line in combined_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                    elif len(doc.paragraphs) > header_count:
                        doc.add_paragraph("")  # 空行保留，三分钟分享各 AI 版本之间可见空白

            doc.save(temp_docx_path)

            try:
                format_chinese_outline_docx(
                    temp_docx_path,
                    traditional_quotes=False,
                    truth_underline_between_markers=(outline_type == "truth"),
                    sharing_all_0000=(outline_type == "sharing"),
                )
            except Exception as e:
                logger.error(f"毛胚纲目格式刷失败: {e}", exc_info=True)

            if with_source and content_lines:
                doc = Document(temp_docx_path)
                self._colorize_ministerialize_source_suffixes(doc, header_count)
                doc.save(temp_docx_path)

            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()

            return {"docx_bytes": docx_bytes, "filename": filename, "error": None}
        except Exception as e:
            logger.error(f"毛胚纲目刷格式失败: {e}", exc_info=True)
            return {"docx_bytes": None, "filename": filename, "error": str(e)}
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    os.unlink(temp_docx_path)
                except Exception:
                    pass

    def format_feast_outline_docx(
        self,
        contents: List[str],
        outline_type: str = "original",
        line1: Optional[str] = None,
        line2: Optional[str] = None,
        line3: Optional[str] = None,
        morning_revival_raw: Optional[str] = None,
        transcript_raw: Optional[str] = None,
        transcript_preface: Optional[str] = None,
        transcript_addendum: Optional[str] = None,
        preface_outline: Optional[str] = None,
        addendum_outline: Optional[str] = None,
    ) -> Dict:
        """
        节期纲目刷格式并下载：将一篇或多篇纲目合并为一个 DOCX，使用节期纲目模板与节期纲目刷格式（按类型）。
        outline_type: original | with_scripture | morning_revival | transcript | composite
        Returns:
            {"docx_bytes": bytes | None, "filename": str, "error": str | None}
        """
        import shutil
        import tempfile
        import os
        from docx import Document

        if not contents:
            return {"docx_bytes": None, "filename": "节期纲目.docx", "error": "内容不能为空"}

        backend_dir = Path(__file__).resolve().parent.parent
        for template_name in ("节期纲目模板.docx", "template.docx"):
            template_path = backend_dir / template_name
            if template_path.exists():
                break
        else:
            return {"docx_bytes": None, "filename": "节期纲目.docx", "error": "节期纲目模板.docx 或 template.docx 不存在"}

        allowed = ("original", "with_scripture", "morning_revival", "transcript", "composite")
        if outline_type not in allowed:
            outline_type = "original"

        # 所有类型（含 with_scripture）均直接使用传入的 contents，不再调用经文汇集或 Claude
        main_contents = "\n\n".join((c or "").rstrip() for c in contents if (c or "").strip())
        l1, l2, l3 = (line1 or "").strip(), (line2 or "").strip(), (line3 or "").strip()
        header = "\n".join([l1, l2, l3]) if (l1 or l2 or l3) else ""

        # 听抄稿、复合稿：序言、添言。优先使用已生成的 preface_outline/addendum_outline；否则用原文请求 Claude 生成
        preface_outline = (preface_outline or "").strip()
        addendum_outline = (addendum_outline or "").strip()
        if outline_type in ("transcript", "composite") and (not preface_outline or not addendum_outline):
            preface_raw = (transcript_preface or "").strip()
            addendum_raw = (transcript_addendum or "").strip()
            if not preface_outline and preface_raw:
                res = self.feast_outline_preface(preface_raw)
                if not res.get("error"):
                    preface_outline = (res.get("outline") or "").strip()
                else:
                    logger.warning("序言 Claude 生成失败: %s", res.get("error"))
            if not addendum_outline and addendum_raw:
                res = self.feast_outline_addendum(addendum_raw)
                if not res.get("error"):
                    addendum_outline = (res.get("outline") or "").strip()
                else:
                    logger.warning("添言 Claude 生成失败: %s", res.get("error"))

        if header:
            combined_text = header + "\n\n" + main_contents
            if outline_type == "transcript" and (preface_outline or addendum_outline):
                # 听抄稿：第4段为读经（contents 第一行），第5段起为序言，再正文、添言
                lines = main_contents.split("\n")
                reading_line = (lines[0].strip() if lines else "") or ""
                rest_main = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                combined_text = header + "\n\n" + reading_line + "\n\n"
                if preface_outline:
                    combined_text += preface_outline + "\n\n"
                combined_text += rest_main
                if addendum_outline:
                    combined_text += "\n\n" + addendum_outline
            elif outline_type == "composite" and (preface_outline or addendum_outline):
                # 复合稿：第4段=读经，第5段起=序言，再正文、添言
                lines = main_contents.split("\n")
                reading_line = (lines[0].strip() if lines else "") or ""
                rest_main = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                combined_text = header + "\n\n" + reading_line + "\n\n"
                if preface_outline:
                    combined_text += preface_outline + "\n\n"
                combined_text += rest_main
                if addendum_outline:
                    combined_text += "\n\n" + addendum_outline
        else:
            combined_text = main_contents
            if outline_type == "transcript" and (preface_outline or addendum_outline):
                parts = []
                if preface_outline:
                    parts.append(preface_outline)
                parts.append(main_contents)
                if addendum_outline:
                    parts.append(addendum_outline)
                combined_text = "\n\n".join(parts)
            elif outline_type == "composite" and (preface_outline or addendum_outline):
                lines = main_contents.split("\n")
                reading_line = (lines[0].strip() if lines else "") or ""
                rest_main = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                parts = [reading_line]
                if preface_outline:
                    parts.append(preface_outline)
                parts.append(rest_main)
                if addendum_outline:
                    parts.append(addendum_outline)
                combined_text = "\n\n".join(parts)
        # 听抄稿/复合稿序言、添言段落范围：第4段=读经，第5段起=序言，再正文、添言
        preface_highlight_end = addendum_highlight_start = addendum_highlight_end = 0
        if outline_type == "transcript" and (preface_outline or addendum_outline):
            lines = main_contents.split("\n")
            rest_main = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
            n_preface = len(preface_outline.split("\n")) if preface_outline else 0
            n_main = len(rest_main.split("\n")) if rest_main else 0
            n_addendum = len(addendum_outline.split("\n")) if addendum_outline else 0
            n_header = 4  # line1, line2, line3, 读经（第一行）
            preface_highlight_end = n_header + n_preface
            addendum_highlight_start = n_header + n_preface + n_main
            addendum_highlight_end = addendum_highlight_start + n_addendum
        elif outline_type == "composite" and (preface_outline or addendum_outline):
            lines = main_contents.split("\n")
            rest_main = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
            n_preface = len(preface_outline.split("\n")) if preface_outline else 0
            n_main = len(rest_main.split("\n")) if rest_main else 0
            n_addendum = len(addendum_outline.split("\n")) if addendum_outline else 0
            n_header = 4  # line1, line2, line3, 读经（第一行）
            preface_highlight_end = n_header + n_preface
            addendum_highlight_start = n_header + n_preface + n_main
            addendum_highlight_end = addendum_highlight_start + n_addendum
        temp_docx_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                temp_docx_path = tmp_file.name
            shutil.copy2(template_path, temp_docx_path)
            doc = Document(temp_docx_path)
            for para in list(doc.paragraphs):
                p_element = para._element
                p_element.getparent().remove(p_element)
            for line in combined_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
                elif len(doc.paragraphs) > 0:
                    doc.add_paragraph("")
            # 刷格式并下载时在第三段末尾追加类型标注（纲目的原文/带经文的纲目/晨兴信息选读的纲目/听抄稿的纲目/复合的纲目）
            FEAST_OUTLINE_TYPE_LABELS = {
                "original": "（纲目的原文）",
                "with_scripture": "（带经文的纲目）",
                "morning_revival": "（晨兴信息选读的纲目）",
                "transcript": "（听抄稿的纲目）",
                "composite": "（复合的纲目）",
            }
            suffix = FEAST_OUTLINE_TYPE_LABELS.get(outline_type)
            if suffix and len(doc.paragraphs) >= 3:
                para = doc.paragraphs[2]
                existing = (para.text or "").rstrip()
                if not existing.endswith(suffix):
                    para.text = existing + suffix
            from docx.enum.text import WD_BREAK
            # 听抄稿：⑥ 纲目末尾先加分页 +「听抄信息：」+ 序言、听抄稿、添言三段合并，⑦ 套听抄信息页样式，再⑧整篇刷格式、⑨高亮
            if outline_type == "transcript" and (
                (transcript_raw or "").strip()
                or (transcript_preface or "").strip()
                or (transcript_addendum or "").strip()
            ):
                _append_transcript_info_section(
                    doc,
                    (transcript_raw or "").strip(),
                    transcript_preface=(transcript_preface or "").strip() or None,
                    transcript_addendum=(transcript_addendum or "").strip() or None,
                )
            # 晨兴信息选读：纲目刷格式后再追加「晨兴圣言信息：」页（刷格式只跑纲目部分）
            doc.save(temp_docx_path)
            try:
                import sys
                if str(backend_dir) not in sys.path:
                    sys.path.insert(0, str(backend_dir))
                from 节期纲目刷格式 import format_feast_outline_docx as apply_feast_format
                apply_feast_format(temp_docx_path, outline_type)
            except ImportError as e:
                logger.warning(f"节期纲目刷格式模块未导入: {e}，跳过格式刷")
            except Exception as e:
                logger.error(f"节期纲目格式刷失败: {e}", exc_info=True)
            if outline_type == "morning_revival" and (morning_revival_raw or "").strip():
                doc2 = Document(temp_docx_path)
                _append_morning_revival_section(doc2, (morning_revival_raw or "").strip())
                doc2.save(temp_docx_path)
            # 听抄稿：⑨【听抄稿添加开始】～【听抄稿添加结束】高亮并删标记；序言整段黄色高亮；添言＝从「添言」到分页符整段黄色高亮
            if outline_type == "transcript":
                try:
                    from docx.enum.text import WD_COLOR_INDEX
                    doc2 = Document(temp_docx_path)
                    _apply_transcript_add_highlight(doc2)
                    n_paras = len(doc2.paragraphs)
                    for idx in range(4, min(preface_highlight_end, n_paras)):
                        for run in doc2.paragraphs[idx].runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    addendum_range = _get_addendum_paragraph_range(doc2, to_page_break=True)
                    if addendum_range:
                        start_a, end_a = addendum_range
                        for idx in range(start_a, min(end_a + 1, n_paras)):
                            for run in doc2.paragraphs[idx].runs:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    doc2.save(temp_docx_path)
                except Exception as e:
                    logger.warning(f"听抄稿添加高亮处理失败: {e}", exc_info=True)
            # 复合稿：序言整段下划线；添言＝从「添言」到文档末尾整段下划线
            elif outline_type == "composite":
                try:
                    doc2 = Document(temp_docx_path)
                    n_paras = len(doc2.paragraphs)
                    if preface_highlight_end:
                        for idx in range(4, min(preface_highlight_end, n_paras)):
                            for run in doc2.paragraphs[idx].runs:
                                run.font.underline = True
                    addendum_range = _get_addendum_paragraph_range(doc2, to_page_break=False)
                    if addendum_range:
                        start_a, end_a = addendum_range
                        for idx in range(start_a, min(end_a + 1, n_paras)):
                            for run in doc2.paragraphs[idx].runs:
                                run.font.underline = True
                    doc2.save(temp_docx_path)
                except Exception as e:
                    logger.warning(f"复合稿序言/添言下划线处理失败: {e}", exc_info=True)
            # 根据第三段与类型生成下载文件名「【类型】序号 内容.docx」，失败则用默认名
            download_filename = "节期纲目.docx"
            try:
                from 节期纲目刷格式 import suggest_feast_outline_filename
                doc_for_name = Document(temp_docx_path)
                if len(doc_for_name.paragraphs) >= 3:
                    third_text = doc_for_name.paragraphs[2].text or ""
                    suggested = suggest_feast_outline_filename(third_text, outline_type)
                    if suggested:
                        download_filename = suggested
            except Exception:
                pass
            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()
            return {"docx_bytes": docx_bytes, "filename": download_filename, "error": None}
        except Exception as e:
            logger.error(f"节期纲目刷格式失败: {e}", exc_info=True)
            return {"docx_bytes": None, "filename": "节期纲目.docx", "error": str(e)}
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    os.unlink(temp_docx_path)
                except Exception:
                    pass

    def feast_outline_collect_scripture(self, outline_text: str) -> str:
        """节期纲目 - 带经文：用经文汇集处理纲目，返回带经文内容的纯文本（用于后续刷格式）。"""
        try:
            from tools.biblecollection import biblecollection
        except ImportError:
            logger.warning("biblecollection 未导入，节期纲目带经文功能不可用")
            return outline_text
        data = biblecollection(outline_text)
        lines = []
        for item in data:
            line = (item.get("text") or "").strip()
            if line:
                lines.append(line)
            vers = item.get("vers") or []
            for v in vers:
                src = (v.get("source") or "").strip()
                txt = (v.get("text") or "").strip()
                if src or txt:
                    lines.append(f"　{src}　{txt}")
        return "\n".join(lines) if lines else outline_text

    def feast_outline_morning_revival(self, content: str) -> Dict[str, Any]:
        """节期纲目 - 晨兴信息选读：用 Claude 根据晨兴内容生成纲目。返回 { outline: str, error: str | None }"""
        try:
            from .feast_outline_prompts import get_morning_revival_prompt
        except ImportError:
            return {"outline": "", "error": "节期纲目 prompt 未找到"}
        prompt = get_morning_revival_prompt(content)
        if not claude_client:
            return {"outline": "", "error": "Claude 客户端未初始化"}
        try:
            message = _call_claude_with_retry(prompt, max_tokens=8192)
            text = _claude_message_text(message)
            try:
                in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                logger.info(f"[Claude节期纲目-晨兴] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                if self.redis:
                    get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
            except Exception:
                pass
            text = _strip_code_fence_for_outline(text) or text
            return {"outline": text or "", "error": None}
        except Exception as e:
            logger.error(f"节期纲目晨兴生成失败: {e}", exc_info=True)
            return {"outline": "", "error": str(e)}

    def feast_outline_transcript(
        self,
        original_outline: str,
        transcript: str,
        transcript_preface: Optional[str] = None,
        transcript_addendum: Optional[str] = None,
    ) -> Dict[str, Any]:
        """节期纲目 - 听抄稿：用 Claude 在原纲目基础上加入听抄稿重点。
        若提供 transcript_preface/transcript_addendum，会一并生成序言/添言纲目并返回。"""
        try:
            from .feast_outline_prompts import get_transcript_prompt
        except ImportError:
            return {"outline": "", "error": "节期纲目 prompt 未找到"}
        prompt = get_transcript_prompt(original_outline, transcript)
        if not claude_client:
            return {"outline": "", "error": "Claude 客户端未初始化"}
        try:
            preface_raw = (transcript_preface or "").strip()
            addendum_raw = (transcript_addendum or "").strip()

            def _main_outline():
                message = _call_claude_with_retry(prompt, max_tokens=8192)
                text = _claude_message_text(message)
                try:
                    in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                    out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                    cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                    logger.info(f"[Claude节期纲目-听抄稿] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                    if self.redis:
                        get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
                except Exception:
                    pass
                return _strip_code_fence_for_outline(text) or text or ""

            def _preface():
                return self.feast_outline_preface(preface_raw) if preface_raw else {"outline": "", "error": None}

            def _addendum():
                return self.feast_outline_addendum(addendum_raw) if addendum_raw else {"outline": "", "error": None}

            if preface_raw or addendum_raw:
                # 主纲目、序言、添言三者并行
                with ThreadPoolExecutor(max_workers=3) as ex:
                    f_main = ex.submit(_main_outline)
                    f_preface = ex.submit(_preface) if preface_raw else None
                    f_addendum = ex.submit(_addendum) if addendum_raw else None
                    main_text = f_main.result()
                    result = {"outline": main_text or "", "error": None}
                    if f_preface:
                        res = f_preface.result()
                        result["preface_outline"] = (res.get("outline") or "").strip() if not res.get("error") else ""
                    if f_addendum:
                        res = f_addendum.result()
                        result["addendum_outline"] = (res.get("outline") or "").strip() if not res.get("error") else ""
            else:
                message = _call_claude_with_retry(prompt, max_tokens=8192)
                text = _claude_message_text(message)
                try:
                    in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                    out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                    cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                    logger.info(f"[Claude节期纲目-听抄稿] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                    if self.redis:
                        get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
                except Exception:
                    pass
                result = {"outline": (_strip_code_fence_for_outline(text) or text or ""), "error": None}
            return result
        except Exception as e:
            logger.error(f"节期纲目听抄稿生成失败: {e}", exc_info=True)
            return {"outline": "", "error": str(e)}

    def feast_outline_composite(self, transcript_outline: str, morning_revival_outline: str) -> Dict[str, Any]:
        """节期纲目 - 复合：用 Claude 将晨兴纲目融入听抄稿纲目。返回 { outline: str, error: str | None }"""
        try:
            from .feast_outline_prompts import get_composite_prompt
        except ImportError:
            return {"outline": "", "error": "节期纲目 prompt 未找到"}
        prompt = get_composite_prompt(transcript_outline, morning_revival_outline)
        if not claude_client:
            return {"outline": "", "error": "Claude 客户端未初始化"}
        try:
            message = _call_claude_with_retry(prompt, max_tokens=8192)
            text = _claude_message_text(message)
            try:
                in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                logger.info(f"[Claude节期纲目-复合] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                if self.redis:
                    get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
            except Exception:
                pass
            text = _strip_code_fence_for_outline(text) or text
            return {"outline": text or "", "error": None}
        except Exception as e:
            logger.error(f"节期纲目复合生成失败: {e}", exc_info=True)
            return {"outline": "", "error": str(e)}

    def feast_outline_preface(self, content: str) -> Dict[str, Any]:
        """节期纲目 - 序言：用 Claude 将序言内容整理成纲目格式。返回 { outline: str, error: str | None }"""
        try:
            from .feast_outline_prompts import get_preface_outline_prompt
        except ImportError:
            return {"outline": "", "error": "节期纲目 prompt 未找到"}
        prompt = get_preface_outline_prompt(content)
        if not claude_client:
            return {"outline": "", "error": "Claude 客户端未初始化"}
        try:
            message = _call_claude_with_retry(prompt, max_tokens=4096)
            text = _claude_message_text(message)
            try:
                in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                logger.info(f"[Claude节期纲目-序言] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                if self.redis:
                    get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
            except Exception:
                pass
            text = _strip_code_fence_for_outline(text) or text
            return {"outline": (text or "").strip(), "error": None}
        except Exception as e:
            logger.error(f"节期纲目序言生成失败: {e}", exc_info=True)
            return {"outline": "", "error": str(e)}

    def feast_outline_addendum(self, content: str) -> Dict[str, Any]:
        """节期纲目 - 添言：用 Claude 将添言内容整理成纲目格式。返回 { outline: str, error: str | None }"""
        try:
            from .feast_outline_prompts import get_addendum_outline_prompt
        except ImportError:
            return {"outline": "", "error": "节期纲目 prompt 未找到"}
        prompt = get_addendum_outline_prompt(content)
        if not claude_client:
            return {"outline": "", "error": "Claude 客户端未初始化"}
        try:
            message = _call_claude_with_retry(prompt, max_tokens=4096)
            text = _claude_message_text(message)
            try:
                in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                logger.info(f"[Claude节期纲目-添言] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                if self.redis:
                    get_monitoring(self.redis).record_tool_usage("feast_outline_claude", "claude", in_tok, out_tok, cost)
            except Exception:
                pass
            text = _strip_code_fence_for_outline(text) or text
            return {"outline": (text or "").strip(), "error": None}
        except Exception as e:
            logger.error(f"节期纲目添言生成失败: {e}", exc_info=True)
            return {"outline": "", "error": str(e)}

    def info_retrieval_export(
        self,
        keyword: str,
        exclude_keywords: str = "",
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Tuple[Optional[bytes], Optional[str], str]:
        """
        信息检索：多关键词 AND、排除关键词 OR。单文件上限 40MB，超出则拆成多个 DOCX（-1、-2…）并打包为 ZIP。
        有结果时：单文件返回 (docx_bytes, filename, log_message)，多文件返回 ([(bytes, filename), ...], None, log_message)；无结果时返回 (None, None, log_message)。
        若提供 progress_callback，会在检索过程中实时回调进度文案。
        """
        def _progress(msg: str) -> None:
            if progress_callback:
                progress_callback(msg)

        raw_keyword = (keyword or "").strip()
        keywords = [k.strip() for k in raw_keyword.split() if k.strip()]
        if not keywords:
            log_message = "请输入至少一个搜索关键词（当前为空或仅空格）。"
            return (None, None, log_message)

        exclude_list = [e.strip() for e in (exclude_keywords or "").split() if e.strip()]
        max_bytes = 40 * 1024 * 1024  # 固定 40MB，超出则拆成多个 DOCX
        _progress(f"开始检索，共 {len(INFO_RETRIEVAL_INDEXES)} 个索引")

        # 用于 wildcard 的转义：* \ ? 需转义
        def _escape_wildcard(s: str) -> str:
            for c, rep in [("\\", "\\\\"), ("*", "\\*"), ("?", "\\?")]:
                s = s.replace(c, rep)
            return s

        # 文件名：仅保留安全字符
        def _safe_filename(s: str) -> str:
            s = re.sub(r'[/\\:*?"<>|]', "_", s)
            return (s.strip() or "export")[:100]

        def _item_bytes(label: str, title: str, content: str) -> int:
            return len((label + title + content).encode("utf-8"))

        def _flush_batch():
            if current_batch:
                batches.append(list(current_batch))
                current_batch.clear()
            return 0

        batches = []  # List[List[(label, title, content)]], 每批 ≤40MB
        current_batch = []
        current_size = 0
        seen_prefix = set()  # (index_name, prefix) 或 (index_name, doc_id) 去重
        seen_map_id = set()  # (index_name, doc_id) map 类去重

        for index_name in INFO_RETRIEVAL_INDEXES:
            index_label = INDEX_LABELS.get(index_name, index_name)
            _progress(f"正在检索 {index_label}…")
            try:
                if index_name in self._MAP_LIKE_INDICES:
                    # map 类：单关键词用 match_phrase（短语相邻），多关键词用 match AND（避免「珍赏职事」拆成「珍赏」「职事」分散命中）
                    if len(keywords) == 1:
                        must_clauses = [{"match_phrase": {"text": keywords[0]}}]
                    else:
                        must_clauses = [{"match": {"text": kw}} for kw in keywords]
                    q = {"bool": {"must": must_clauses}}
                    if exclude_list:
                        q["bool"]["must_not"] = [{"match_phrase": {"text": ex}} for ex in exclude_list]
                    body = {
                        "query": q,
                        "size": 10000,
                        "_source": ["id", "text", "msg", "source", "sn", "bookname", "title", "bookname2"],
                    }
                else:
                    # 非 map：title wildcard 多关键词 AND，排除词 OR
                    must_clauses = [{"wildcard": {"title": "*%s*" % _escape_wildcard(kw)}} for kw in keywords]
                    q = {"bool": {"must": must_clauses}}
                    if exclude_list:
                        q["bool"]["must_not"] = [{"wildcard": {"title": "*%s*" % _escape_wildcard(ex)}} for ex in exclude_list]
                    body = {
                        "query": q,
                        "size": 10000,
                        "_source": ["id", "type", "text", "title", "book", "chapter", "verse"],
                    }
                resp = self.es.search(
                    index=index_name,
                    body=body,
                    request_timeout=60,
                    ignore_unavailable=True,
                )
            except Exception as e:
                logger.warning(f"信息检索索引 {index_name} 查询失败: {e}")
                _progress(f"索引 {index_label} 查询失败")
                continue

            hits = resp.get("hits", {}).get("hits", [])

            if index_name in self._MAP_LIKE_INDICES:
                for h in hits:
                    source = h.get("_source", {})
                    doc_id = str(source.get("id") or h.get("_id") or "")
                    if (index_name, doc_id) in seen_map_id:
                        continue
                    seen_map_id.add((index_name, doc_id))
                    # 返回同 id 文档的 msg 中所有 text 的拼接
                    msg_list = source.get("msg") or []
                    parts = [t for m in msg_list for t in [str(m.get("text") or "").strip()] if t]
                    content = "\n".join(parts)
                    if not content.strip():
                        continue
                    # 标题与 AI 搜索引用来源一致
                    title = self._get_map_note_reference_from_hit(source, h, index_name)
                    # 多关键词时只保留篇题同时包含所有关键词的结果（真 AND）
                    if len(keywords) > 1 and not all(kw in title for kw in keywords):
                        continue
                    label = INDEX_LABELS.get(index_name, index_name)
                    item_size = _item_bytes(label, title, content)
                    if current_size + item_size > max_bytes and current_batch:
                        _flush_batch()
                        current_size = 0
                    current_batch.append((label, title, content))
                    current_size += item_size
            else:
                for h in hits:
                    source = h.get("_source", {})
                    doc_id = str(source.get("id") or h.get("_id", ""))
                    prefix, _ = self._parse_doc_id(doc_id)
                    if prefix:
                        if (index_name, prefix) in seen_prefix:
                            continue
                        seen_prefix.add((index_name, prefix))
                        docs = self._fetch_message_docs(index_name, prefix)
                        if not docs:
                            continue
                        parts = []
                        for d in docs:
                            t = str(d.get("text") or "").strip()
                            if t:
                                parts.append(t)
                        content = "\n".join(parts)
                        doc_title = self._format_reference(docs[0])
                    else:
                        # id 无 "-" 时视为单条文档，不按 prefix 拉整篇
                        if (index_name, doc_id) in seen_prefix:
                            continue
                        seen_prefix.add((index_name, doc_id))
                        content = str(source.get("text") or "").strip()
                        doc_title = self._format_reference(source)
                    if not content.strip():
                        continue
                    # 多关键词时只保留篇题同时包含所有关键词的结果（真 AND）
                    if len(keywords) > 1 and not all(kw in doc_title for kw in keywords):
                        continue
                    label = INDEX_LABELS.get(index_name, index_name)
                    item_size = _item_bytes(label, doc_title, content)
                    if current_size + item_size > max_bytes and current_batch:
                        _flush_batch()
                        current_size = 0
                    current_batch.append((label, doc_title, content))
                    current_size += item_size
            total_items = sum(len(b) for b in batches) + len(current_batch)
            _progress(f"完成 {index_label}，命中 {len(hits)} 条，当前累计 {total_items} 条")

        if current_batch:
            batches.append(current_batch)

        if not batches:
            log_message = "未找到匹配的文档。请检查关键词或排除词，或稍后重试（部分索引可能暂时不可用）。"
            return (None, None, log_message)

        _progress("正在生成 DOCX…")
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise RuntimeError("请安装 python-docx: pip install python-docx")

        name_base = "_".join(keywords) if keywords else raw_keyword
        if exclude_list:
            name_base += "（过滤" + "、".join(exclude_list) + "）"
        base_name = _safe_filename(name_base)

        # 中文字体名，用于 DOCX 正文与标题，避免手机等设备打开时缺字出现方框乱码
        _DOCX_CJK_FONT = "Microsoft YaHei"

        def _set_docx_cjk_font(doc: "Document") -> None:
            """为文档的 Normal 与 Heading 1 设置支持中文的字体（含东亚 w:eastAsia），减少手机打开时方框乱码。"""
            for style_name in ("Normal", "Heading 1"):
                style = doc.styles[style_name]
                style.font.name = _DOCX_CJK_FONT
                try:
                    if style._element.rPr is not None and style._element.rPr.rFonts is not None:
                        style._element.rPr.rFonts.set(qn("w:eastAsia"), _DOCX_CJK_FONT)
                except (AttributeError, Exception):
                    pass

        def _make_docx(items_batch: List[Tuple[str, str, str]]) -> bytes:
            doc = Document()
            _set_docx_cjk_font(doc)
            for i, (label, title, content) in enumerate(items_batch):
                doc.add_heading(f"[{label}] {title}", level=1)
                doc.add_paragraph(content)
                if i < len(items_batch) - 1:
                    doc.add_paragraph()
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()

        total_items = sum(len(b) for b in batches)
        if len(batches) == 1:
            filename = base_name + ".docx"
            log_message = f"共导出 {total_items} 条，已生成 {filename}"
            return (_make_docx(batches[0]), filename, log_message)

        # 多文件：返回 [(bytes, filename), ...]，由路由以 JSON 返回，前端逐个下载
        files_list = [
            (_make_docx(batch), f"{base_name}-{i}.docx")
            for i, batch in enumerate(batches, start=1)
        ]
        log_message = f"共导出 {total_items} 条，已生成 {len(batches)} 个文件：{base_name}-1.docx ～ {base_name}-{len(batches)}.docx"
        return (files_list, None, log_message)

    def health_check(self) -> Dict:
        """健康检查"""
        status = {
            "elasticsearch": False,
            "redis": False,
            "claude": False,
            "gemini": False,
            "overall": False
        }

        try:
            # 检查ES
            status["elasticsearch"] = self.es.ping()
        except Exception:
            pass

        try:
            # 检查Redis
            if self.redis:
                status["redis"] = self.redis.ping()
        except Exception:
            pass

        try:
            # 检查Claude（通过检查API key是否存在）
            status["claude"] = bool(CLAUDE_API_KEY)
        except Exception:
            pass

        try:
            # 检查Gemini（英文翻译，可选）
            status["gemini"] = bool(GEMINI_API_KEY and gemini_client)
        except Exception:
            pass

        # 核心依赖为 ES + Claude；Redis、Gemini 为可选
        status["overall"] = status["elasticsearch"] and status["claude"]

        return status

    def get_rough_outline_ai_counts(self) -> Dict[str, int]:
        """返回每种纲目类型对应的 AI 数量（即该类型需调用几次 API）。"""
        try:
            from .rough_outline_prompts import get_ai_configs
            types = ("polish", "beginner", "youth", "truth", "sharing")
            return {t: len(get_ai_configs(t)) for t in types}
        except Exception as e:
            logger.error(f"get_rough_outline_ai_counts 失败: {e}", exc_info=True)
            return {}

    def generate_rough_outline(
        self,
        outline_type: str,
        content: str,
        ai_index: int = 0,
    ) -> Dict:
        """
        生成毛胚纲目。每次只调用一个 AI，生成一篇。
        
        Args:
            outline_type: 纲目类型 ("polish", "beginner", "youth", "truth", "sharing")
            content: 原始纲目内容
            ai_index: 该类型下第几个 AI（0 起），只调用这一个
        
        Returns:
            {
                "results": [{"type": str, "content": str, "ai_model": str}],  # 长度 0 或 1
                "error": str | None,
            }
        """
        try:
            from .rough_outline_prompts import get_prompt_template, get_ai_configs
            
            ai_configs = get_ai_configs(outline_type)
            if not ai_configs:
                return {
                    "results": [],
                    "error": f"不支持的纲目类型: {outline_type}",
                }
            
            if ai_index < 0 or ai_index >= len(ai_configs):
                return {
                    "results": [],
                    "error": f"ai_index {ai_index} 超出范围 (0~{len(ai_configs) - 1})",
                }
            
            ai_config = ai_configs[ai_index]
            # 润色版等类型可能按 AI 使用不同 prompt（prompt_key：如 polish_gemini、polish_claude）
            prompt_key = ai_config.get("prompt_key")
            prompt_template = get_prompt_template(outline_type, prompt_key=prompt_key)
            if not prompt_template:
                return {
                    "results": [],
                    "error": f"未找到 prompt 模板 (类型: {outline_type}, prompt_key: {prompt_key})",
                }
            
            prompt = prompt_template.replace("{content}", content)
            ai_name = ai_config.get("name", "Unknown")
            
            logger.info(f"使用 {ai_name} 生成毛胚纲目 (类型: {outline_type}, ai_index: {ai_index})")
            
            generated_content = self._call_ai_for_rough_outline(ai_config, prompt)
            content = None
            tokens = None
            if generated_content is not None:
                if isinstance(generated_content, tuple):
                    content = generated_content[0]
                    tokens = generated_content[1] if len(generated_content) > 1 else None
                else:
                    content = generated_content
            if content:
                content = _strip_code_fence_for_outline(content) or content
            if content:
                return {
                    "results": [{
                        "type": outline_type,
                        "content": content,
                        "ai_model": ai_name,
                        "tokens": tokens or {"input": 0, "output": 0, "cost": 0},
                    }],
                    "error": None,
                }
            return {
                "results": [],
                "error": f"{ai_name} 生成失败",
            }
        except Exception as e:
            logger.error(f"generate_rough_outline 失败: {e}", exc_info=True)
            return {
                "results": [],
                "error": f"生成失败: {str(e)}",
            }


    def _call_ai_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[tuple]:
        """调用指定的 AI API 生成内容，返回 (content, tokens_dict) 或 None"""
        ai_type = ai_config.get("type")
        ai_name = ai_config.get("name", "Unknown")
        
        try:
            if ai_type == "claude":
                return self._call_claude_for_rough_outline(ai_config, prompt)
            elif ai_type == "gemini":
                return self._call_gemini_for_rough_outline(ai_config, prompt)
            elif ai_type == "deepseek":
                return self._call_deepseek_for_rough_outline(ai_config, prompt)
            elif ai_type == "perplexity":
                return self._call_perplexity_for_rough_outline(ai_config, prompt)
            elif ai_type == "chatgpt":
                return self._call_chatgpt_for_rough_outline(ai_config, prompt)
            elif ai_type == "grok":
                return self._call_grok_for_rough_outline(ai_config, prompt)
            else:
                logger.error(f"不支持的 AI 类型: {ai_type}")
                return None
        except Exception as e:
            logger.error(f"调用 {ai_name} 失败: {e}", exc_info=True)
            return None

    def _call_claude_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """调用 Claude API"""
        if not claude_client:
            logger.error("Claude 客户端未初始化")
            return None
        
        model = ai_config.get("model", CLAUDE_MODEL)
        max_tokens = ai_config.get("max_tokens", 8192)
        
        try:
            with CLAUDE_SEMAPHORE:
                message = claude_client.messages.create(
                    model=model,
                    max_tokens=max_tokens,
                    messages=[{"role": "user", "content": prompt}],
                )
                text = message.content[0].text
                try:
                    in_tok = int(getattr(message.usage, "input_tokens", 0) or 0)
                    out_tok = int(getattr(message.usage, "output_tokens", 0) or 0)
                    cost = (in_tok * 3 + out_tok * 15) / 1_000_000
                    logger.info(f"[Claude毛胚] 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                    if self.redis:
                        get_monitoring(self.redis).record_tool_usage("rough_outline_claude", model, in_tok, out_tok, cost)
                    tokens_out = {"input": in_tok, "output": out_tok, "cost": cost}
                except Exception:
                    tokens_out = {"input": 0, "output": 0, "cost": 0}
                return (text, tokens_out)
        except Exception as e:
            logger.error(f"Claude API 调用失败: {e}", exc_info=True)
            return None

    def _gemini_rough_outline_generate(
        self, model: str, prompt: str, max_tokens: int, max_retries: int = 0
    ) -> Optional[tuple]:
        """调用 Gemini 生成毛胚纲目；max_retries 为同模型 503/429 时的额外重试次数。"""
        backoff_seconds = (8, 15)
        for attempt in range(max_retries + 1):
            try:
                from google.genai import types
                with GEMINI_SEMAPHORE:
                    response = gemini_client.models.generate_content(
                        model=model,
                        contents=prompt,
                        config=types.GenerateContentConfig(max_output_tokens=max_tokens),
                    )
                text = None
                tokens_out = None
                if hasattr(response, "text"):
                    text = response.text
                elif hasattr(response, "candidates") and response.candidates:
                    text = response.candidates[0].content.parts[0].text
                if text is not None:
                    try:
                        usage_meta = response.usage_metadata
                        in_tok = int(getattr(usage_meta, "prompt_token_count", 0) or 0)
                        out_tok = int(getattr(usage_meta, "candidates_token_count", 0) or 0)
                        cost = (in_tok * 1.25 + out_tok * 10) / 1_000_000
                        logger.info(
                            f"[Gemini毛胚] model={model} | 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}"
                        )
                        if self.redis:
                            get_monitoring(self.redis).record_tool_usage(
                                "rough_outline_gemini", model, in_tok, out_tok, cost
                            )
                        tokens_out = {"input": in_tok, "output": out_tok, "cost": cost}
                    except Exception:
                        pass
                    return (text, tokens_out)
                logger.error(f"Gemini 响应格式异常 (model={model}): {response}")
                return None
            except Exception as e:
                status = getattr(e, "status_code", None) or (e.args[0] if e.args else None)
                retryable = status in (503, 429) or "503" in str(e) or "429" in str(e)
                if retryable and attempt < max_retries:
                    wait = backoff_seconds[attempt] if attempt < len(backoff_seconds) else 40
                    logger.warning(
                        "Gemini 暂时不可用 (503/429)，%s 秒后重试 (%s/%s)，model=%s",
                        wait, attempt + 1, max_retries, model,
                    )
                    time.sleep(wait)
                else:
                    logger.error(f"Gemini API 调用失败 (model={model}): {e}", exc_info=True)
                    return None
        return None

    def _call_gemini_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """
        毛胚纲目 Gemini：首次 gemini-3.1-pro-preview，失败后改用 gemini-2.5-pro 重试（含 503/429 退避）。
        可通过 ROUGH_OUTLINE_GEMINI_MODEL / ROUGH_OUTLINE_GEMINI_FALLBACK_MODEL 覆盖。
        """
        if not gemini_client:
            logger.error("Gemini 客户端未初始化")
            return None

        primary = (ROUGH_OUTLINE_GEMINI_MODEL or "").strip() or ai_config.get(
            "model", "gemini-3.1-pro-preview"
        )
        fallback = (ROUGH_OUTLINE_GEMINI_FALLBACK_MODEL or "").strip() or "gemini-2.5-pro"
        max_tokens = ai_config.get("max_tokens", 8192)

        result = self._gemini_rough_outline_generate(primary, prompt, max_tokens, max_retries=0)
        if result:
            return result

        if fallback and fallback != primary:
            logger.warning("Gemini 3.1 Pro 不可用，改用备用模型重试: %s", fallback)
            return self._gemini_rough_outline_generate(fallback, prompt, max_tokens, max_retries=2)
        return None

    def _call_openai_compatible_rough_outline(
        self,
        api_key: str,
        prompt: str,
        model: str,
        max_tokens: int = 8192,
        base_url: Optional[str] = None,
        use_max_completion_tokens: bool = False,
    ) -> Optional[str]:
        """通用：OpenAI 兼容接口（Deep Seek / OpenAI / Perplexity / xAI Grok 均兼容）。OpenAI 新模型需传 max_completion_tokens。"""
        try:
            from openai import OpenAI
            kwargs = {"api_key": api_key, "timeout": 120.0}
            if base_url:
                kwargs["base_url"] = base_url
            client = OpenAI(**kwargs)
            create_kw = {"model": model, "messages": [{"role": "user", "content": prompt}]}
            if use_max_completion_tokens:
                create_kw["max_completion_tokens"] = max_tokens
            else:
                create_kw["max_tokens"] = max_tokens
            r = client.chat.completions.create(**create_kw)
            content = None
            if r.choices and len(r.choices) > 0 and getattr(r.choices[0].message, "content", None):
                content = r.choices[0].message.content
            tokens_out = None
            if content is not None:
                try:
                    in_tok = int(getattr(r.usage, "prompt_tokens", 0) or 0)
                    out_tok = int(getattr(r.usage, "completion_tokens", 0) or 0)
                    if "deepseek" in model.lower():
                        provider = "deepseek"
                    elif "grok" in model.lower() or "grok" in (base_url or "").lower():
                        provider = "grok"
                    elif "perplexity" in (base_url or "").lower() or "sonar" in (model or "").lower():
                        provider = "perplexity"
                    else:
                        provider = "chatgpt"
                    price = ROUGH_OUTLINE_PRICES.get(provider, {"in": 2.50, "out": 15.00})
                    cost = (in_tok * price["in"] + out_tok * price["out"]) / 1_000_000
                    logger.info(f"[OpenAI兼容毛胚] model={model} | 输入={in_tok} tokens | 输出={out_tok} tokens | 费用=${cost:.6f}")
                    if self.redis:
                        if provider == "deepseek":
                            tool = "rough_outline_deepseek"
                        elif provider == "grok":
                            tool = "rough_outline_grok"
                        elif provider == "perplexity":
                            tool = "rough_outline_perplexity"
                        else:
                            tool = "rough_outline_openai"
                        get_monitoring(self.redis).record_tool_usage(tool, model, in_tok, out_tok, cost)
                    tokens_out = {"input": in_tok, "output": out_tok, "cost": cost}
                except Exception:
                    pass
            if content is not None:
                return (content, tokens_out)
            return (None, None)
        except Exception as e:
            logger.error(f"OpenAI 兼容 API 调用失败 (model={model}): {e}", exc_info=True)
            return None

    def _call_deepseek_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """调用 Deep Seek API（OpenAI 兼容，需在 .env 中配置 DEEPSEEK_API_KEY）"""
        if not DEEPSEEK_API_KEY:
            logger.warning("Deep Seek 未配置: 请在 .env 中填写 DEEPSEEK_API_KEY")
            return None
        model = ai_config.get("model", DEEPSEEK_MODEL)
        if model == "deepseek-v3.2":
            model = "deepseek-chat"
        max_tokens = ai_config.get("max_tokens", 8192)
        return self._call_openai_compatible_rough_outline(
            api_key=DEEPSEEK_API_KEY,
            prompt=prompt,
            model=model,
            max_tokens=max_tokens,
            base_url="https://api.deepseek.com",
        )

    def _call_perplexity_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """调用 Perplexity API（OpenAI 兼容，需在 .env 中配置 PERPLEXITY_API_KEY）"""
        if not PERPLEXITY_API_KEY:
            logger.warning("Perplexity 未配置: 请在 .env 中填写 PERPLEXITY_API_KEY")
            return None
        model = ai_config.get("model", PERPLEXITY_MODEL)
        if model and "pplx-" in model:
            model = "sonar-pro"
        max_tokens = ai_config.get("max_tokens", 8192)
        return self._call_openai_compatible_rough_outline(
            api_key=PERPLEXITY_API_KEY,
            prompt=prompt,
            model=model or "sonar-pro",
            max_tokens=max_tokens,
            base_url="https://api.perplexity.ai",
        )

    def _call_chatgpt_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """调用 OpenAI / ChatGPT API（需在 .env 中配置 OPENAI_API_KEY）。新模型使用 max_completion_tokens。"""
        if not OPENAI_API_KEY:
            logger.warning("OpenAI 未配置: 请在 .env 中填写 OPENAI_API_KEY")
            return None
        model = ai_config.get("model", OPENAI_MODEL)
        max_tokens = ai_config.get("max_tokens", 8192)
        return self._call_openai_compatible_rough_outline(
            api_key=OPENAI_API_KEY,
            prompt=prompt,
            model=model,
            max_tokens=max_tokens,
            base_url=None,
            use_max_completion_tokens=True,
        )

    def _call_grok_for_rough_outline(self, ai_config: Dict, prompt: str) -> Optional[str]:
        """调用 xAI Grok API（OpenAI 兼容，需在 .env 中配置 XAI_API_KEY）"""
        if not XAI_API_KEY:
            logger.warning("xAI Grok 未配置: 请在 .env 中填写 XAI_API_KEY")
            return None
        model = ai_config.get("model", XAI_MODEL)
        max_tokens = ai_config.get("max_tokens", 8192)
        return self._call_openai_compatible_rough_outline(
            api_key=XAI_API_KEY,
            prompt=prompt,
            model=model,
            max_tokens=max_tokens,
            base_url="https://api.x.ai/v1",
        )


def _paragraph_has_page_break(para) -> bool:
    """判断段落内是否包含分页符（任一 run 内有 w:br type=page）。"""
    try:
        from docx.oxml.ns import qn
        for run in para.runs:
            for child in run._element:
                if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
                    return True
    except Exception:
        pass
    return False


def _get_addendum_paragraph_range(doc: "Document", to_page_break: bool):
    """
    按内容定位「添言」：找到第一个包含「添言」的段落下标 start。
    to_page_break=True（听抄稿）时 end 为 start 之后第一个含分页符的段落的前一段（不含分页符段）；
    to_page_break=False（复合稿）时 end 为文档末尾。
    返回 (start, end) 或 None（未找到「添言」）。
    """
    start_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if "添言" in (para.text or ""):
            start_idx = idx
            break
    if start_idx is None:
        return None
    n = len(doc.paragraphs)
    if to_page_break:
        # 听抄稿：到分页符前一段为止，分页符段落不高亮
        end_idx = n - 1
        for idx in range(start_idx + 1, n):
            if _paragraph_has_page_break(doc.paragraphs[idx]):
                end_idx = idx - 1
                break
    else:
        end_idx = n - 1
    return (start_idx, end_idx)


def _apply_transcript_add_highlight(doc: "Document") -> None:
    """
    听抄稿纲目：与工具箱-毛胚纲目-真理加强版同一逻辑，仅改为黄色高亮（真理加强版为下划线）。
    定位所有【听抄稿添加开始】与【听抄稿添加结束】配对（按出现顺序），
    对每对之间的段落整段黄色高亮，并删除所有标记段落。支持简繁体标记。
    """
    from docx.enum.text import WD_COLOR_INDEX

    start_marker_variants = ("【听抄稿添加开始】", "【聽抄稿添加開始】")
    end_marker_variants = ("【听抄稿添加结束】", "【聽抄稿添加結束】")
    all_markers = start_marker_variants + end_marker_variants

    def _is_pure_marker_line(text: str) -> bool:
        """仅含标记（可带句读、空白）的段落才作为删除目标，不含其他正文。"""
        t = text.strip()
        if not any(m in t for m in all_markers):
            return False
        remainder = t
        for m in all_markers:
            remainder = remainder.replace(m, "")
        remainder = remainder.replace("。", "").strip()
        return not remainder

    start_indices = []
    pairs = []  # [(start_idx, end_idx), ...]
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(m in text for m in start_marker_variants):
            # 同行另有正文时不作为区间起点，避免高亮范围偏移
            if _is_pure_marker_line(text):
                start_indices.append(idx)
        if any(m in text for m in end_marker_variants):
            if start_indices:
                pairs.append((start_indices.pop(), idx))
    if not pairs:
        return
    # 从最后一对往第一对处理，删除段落时不会影响前面配对的下标
    for (start_idx, end_idx) in reversed(pairs):
        if start_idx >= end_idx:
            continue
        # 对两标记之间的段落整段黄色高亮；含标记的行不纳入高亮
        for idx in range(start_idx + 1, end_idx):
            para_text = doc.paragraphs[idx].text
            if any(m in para_text for m in all_markers):
                continue
            for run in doc.paragraphs[idx].runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        # 仅删除纯标记段落（先删结束，再删开始）
        for del_idx in (end_idx, start_idx):
            if del_idx < len(doc.paragraphs) and _is_pure_marker_line(doc.paragraphs[del_idx].text):
                el = doc.paragraphs[del_idx]._element
                el.getparent().remove(el)


def _get_feast_body_style(doc: "Document"):
    """节期纲目模板正文段落样式：优先 0000模板，其次 9职事信息摘录，最后 Normal。"""
    for name in ("0000模板", "9职事信息摘录", "Normal"):
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return None


def _append_morning_revival_section(doc: "Document", morning_revival_raw: str) -> None:
    """晨兴信息选读：按节期纲目刷格式逻辑，分页 +「晨兴圣言信息：」(9职事信息摘录) + 小标题(81级标题)/正文(0000模板)。"""
    from docx.enum.text import WD_BREAK

    body_style = _get_feast_body_style(doc)
    title_ends = ('。', '！', '？', '…', '"', '\'', '）', '：', '』')
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    title_para = doc.add_paragraph("晨兴圣言信息：")
    try:
        title_para.style = doc.styles["9职事信息摘录"]
    except KeyError:
        try:
            title_para.style = doc.styles["81级标题"]
        except KeyError:
            pass
    for line in morning_revival_raw.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            if not (line.rstrip().endswith(title_ends)):
                try:
                    p.style = doc.styles["81级标题"]
                except KeyError:
                    if body_style is not None:
                        p.style = body_style
            elif body_style is not None:
                p.style = body_style
        else:
            doc.add_paragraph("")


def _append_transcript_info_section(
    doc: "Document",
    transcript_raw: str,
    transcript_preface: Optional[str] = None,
    transcript_addendum: Optional[str] = None,
) -> None:
    """听抄稿：分页 +「听抄信息：」+ 序言、听抄稿、添言三段合并内容（按节期纲目刷格式逻辑套样式）。"""
    from docx.enum.text import WD_BREAK

    parts = []
    if (transcript_preface or "").strip():
        parts.append((transcript_preface or "").strip())
    if (transcript_raw or "").strip():
        parts.append((transcript_raw or "").strip())
    if (transcript_addendum or "").strip():
        parts.append((transcript_addendum or "").strip())
    combined = "\n\n".join(parts) if parts else ""

    body_style = _get_feast_body_style(doc)
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    title_para = doc.add_paragraph("听抄信息：")
    try:
        title_para.style = doc.styles["9职事信息摘录"]
    except KeyError:
        try:
            title_para.style = doc.styles["81级标题"]
        except KeyError:
            pass
    try:
        import sys
        from pathlib import Path
        _backend = Path(__file__).resolve().parent.parent
        if str(_backend) not in sys.path:
            sys.path.insert(0, str(_backend))
        from 节期纲目刷格式 import apply_custom_92_style
    except ImportError:
        apply_custom_92_style = None
    for line in combined.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            if '。' not in line and apply_custom_92_style is not None:
                apply_custom_92_style(p)
            elif body_style is not None:
                p.style = body_style
        else:
            doc.add_paragraph("")


def format_english_bibco_docx(contents: str, filename: str) -> Dict:
    """
    将英文经文汇集纯文本写入 DOCX 并刷格式，返回 base64 编码结果。

    Args:
        contents: 多行文本（纲目行 + 经文行）
        filename: 下载文件名（不含 .docx 后缀）

    Returns:
        {"docx_base64": str, "filename": str} 或 {"error": str}
    """
    import base64
    import shutil
    import tempfile
    from docx import Document

    if format_english_outline_docx is None:
        return {"error": "格式刷函数未导入"}

    backend_dir = Path(__file__).resolve().parent.parent
    template_path = backend_dir / "英文纲目模板.docx"
    if not template_path.exists():
        return {"error": f"模板文件不存在: 英文纲目模板.docx"}

    out_name = filename if filename.endswith(".docx") else f"{filename}.docx"
    temp_docx_path = None

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            temp_docx_path = tmp_file.name

        shutil.copy2(template_path, temp_docx_path)
        doc = Document(temp_docx_path)

        for para in list(doc.paragraphs):
            p_element = para._element
            p_element.getparent().remove(p_element)

        for line in contents.split("\n"):
            doc.add_paragraph(line)

        doc.save(temp_docx_path)
        format_english_outline_docx(temp_docx_path)

        with open(temp_docx_path, "rb") as f:
            docx_bytes = f.read()

        return {
            "docx_base64": base64.b64encode(docx_bytes).decode("utf-8"),
            "filename": out_name,
        }
    except Exception as e:
        logger.error(f"英文经文汇集刷格式失败: {e}", exc_info=True)
        return {"error": str(e)}
    finally:
        if temp_docx_path:
            try:
                os.remove(temp_docx_path)
            except OSError:
                pass


# 创建全局服务实例
ai_service = AISearchService()
