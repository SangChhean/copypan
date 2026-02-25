# ===== 纲目一键刷格式 5.5版本 =====
# 基于5.0版本，修改双出处纲目星号处理逻辑
# 主要更新：
# 1. 双出处纲目中的星号(*)保留而非去除
# 2. 新增句中句号改为分号功能（参考5.1版本）
# 3. 新增Config统一配置管理（参考5.4版本）
# 4. 新增PerformanceOptimizer性能优化器（参考5.4版本）
# 5. 保持5.0版本的句末标点处理逻辑

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
import sys
import shutil

# 可选依赖：无 GUI 环境（如 Web 后端）下可跳过
try:
    import winsound
except ImportError:
    winsound = None
try:
    import cn2an
except ImportError:
    cn2an = None
try:
    import tkinter as tk
    from tkinter import messagebox
except ImportError:
    tk = None
    messagebox = None

# ===== 5.5版本配置和常量管理 =====
class Config:
    """统一的配置和常量管理类"""
    
    # 版本信息
    VERSION = "5.5"
    DEBUG_MODE = False
    
    # 文档标记和标识符
    AFTER_MARKER_LIST = ['职事信息摘录：', '听抄信息：', '晨兴圣言信息：']
    COMPOUND_OUTLINE_START = "【添加开始】"
    COMPOUND_OUTLINE_END = "【添加结束】"
    
    # 文档类型识别
    CATEGORY_PATTERNS = {
        '（听抄稿的纲目）': '【听抄稿的纲目】',
        '（晨兴信息选读的纲目）': '【晨兴信息选读的纲目】',
        '（复合的纲目）': '【复合的纲目】',
        '（纲目的原文）': '【纲目的原文】',
        '（带经文的纲目）': '【带经文的纲目】',
        '（纲目的双出处）': '【纲目的双出处】',
    }
    
    # 纲目层级匹配模式
    OUTLINE_LEVEL_PATTERNS = [
        (re.compile(r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t'), 2),   # 2 大点
        (re.compile(r'^([一二三四五六七八九十百千万亿]+)\t'),             3),   # 3 中点
        (re.compile(r'^(\d+)\t'),                                         4),   # 4 小点
        (re.compile(r'^([a-z])\t'),                                       5),   # 5 a点
        (re.compile(r'^\(([一二三四五六七八九十百千万亿]+)\)\t'),          6),   # 6 （一）
        (re.compile(r'^\((\d+)\)\t'),                                     7),   # 7 （1）
        (re.compile(r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]\t'),            7),   # 7 带圈数字
        (re.compile(r'^\(([a-z])\)\t'),                                   8),   # 8 （a）
    ]
    
    # 正则表达式模式
    PUNCT_AT_END_RE = re.compile(r'[。！？；.,!?;,，]$')
    
    # 性能优化配置
    ENABLE_CACHE = True
    MAX_CACHE_SIZE = 1000
    BATCH_SIZE = 100

# 为向后兼容性保留的全局变量
after_marker_list = Config.AFTER_MARKER_LIST
DEBUG_MODE = Config.DEBUG_MODE  # 5.5版本：使用统一配置管理

# ===== 5.5版本性能优化器 =====
class PerformanceOptimizer:
    """性能优化器，提供缓存和批量处理功能"""
    
    def __init__(self):
        self.outline_level_cache = {}
        self.cache_hits = 0
        self.cache_misses = 0
        
    def get_outline_level_cached(self, text):
        """缓存的纲目层级识别"""
        if not Config.ENABLE_CACHE:
            return self._detect_outline_level_direct(text)
            
        # 缓存查找
        if text in self.outline_level_cache:
            self.cache_hits += 1
            return self.outline_level_cache[text]
        
        # 缓存未命中，计算结果
        self.cache_misses += 1
        level = self._detect_outline_level_direct(text)
        
        # 添加到缓存（限制缓存大小）
        if len(self.outline_level_cache) < Config.MAX_CACHE_SIZE:
            self.outline_level_cache[text] = level
        
        return level
    
    def _detect_outline_level_direct(self, text):
        """直接检测纲目层级（无缓存）"""
        text = text.strip()
        if not text:
            return None
        
        # === 特殊处理 ===
        if text.startswith("序言") or text.startswith("添言"):
            return 2  # 当作 2大点
        if text.startswith("前言"):
            return 3  # 当作 3中点
        # === 特殊处理结束 ===
            
        # 使用配置中的层级匹配模式
        for pattern, level in Config.OUTLINE_LEVEL_PATTERNS:
            if pattern.match(text):
                return level
        
        return None
    
    def batch_process_paragraphs(self, doc, processor_func):
        """批量处理段落，提高大文档处理效率"""
        paragraphs = doc.paragraphs
        total_paras = len(paragraphs)
        processed_count = 0
        
        # 分批处理
        for i in range(0, total_paras, Config.BATCH_SIZE):
            batch_end = min(i + Config.BATCH_SIZE, total_paras)
            batch = paragraphs[i:batch_end]
            
            # 处理当前批次
            for para_index, para in enumerate(batch, start=i):
                try:
                    if processor_func(para, para_index):
                        processed_count += 1
                except Exception as e:
                    if DEBUG_MODE:
                        print(f"处理段落{para_index}时发生错误: {e}")
                        
        return processed_count
    
    def clear_cache(self):
        """清空缓存"""
        self.outline_level_cache.clear()
        self.cache_hits = 0
        self.cache_misses = 0
        
    def get_cache_stats(self):
        """获取缓存统计信息"""
        total_requests = self.cache_hits + self.cache_misses
        hit_rate = (self.cache_hits / total_requests * 100) if total_requests > 0 else 0
        
        return {
            'cache_hits': self.cache_hits,
            'cache_misses': self.cache_misses,
            'total_requests': total_requests,
            'hit_rate': f"{hit_rate:.1f}%",
            'cache_size': len(self.outline_level_cache)
        }

# 全局性能优化器实例
optimizer = PerformanceOptimizer()

def remove_duplicate_chinese_punctuation(text):
    """
    将连续出现的相同中文标点符号替换为单个标点符号
    """
    # 简单的替换常见重复标点
    text = text.replace('，，', '，')
    text = text.replace('。。', '。')
    text = text.replace('！！', '！')
    text = text.replace('？？', '？')
    text = text.replace('：：', '：')
    text = text.replace('；；', '；')
    
    return text

def remove_text_highlights(doc):
    """
    去除文档中所有文本的突出显示颜色（高亮背景）
    返回处理的run数量
    """
    removed_count = 0
    
    for para in doc.paragraphs:
        for run in para.runs:
            # 检查是否有高亮背景
            if has_highlight_background(run):
                # 移除高亮背景
                remove_highlight_background(run)
                removed_count += 1
    
    return removed_count

def has_highlight_background(run):
    """
    检查run是否有高亮背景
    """
    try:
        # 检查标准高亮
        if run.font.highlight_color is not None:
            return True
        
        # 检查XML中的高亮设置
        if hasattr(run, '_element'):
            rPr = run._element.find('.//w:rPr', run._element.nsmap)
            if rPr is not None:
                highlight_elem = rPr.find('.//w:highlight', rPr.nsmap)
                if highlight_elem is not None:
                    return True
                
                # 检查底纹设置
                shd_elem = rPr.find('.//w:shd', rPr.nsmap)
                if shd_elem is not None:
                    # 检查是否有背景色设置
                    fill_attr = shd_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill_attr and fill_attr.lower() not in ['auto', 'ffffff', 'white']:
                        return True
        
        return False
    except Exception as e:
        return False

def remove_highlight_background(run):
    """
    移除run的高亮背景
    """
    try:
        # 移除标准高亮
        if run.font.highlight_color is not None:
            run.font.highlight_color = None
        
        # 移除XML中的高亮设置
        if hasattr(run, '_element'):
            rPr = run._element.find('.//w:rPr', run._element.nsmap)
            if rPr is not None:
                # 移除highlight元素
                highlight_elem = rPr.find('.//w:highlight', rPr.nsmap)
                if highlight_elem is not None:
                    rPr.remove(highlight_elem)
                
                # 移除底纹元素
                shd_elem = rPr.find('.//w:shd', rPr.nsmap)
                if shd_elem is not None:
                    rPr.remove(shd_elem)
        
    except Exception as e:
        print(f"移除高亮时出错：{e}")

# ===== 修改后的copy_content_to_template函数 =====
def copy_content_to_template_with_highlight_removal(source_doc_path, template_path, output_path):
    """
    将源文档的内容复制到模板文档中，并去除高亮背景
    """
    try:
        # 打开源文档和模板
        source_doc = Document(source_doc_path)
        template_doc = Document(template_path)
        
        # 清空模板的现有内容（保留样式）
        for i in range(len(template_doc.paragraphs) - 1, 0, -1):
            p = template_doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # 清空第一个段落
        if template_doc.paragraphs:
            template_doc.paragraphs[0].clear()
        
        # 复制源文档的所有段落到模板
        for para in source_doc.paragraphs:
            # 在模板中创建新段落
            new_para = template_doc.add_paragraph()
            
            # 复制段落内容和基本格式
            if para.runs:
                # 如果原段落有runs，逐个复制
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    
                    # 复制基本格式
                    if run.font.bold is not None:
                        new_run.font.bold = run.font.bold
                    if run.font.italic is not None:
                        new_run.font.italic = run.font.italic
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
                    if run.font.name is not None:
                        new_run.font.name = run.font.name
                    
                    # 复制字体颜色（保留红色等字体颜色）
                    if run.font.color.rgb is not None:
                        new_run.font.color.rgb = run.font.color.rgb
                    
                    # 注意：这里不复制高亮背景，实现去除高亮的效果
                    # 原来的代码：
                    # if run.font.highlight_color is not None:
                    #     new_run.font.highlight_color = run.font.highlight_color
                    # 现在跳过高亮背景的复制
            else:
                # 如果没有runs，直接复制文本
                new_para.add_run(para.text)
        
        # 删除第一个空段落
        if template_doc.paragraphs and not template_doc.paragraphs[0].text.strip():
            p = template_doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # 保存到输出路径
        template_doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"复制内容到模板时出错：{e}")
        return False
    
# ===== 双出处处理函数 =====
def detect_dual_source_outline(doc):
    """
    检测文档是否为双出处纲目
    判断标准：文档开头是否包含"节期纲目的出处"
    """
    # 检查前几个段落（通常在开头）
    for i, para in enumerate(doc.paragraphs[:8]):  # 检查前8段
        if "节期纲目的出处" in para.text:
            return True
    return False

def add_dual_source_marker_to_title(doc):
    """
    在篇题之后添加"（纲目的双出处）"标记
    """
    if len(doc.paragraphs) >= 3:
        # 第三段是篇题
        title_para = doc.paragraphs[2]
        current_text = title_para.text.strip()
        
        # 检查是否已经包含双出处标记，避免重复添加
        if "（纲目的双出处）" not in current_text:
            # 添加双出处标记
            new_text = current_text + "（纲目的双出处）"
            title_para.text = new_text
            return True
    return False
def copy_content_to_template(source_doc_path, template_path, output_path):
    """
    将源文档的内容复制到模板文档中
    """
    try:
        # 打开源文档和模板
        source_doc = Document(source_doc_path)
        template_doc = Document(template_path)
        
        # 清空模板的现有内容（保留样式）
        # 删除所有段落，但保留第一个空段落作为起点
        for i in range(len(template_doc.paragraphs) - 1, 0, -1):
            p = template_doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # 清空第一个段落
        if template_doc.paragraphs:
            template_doc.paragraphs[0].clear()
        
        # 复制源文档的所有段落到模板
        for para in source_doc.paragraphs:
            # 在模板中创建新段落
            new_para = template_doc.add_paragraph()
            
            # 复制段落内容和基本格式
            if para.runs:
                # 如果原段落有runs，逐个复制
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    # 复制基本格式
                    if run.font.bold is not None:
                        new_run.font.bold = run.font.bold
                    if run.font.italic is not None:
                        new_run.font.italic = run.font.italic
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
                    if run.font.name is not None:
                        new_run.font.name = run.font.name
                    # 复制颜色
                    if run.font.color.rgb is not None:
                        new_run.font.color.rgb = run.font.color.rgb
                    # 复制高亮
                    if run.font.highlight_color is not None:
                        new_run.font.highlight_color = run.font.highlight_color
            else:
                # 如果没有runs，直接复制文本
                new_para.add_run(para.text)
        
        # 删除第一个空段落
        if template_doc.paragraphs and not template_doc.paragraphs[0].text.strip():
            p = template_doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # 保存到输出路径
        template_doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"复制内容到模板时出错：{e}")
        return False

def process_dual_source_with_template_no_highlight(filename, script_dir):
    """
    使用模板处理双出处纲目文档，并去除高亮背景
    """
    source_path = os.path.join(script_dir, filename)
    template_path = os.path.join(script_dir, 'template.docx')
    
    # 检查template.docx是否存在
    if not os.path.exists(template_path):
        print(f"警告：未找到template.docx文件，将按普通方式处理 {filename}")
        return None
    
    # 生成临时文件名
    temp_filename = f"temp_{filename}"
    temp_path = os.path.join(script_dir, temp_filename)
    
    # 复制内容到模板（已包含去除高亮的功能）
    if copy_content_to_template_with_highlight_removal(source_path, template_path, temp_path):
        print(f"已使用template.docx处理双出处文档并去除高亮：{filename}")
        return temp_path
    else:
        print(f"使用模板处理失败，将按普通方式处理：{filename}")
        return None
    
def convert_outline_spaces_to_tabs(doc):
    """
    将段首纲目结构的半角空格转换为制表符
    只处理段首的特定纲目模式，保持其他内容不变
    """
    # 定义段首纲目模式（使用^确保在段首，\s+匹配一个或多个空格）
    outline_patterns = [
        r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\s+',      # 段首：大写中文数字 + 空格
        r'^([一二三四五六七八九十百千万亿]+)\s+',                    # 段首：小写中文数字 + 空格  
        r'^(\d+)\s+',                                              # 段首：阿拉伯数字 + 空格
        r'^([a-z])\s+',                                            # 段首：小写字母 + 空格
        r'^\(([一二三四五六七八九十百千万亿]+)\)\s+',                 # 段首：(中文数字) + 空格
        r'^\((\d+)\)\s+',                                          # 段首：(阿拉伯数字) + 空格
        r'^\(([a-z])\)\s+',                                        # 段首：(小写字母) + 空格
    ]
    
    converted_count = 0
    
    for para in doc.paragraphs:
        original_text = para.text
        
        # 只处理非空段落
        if not original_text.strip():
            continue
            
        # 检查每个段首纲目模式并替换
        for pattern in outline_patterns:
            if re.match(pattern, original_text):
                # 将匹配的空格替换为制表符
                modified_text = re.sub(pattern, lambda m: m.group(0).replace(' ', '\t', 1), original_text)
                if modified_text != original_text:
                    para.text = modified_text
                    converted_count += 1
                break  # 找到匹配后跳出循环
    
    return converted_count


# ===== 讲者检查函数 =====
def check_speaker_for_dictation(text, filename=""):
    """
    检查听抄稿的纲目后是否有讲者缩写
    返回: (是否有讲者, 处理后的文本)
    """
    # 检查是否包含"（听抄稿的纲目）"
    if "（听抄稿的纲目）" not in text:
        return True, text  # 不是听抄稿，跳过检查
    
    # 查找"（听抄稿的纲目）"的位置
    pattern_pos = text.find("（听抄稿的纲目）")
    after_pattern = text[pattern_pos + len("（听抄稿的纲目）"):].strip()
    
    # 检查后面是否有讲者格式，支持：
    # 单人：（M. C.）、(R. K.)
    # 多人：（A. B.,C. D.）、(M. C., R. K.)
    # 集体：(Various)、（Various）
    speaker_pattern = r'^\s*[（(]\s*(?:[A-Z]\.\s*[A-Z]\.\s*(?:,\s*[A-Z]\.\s*[A-Z]\.\s*)*|Various)\s*[）)]\s*'
    has_speaker = bool(re.match(speaker_pattern, after_pattern))
    
    return has_speaker, text

def show_speaker_reminder(third_para_text=""):
    """显示讲者提醒窗口，包含第三行内容（无 GUI 时仅 print）"""
    if tk is None or messagebox is None:
        print("提醒：听抄稿请在 (听抄稿的纲目) 后面添加讲者缩写，例如：(M. C.)")
        return
    root = tk.Tk()
    root.withdraw()
    if third_para_text:
        message = f"您忘记加讲者了！\n\n当前第三行内容：\n{third_para_text}\n\n请在 (听抄稿的纲目) 后面添加讲者缩写，\n例如：(听抄稿的纲目)(M. C.)"
    else:
        message = "您忘记加讲者了！\n\n请在 (听抄稿的纲目) 后面添加讲者缩写，\n例如：(听抄稿的纲目)(M. C.)"
    messagebox.showwarning("提醒", message, icon="warning")
    root.destroy()

# ===== 背景色处理函数 =====
def has_yellow_background(run):
    """检查run是否有黄色背景 - 只检测背景高亮"""
    try:
        # 检查标准高亮
        if run.font.highlight_color is not None:
            highlight = run.font.highlight_color
            # 这里可以更精确地检测是否为黄色高亮
            from docx.enum.text import WD_COLOR_INDEX
            if hasattr(WD_COLOR_INDEX, 'YELLOW') and highlight == WD_COLOR_INDEX.YELLOW:
                return True
            # 或者假设任何高亮都是黄色背景（如果您确定只使用黄色高亮）
            return True
        
        # 检查XML中的高亮设置
        if hasattr(run, '_element'):
            rPr = run._element.find('.//w:rPr', run._element.nsmap)
            if rPr is not None:
                highlight_elem = rPr.find('.//w:highlight', rPr.nsmap)
                if highlight_elem is not None:
                    highlight_val = highlight_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if highlight_val:
                        if highlight_val.lower() in ['yellow', 'lightyellow']:
                            return True
        
        return False
    except Exception as e:
        return False

def apply_yellow_background(run):
    """为run应用黄色背景"""
    try:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except:
        # 如果高亮不可用，尝试其他方法
        pass

def preserve_text_simple(para, new_text):
    """简化版本：如果整个段落有黄色背景，记录但先不处理"""
    # 检查段落是否有黄色背景
    has_any_background = any(has_yellow_background(run) for run in para.runs)
    
    # 先正常更新文本
    para.text = new_text
    
    # 返回背景状态信息
    return has_any_background

def apply_custom_92_style(para):
    """应用自定义的92级标题格式"""
    # 设置字体
    for run in para.runs:
        run.font.name = '方正书宋_GBK'
        run.font.size = Pt(14)  # 四号字体
        run.font.bold = True
    
    # 设置段落格式
    from docx.shared import Cm
    para_format = para.paragraph_format
    para_format.left_indent = Cm(0.5)    # 左缩进1字符（约0.35英寸）
    para_format.right_indent = Cm(0.5)   # 右缩进1字符（约0.35英寸）
    para_format.first_line_indent = Cm(-1)     # 悬挂缩进：首行向左偏移1厘米
    para_format.space_before = Pt(0.3 * 12)   # 段前0.3行（0.3 * 12磅）
    para_format.space_after = Pt(0.3 * 12)    # 段后0.3行（0.3 * 12磅）
    para_format.line_spacing = Pt(23)         # 行距23磅


def create_scripture_styles_fixed(doc):
    """
    简化版：由于经节内容统一使用2大点样式，这个函数现在只做验证
    """
    styles = doc.styles
    
    try:
        # 检查2大点样式是否存在
        two_points_style = styles["2大点"]
        return {"2大点": two_points_style}
    except KeyError:
        return {}

def apply_scripture_outline_styles_fixed(doc):
    """
    修复版：经节内容统一应用2大点样式，纲目行应用对应样式
    """
    
    # 纲目识别模式和对应的样式
    outline_patterns_with_styles = [
        (r'^读经：', '11读经'),
        (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t', '2大点'),
        (r'^([一二三四五六七八九十百千万亿]+)\t', '3中点'),
        (r'^(\d+)\t', '4小点'),
        (r'^([a-z])\t', '5a点'),
        (r'^\(([一二三四五六七八九十百千万亿]+)\)\t', '6（一）'),
        (r'^\((\d+)\)\t', '7（1）'),
        (r'^\(([a-z])\)\t', '8（a）'),
        # 特殊字符直接应用对应的样式
        (r'^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]', '4小点'),  # 与数字同样式
        (r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]', '7（1）'),  # 与(1)同样式
        (r'^[㈠㈡㈢㈣㈤㈥㈦㈧㈨㈩]', '6（一）'),              # 与(一)同样式
        (r'^[ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟⓠⓡⓢⓣⓤⓥⓦⓧⓨⓩ]', '8（a）'),  # 与(a)同样式
        (r'^[ⒶⒷⒸⒹⒺⒻⒼⒽⒾⒿⓀⓁⓂⓃⓄⓅⓆⓇⓈⓉⓊⓋⓌⓍⓎⓏ]', '8（a）'),  # 与(A)同样式
    ]
    
    # 用于跟踪是否在经节区域
    in_scripture_section = False
    applied_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:  # 跳过空段落
            continue
            
        # 检查是否匹配纲目模式
        is_outline = False
        for pattern, style_name in outline_patterns_with_styles:
            if re.match(pattern, text):
                is_outline = True
                in_scripture_section = True
                # 对纲目行应用对应的样式
                try:
                    para.style = style_name
                except:
                    pass
                break
        
        # 如果不是纲目行，且在经节区域内，则应用2大点样式
        if not is_outline and in_scripture_section:
            # 检查是否是其他结构性内容
            if any(marker in text for marker in ['职事信息摘录：', '听抄信息：', '晨兴圣言信息：']):
                in_scripture_section = False
                continue
            
            try:
                # 经节内容统一应用2大点样式
                para.style = "2大点"
                applied_count += 1
            except Exception as e:
                pass
    
    return applied_count


def force_scripture_font_final(doc):
    """调整使用2大点样式的经节段落的字体和缩进"""
    font_fixed_count = 0
    
    # 用于跟踪是否在经节区域和当前缩进级别
    in_scripture_section = False
    current_indent_level = 1  # 默认缩进级别
    
    # 纲目识别模式和对应的缩进级别
    outline_patterns = [
        (r'^读经：', 1),
        (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t', 1),      # 1级缩进
        (r'^([一二三四五六七八九十百千万亿]+)\t', 7),                    # 2级缩进
        (r'^(\d+)\t', 2),                                              # 2级缩进
        (r'^([a-z])\t', 3),                                            # 3级缩进
        (r'^\(([一二三四五六七八九十百千万亿]+)\)\t', 4),                 # 4级缩进
        (r'^\((\d+)\)\t', 5),                                          # 5级缩进
        (r'^\(([a-z])\)\t', 6),                                        # 6级缩进
    ]
    
    # 缩进级别对应的具体缩进值（cm）
    indent_mapping = {
        1: 1.0,    # 1级经节
        2: 1.75,   # 2级经节
        3: 2.0,    # 3级经节
        4: 2.45,   # 4级经节
        5: 2.75,   # 5级经节
        6: 3.25,   # 6级经节
        7: 1.5,   # 6级经节
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:  # 跳过空段落
            continue
            
        # 检查是否匹配纲目模式，更新当前缩进级别
        is_outline = False
        for pattern, level in outline_patterns:
            if re.match(pattern, text):
                is_outline = True
                in_scripture_section = True  # 发现纲目，进入经节区域
                current_indent_level = level  # 更新当前缩进级别
                break
        
        # 如果不是纲目行，且在经节区域内，且使用2大点样式，则调整字体和缩进
        if not is_outline and in_scripture_section:
            # 检查是否是其他结构性内容，如果是则退出经节区域
            if any(marker in text for marker in ['职事信息摘录：', '听抄信息：', '晨兴圣言信息：']):
                in_scripture_section = False
                continue
            
            # 检查是否使用2大点样式
            if para.style and para.style.name == "2大点":
                # 1. 调整字体为楷体
                for run in para.runs:
                    run.font.name = '方正楷体_GBK'  # 楷体
                    run.font.size = Pt(11)         # 11号字体
                    run.font.bold = False          # 不加粗
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                
                # 2. 调整段落缩进格式
                target_indent_cm = indent_mapping.get(current_indent_level, 1.75)  # 默认1.75cm
                hanging_distance = 2.75  # 悬挂距离固定为2.75cm
                total_left_indent = target_indent_cm + hanging_distance
                
                para_format = para.paragraph_format
                para_format.left_indent = Cm(total_left_indent)           # 总的左侧缩进
                para_format.first_line_indent = Cm(-hanging_distance)     # 首行向左突出（悬挂缩进）
                para_format.right_indent = Cm(0)                          # 右侧缩进0
                para_format.line_spacing = Pt(18)                         # 行距固定值18磅
                para_format.space_before = Pt(0)                          # 段前0
                para_format.space_after = Pt(0)                           # 段后0
                
                font_fixed_count += 1
    
    return font_fixed_count

# ===== 文件名处理函数 =====
def clean_text_for_filename(text):
    """清理文本中的换行符和其他不可见字符，用于文件名处理"""
    # 移除各种换行符和不可见字符
    cleaned = text.replace('\n', '').replace('\r', '').replace('\x0b', '').replace('\x0c', '')
    # 移除多余的空白字符
    cleaned = ' '.join(cleaned.split())
    return cleaned.strip()

# ===== 原有的文本处理函数 =====
def fullwidth_to_halfwidth(text):
    result = ''
    for char in text:
        code = ord(char)
        if 0xFF10 <= code <= 0xFF19:
            char = chr(code - 0xFF10 + ord('0'))
        elif 0xFF21 <= code <= 0xFF3A:
            char = chr(code - 0xFF21 + ord('A'))
        elif 0xFF41 <= code <= 0xFF5A:
            char = chr(code - 0xFF41 + ord('a'))
        result += char
    return result

def correct_chinese_quotes_nesting_improved(text):
    """
    直接处理中文引号，不用临时标记
    规则：双引号内的双引号改为单引号
    """
    # 定义引号字符
    left_double = '\u201c'   # "
    right_double = '\u201d'  # "
    left_single = '\u2018'   # '
    right_single = '\u2019'  # '
    
    result = []
    quote_level = 0  # 跟踪嵌套层次
    
    for char in text:
        if char == left_double:
            # 遇到左双引号
            if quote_level == 0:
                result.append(left_double)  # 第一层，保持双引号
            else:
                result.append(left_single)  # 嵌套层，改为单引号
            quote_level += 1
            
        elif char == right_double:
            # 遇到右双引号
            quote_level -= 1
            if quote_level == 0:
                result.append(right_double)  # 第一层结束，保持双引号
            else:
                result.append(right_single)  # 嵌套层结束，改为单引号
                
        else:
            result.append(char)
    
    return ''.join(result)


def replace_english_punctuation(text, preserve_asterisk=False):
    """
    改进版本的标点符号处理函数，包含优化的引号嵌套处理
    
    Args:
        text: 待处理的文本
        preserve_asterisk: 是否保留星号（双出处纲目时为True）
    """
    # 基础标点符号映射（移除括号，单独处理）
    punctuation_mapping = {
        '\u2018': '\u201c', '\u2019': '\u201d',  # 使用Unicode编码：' -> ", ' -> "
        '\'': '"', '\'': '"',
        '∶': '：', ';': '；',',': '，', ';': '；', ':': '：', '?': '？', '!': '！',
        '[': '【', ']': '】', '{': '｛', '}': '｝', '#': '',
        '经节出处：': '读经：', '经文出处：': '读经：', '今日晨兴/今日纲目/页首': '', 
        '─': '—', '——': '—','—': '—','─引用经文': '', '—引用经文': '','～': '~', '晨兴-大纲|Outline|对照-听抄-目录': '',
        
        # 带圆圈的中文数字映射
        '㈠': '(一)', '㈡': '(二)', '㈢': '(三)', '㈣': '(四)', '㈤': '(五)',
        '㈥': '(六)', '㈦': '(七)', '㈧': '(八)', '㈨': '(九)', '㈩': '(十)',

        
        # 带圆圈的阿拉伯数字映射
        '①': '(1)', '②': '(2)', '③': '(3)', '④': '(4)', '⑤': '(5)',
        '⑥': '(6)', '⑦': '(7)', '⑧': '(8)', '⑨': '(9)', '⑩': '(10)',
        '⑪': '(11)', '⑫': '(12)', '⑬': '(13)', '⑭': '(14)', '⑮': '(15)',
        '⑯': '(16)', '⑰': '(17)', '⑱': '(18)', '⑲': '(19)', '⑳': '(20)',
        
        # 带圆圈的英文字母映射（大写）
        'Ⓐ': '(A)', 'Ⓑ': '(B)', 'Ⓒ': '(C)', 'Ⓓ': '(D)', 'Ⓔ': '(E)', 'Ⓕ': '(F)',
        'Ⓖ': '(G)', 'Ⓗ': '(H)', 'Ⓘ': '(I)', 'Ⓙ': '(J)', 'Ⓚ': '(K)', 'Ⓛ': '(L)',
        'Ⓜ': '(M)', 'Ⓝ': '(N)', 'Ⓞ': '(O)', 'Ⓟ': '(P)', 'Ⓠ': '(Q)', 'Ⓡ': '(R)',
        'Ⓢ': '(S)', 'Ⓣ': '(T)', 'Ⓤ': '(U)', 'Ⓥ': '(V)', 'Ⓦ': '(W)', 'Ⓧ': '(X)',
        'Ⓨ': '(Y)', 'Ⓩ': '(Z)',
        
        # 带圆圈的英文字母映射（小写）
        'ⓐ': '(a)', 'ⓑ': '(b)', 'ⓒ': '(c)', 'ⓓ': '(d)', 'ⓔ': '(e)', 'ⓕ': '(f)',
        'ⓖ': '(g)', 'ⓗ': '(h)', 'ⓘ': '(i)', 'ⓙ': '(j)', 'ⓚ': '(k)', 'ⓛ': '(l)',
        'ⓜ': '(m)', 'ⓝ': '(n)', 'ⓞ': '(o)', 'ⓟ': '(p)', 'ⓠ': '(q)', 'ⓡ': '(r)',
        'ⓢ': '(s)', 'ⓣ': '(t)', 'ⓤ': '(u)', 'ⓥ': '(v)', 'ⓦ': '(w)', 'ⓧ': '(x)',
        'ⓨ': '(y)', 'ⓩ': '(z)','。。': '。', '：：': '：',
    }
    
    # 先处理基础标点符号
    for en_punc, zh_punc in punctuation_mapping.items():
        text = text.replace(en_punc, zh_punc)
    
    # 处理星号：根据preserve_asterisk参数决定是否保留
    if not preserve_asterisk:
        text = text.replace('*', '')
    
    # 处理句号和连字符
    text = re.sub(r'(?<![A-Za-z])\.(?![A-Za-z])', '。', text)
    text = re.sub(r'(?<![A-Za-z])\-(?![A-Za-z])', '～', text)

    # 处理括号 - 基于纯中文数字且在段首的转换
    def replace_parentheses(text):
        """
        改进版：处理所有类型的括号组合，支持中文数字、阿拉伯数字、英文字母
        """
        # 字符集定义
        chinese_numbers = '零一二三四五六七八九十百千万亿壹贰叁肆伍陆柒捌玖拾佰仟'
        arabic_numbers = '0123456789'
        english_letters = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        
        # 检查是否为纯中文数字
        def is_pure_chinese_number(content):
            content = content.strip()
            return content and all(char in chinese_numbers for char in content)
        
        # 检查是否为纯阿拉伯数字
        def is_pure_arabic_number(content):
            content = content.strip()
            return content and all(char in arabic_numbers for char in content)
        
        # 检查是否为纯英文字母
        def is_pure_english_letter(content):
            content = content.strip()
            return content and all(char in english_letters for char in content)
        
        # 检查是否为纲目标记内容（中文数字、阿拉伯数字或英文字母）
        def is_outline_marker(content):
            return (is_pure_chinese_number(content) or 
                    is_pure_arabic_number(content) or 
                    is_pure_english_letter(content))
        
        # 检查括号是否在段首
        def is_at_line_start(start_pos):
            text_before = text[:start_pos]
            lines = text_before.split('\n')
            current_line_before = lines[-1] if lines else ''
            return current_line_before.strip() == ''
        
        # 统一匹配所有括号组合：() （） (） （)
        bracket_pattern = r'([（(])([^）)]*?)([）)])'
        
        def replace_bracket_match(match):
            content = match.group(2)
            start_pos = match.start()
            
            # 如果不是纲目标记内容，统一使用中文括号
            if not is_outline_marker(content):
                return f'（{content}）'
            
            # 如果是纲目标记内容，根据位置决定括号类型
            if is_at_line_start(start_pos):
                return f'({content})'  # 段首：英文括号
            else:
                return f'（{content}）'  # 非段首：中文括号
        
        return re.sub(bracket_pattern, replace_bracket_match, text)

    text = replace_parentheses(text)

    # 改进的引号处理函数
    def replace_quotes_improved(text):
        # 定义中文引号字符 - 使用Unicode编码
        left_double = '\u201c'   # "
        right_double = '\u201d'  # "
        left_single = '\u2018'   # '
        right_single = '\u2019'  # '
        
        # 首先统一所有英文双引号为标准双引号
        text = re.sub(r'["""]', '"', text)
        # 统一所有英文单引号为标准单引号
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        
        # 处理双引号配对 - 使用栈来跟踪引号状态
        result = []
        quote_stack = []
        i = 0
        
        while i < len(text):
            char = text[i]
            if char == '"':
                # 如果栈为空或者栈顶不是双引号，说明这是开始引号
                if not quote_stack or quote_stack[-1] != 'double':
                    result.append(left_double)
                    quote_stack.append('double')
                else:
                    # 栈顶是双引号，说明这是结束引号
                    result.append(right_double)
                    quote_stack.pop()
            elif char == "'":
                # 检查是否有配对的单引号
                next_single = text.find("'", i + 1)
                if next_single != -1:
                    # 找到配对，添加中文单引号
                    result.append(left_single)
                    result.append(text[i+1:next_single])
                    result.append(right_single)
                    i = next_single  # 跳到配对引号位置
                else:
                    # 没有配对，保持原样
                    result.append(char)
            else:
                result.append(char)
            i += 1
        
        return ''.join(result)
    
    # 第一步：英文引号转中文引号
    text = replace_quotes_improved(text)
    
    # 第二步：使用改进的中文引号嵌套矫正
    text = correct_chinese_quotes_nesting_improved(text)
    
    return text

# ===== 5.5版本新增：句中句号改分号功能（参考5.1版本） =====
def is_outline_paragraph(para):
    """
    判断段落是否为纲目结构（而非内容部分）
    只对纲目行进行标点处理，完全跳过内容段落
    直接使用detect_outline_level函数来判断，确保一致性
    """
    level = detect_outline_level(para.text.strip())
    is_outline = level is not None
    return is_outline

def is_in_protected_content_area(doc, para_index):
    """
    判断段落是否在需要保护的内容区域（听抄、晨兴、职事摘录）
    这些区域完全不进行标点处理
    """
    protected_markers = ['听抄信息：', '晨兴圣言信息：', '职事信息摘录：']
    
    # 向前查找是否在保护区域内
    for i in range(para_index - 1, -1, -1):
        if i < len(doc.paragraphs):
            para_text = doc.paragraphs[i].text.strip()
            
            # 如果遇到保护标记，说明当前段落在保护区域内
            for marker in protected_markers:
                if marker in para_text:
                    return True
                    
            # 如果遇到新的纲目标识符，说明已经脱离保护区域
            if is_outline_paragraph(doc.paragraphs[i]):
                return False
    
    return False

def convert_mid_sentence_periods_to_semicolons(text):
    """
    将句中的句号改为中文分号（保留句末句号）
    只处理中间位置的句号，不处理句末句号
    """
    # 使用正则表达式匹配句中的句号（句号后面还有内容）
    # 负前瞻 (?!$) 确保句号不在行尾
    # 负前瞻 (?!\s*$) 确保句号后不是只有空白字符
    text = re.sub(r'。(?!\s*$)', '；', text)
    
    return text

def delete_empty_paragraphs(doc):
    empty_paras = [para for para in doc.paragraphs if not para.text.strip()]
    for para in empty_paras:
        p_element = para._element
        p_element.getparent().remove(p_element)

def extract_category_info_with_dual_source(third_para_text, filename="", is_dual_source=False):
    """
    修改版的类型信息提取函数，支持双出处处理
    """
    # 首先检查听抄稿的讲者
    has_speaker, processed_text = check_speaker_for_dictation(third_para_text, filename)
    if not has_speaker:
        raise FileSkipException(f"缺少讲者信息")
    
    category_patterns = {
        '（听抄稿的纲目）': '【听抄稿的纲目】',
        '（晨兴信息选读的纲目）': '【晨兴信息选读的纲目】',
        '（复合的纲目）': '【复合的纲目】',
        '（纲目的原文）': '【纲目的原文】',
        '（带经文的纲目）': '【带经文的纲目】',
        '（纲目的双出处）': '【纲目的双出处】',
    }
    
    category_prefix = ''
    speaker_info = ""
    
    # 如果是双出处纲目，直接设置前缀
    if is_dual_source:
        category_prefix = '【纲目的双出处】'
        # 移除文本中的双出处标记，因为要用于文件名
        processed_text = processed_text.replace('（纲目的双出处）', '').strip()
        return category_prefix, processed_text, speaker_info
    
    # 以下是原有的extract_category_info函数的内容，直接复制过来
    # 优先级1：检查是否以任何模式结尾（但需要考虑讲者）
    for pattern, replacement in category_patterns.items():
        if pattern == '（听抄稿的纲目）':
            # 特殊处理听抄稿：需要考虑后面可能有讲者
            pattern_pos = processed_text.find(pattern)
            if pattern_pos != -1:
                after_pattern = processed_text[pattern_pos + len(pattern):].strip()
                # 检查是否只有讲者信息 - 修改正则以捕获完整的括号格式
                speaker_only_pattern = r'^\s*([（(]\s*[A-Z]\.\s*[A-Z]\.\s*[）)])\s*$'
                speaker_match = re.match(speaker_only_pattern, after_pattern)
                
                if speaker_match:
                    # 有讲者信息 - 保持原始括号格式
                    category_prefix = replacement
                    speaker_info = speaker_match.group(1)  # 使用完整匹配，保持原始括号
                    processed_text = processed_text[:pattern_pos].strip()
                    return category_prefix, processed_text, speaker_info
                elif after_pattern == '':
                    # 没有讲者信息
                    category_prefix = replacement
                    processed_text = processed_text[:pattern_pos].strip()
                    return category_prefix, processed_text, speaker_info
        else:
            # 其他模式的正常处理
            if processed_text.endswith(pattern):
                category_prefix = replacement
                processed_text = processed_text[:-len(pattern)].strip()
                return category_prefix, processed_text, speaker_info
    
    # 优先级2：找最后出现的模式
    last_match = None
    last_index = -1
    
    for pattern, replacement in category_patterns.items():
        index = processed_text.rfind(pattern)
        if index > last_index:
            last_index = index
            last_match = {'pattern': pattern, 'replacement': replacement, 'index': index}
    
    if last_match:
        category_prefix = last_match['replacement']
        pattern_start = last_match['index']
        pattern_end = pattern_start + len(last_match['pattern'])
        
        if last_match['pattern'] == '（听抄稿的纲目）':
            # 特殊处理：也要移除后面的讲者信息
            remaining_text = processed_text[pattern_end:]
            # 修改正则以捕获完整的括号格式 - 支持多种讲者格式
            speaker_pattern = r'^\s*([（(]\s*(?:[A-Z]\.\s*[A-Z]\.\s*(?:,\s*[A-Z]\.\s*[A-Z]\.\s*)*|Various)\s*[）)])'
            match = re.match(speaker_pattern, remaining_text)
            if match:
                speaker_info = match.group(1)  # 使用完整匹配，保持原始括号
                pattern_end += match.end()
        
        processed_text = (processed_text[:pattern_start] + processed_text[pattern_end:]).strip()
        return category_prefix, processed_text, speaker_info
    
    return '', processed_text, speaker_info


# 自定义异常类，用于跳过文件
class FileSkipException(Exception):
    pass

def show_final_error_summary(skipped_files, error_files):
    """在所有文件处理完成后显示错误总结（无 GUI 时仅 print）"""
    if not skipped_files and not error_files:
        return
    message_parts = ["文件处理完成！"]
    if skipped_files:
        message_parts.append(f"以下 {len(skipped_files)} 个文件被跳过：")
        for filename, reason in skipped_files:
            message_parts.append(f"  • {filename} - {reason}")
    if error_files:
        message_parts.append(f"以下 {len(error_files)} 个文件处理出错：")
        for filename, error in error_files:
            message_parts.append(f"  • {filename} - {error}")
    final_message = "\n".join(message_parts)
    if tk is None or messagebox is None:
        print(final_message)
        return
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("处理结果总结", final_message, icon="info")
    root.destroy()

# ===== 新增：红色字体检测和应用函数 =====
def has_red_font(run):
    """判断run中的文字是否为红色 - 只检测字体颜色，不检测高亮"""
    try:
        if run.font.color.rgb:
            rgb_color = run.font.color.rgb
            hex_color = str(rgb_color)
            
            # 将十六进制字符串转换为RGB值
            if len(hex_color) == 6:  # 确保是有效的6位十六进制
                try:
                    red_value = int(hex_color[0:2], 16)
                    green_value = int(hex_color[2:4], 16)
                    blue_value = int(hex_color[4:6], 16)
                    
                    # 检测红色：红色值高，绿蓝值低
                    if red_value > 200 and green_value < 100 and blue_value < 100:
                        return True
                    else:
                        return False
                except ValueError as e:
                    pass
        
        return False
            
    except Exception as e:
        return False

def apply_red_font(run):
    """为run应用红色字体"""
    try:
        run.font.color.rgb = RGBColor(255, 0, 0)  # 纯红色
    except:
        pass
def has_red_font_in_paragraph(para):
    """检测段落中是否包含红色字体"""
    for run in para.runs:
        if has_red_font(run):
            return True
    return False

def has_red_font_strict(run):
    """严格检测红色字体 - 只有明确的红色RGB值才返回True"""
    try:
        if run.font.color.rgb:
            rgb_color = run.font.color.rgb
            hex_color = str(rgb_color)
            
            # 只有明确的红色十六进制值才被认为是红色
            known_red_colors = ['FF0000', 'ff0000', 'C00000', 'c00000', '800000']
            
            if hex_color.upper() in [color.upper() for color in known_red_colors]:
                return True
            
            # 或者使用RGB值检测
            if len(hex_color) == 6:
                try:
                    red_value = int(hex_color[0:2], 16)
                    green_value = int(hex_color[2:4], 16)
                    blue_value = int(hex_color[4:6], 16)
                    
                    # 更严格的红色检测：红色值必须很高，绿蓝值必须很低
                    if red_value >= 200 and green_value <= 50 and blue_value <= 50:
                        return True
                    else:
                        return False
                except ValueError as e:
                    pass
        
        return False
            
    except Exception as e:
        return False

# ===== 修改：扩展原有的背景存储函数以支持红色字体 =====
def store_paragraph_formatting(para):
    """存储段落的详细格式信息，精确到每个字符"""
    # 存储每个字符的格式信息
    char_formatting = []
    char_index = 0
    
    for run_idx, run in enumerate(para.runs):
        run_text = run.text
        has_background = has_yellow_background(run)
        has_red = has_red_font(run)
        
        # 为这个run中的每个字符记录格式
        for char in run_text:
            char_formatting.append({
                'char': char,
                'position': char_index,
                'has_background': has_background,
                'has_red_font': has_red
            })
            char_index += 1
    
    return {
        'original_text': para.text,
        'char_formatting': char_formatting
    }


def restore_paragraph_formatting(para, formatting_info):
    """精确恢复每个字符的格式"""
    if not formatting_info or not formatting_info.get('char_formatting'):
        return
    
    current_text = para.text
    char_formatting = formatting_info['char_formatting']
    
    # 清空段落内容
    para.clear()
    
    # 按格式分组连续的字符
    current_group = {'text': '', 'has_background': False, 'has_red_font': False}
    
    for i, char in enumerate(current_text):
        # 找到对应的格式信息
        char_format = None
        if i < len(char_formatting):
            char_format = char_formatting[i]
        
        # 判断是否需要开始新的run
        need_new_run = False
        if char_format:
            if (char_format['has_background'] != current_group['has_background'] or 
                char_format['has_red_font'] != current_group['has_red_font']):
                need_new_run = True
        
        # 如果需要新run或者这是最后一个字符
        if need_new_run or i == len(current_text) - 1:
            # 先处理当前组
            if current_group['text']:
                run = para.add_run(current_group['text'])
                if current_group['has_background']:
                    apply_yellow_background(run)
                if current_group['has_red_font']:
                    apply_red_font(run)
            
            # 开始新组
            current_group = {
                'text': char,
                'has_background': char_format['has_background'] if char_format else False,
                'has_red_font': char_format['has_red_font'] if char_format else False
            }
        else:
            # 继续当前组
            current_group['text'] += char
    
    # 处理最后一组
    if current_group['text']:
        run = para.add_run(current_group['text'])
        if current_group['has_background']:
            apply_yellow_background(run)
        if current_group['has_red_font']:
            apply_red_font(run)

def preserve_text_simple(para, new_text):
    """简化版本：记录黄色背景和红色字体状态"""
    # 检查段落是否有黄色背景和红色字体
    has_any_background = any(has_yellow_background(run) for run in para.runs)
    has_any_red_font = any(has_red_font(run) for run in para.runs)  # 新增
    
    # 先正常更新文本
    para.text = new_text
    
    # 返回格式状态信息
    return {
        'has_background': has_any_background,
        'has_red_font': has_any_red_font  # 新增
    }

# 添加辅助函数：从存储的格式信息中检测红色字体
def has_red_font_in_stored_formatting(formatting_info):
    """从存储的格式信息中检测是否包含红色字体"""
    if not formatting_info or not formatting_info.get('char_formatting'):
        return False
    
    for char_info in formatting_info['char_formatting']:
        if char_info.get('has_red_font', False):
            return True
    return False
def detect_outline_level(text: str):
    """匹配并返回纲目层级（int）；匹配不到返回 None。使用性能优化器缓存"""
    return optimizer.get_outline_level_cached(text)

# 句尾可替换为冒号的标点（中英标点）
PUNCT_AT_END_RE = Config.PUNCT_AT_END_RE

# ===== 纲目起始匹配（所有可能的"下一级"开头）=====
OUTLINE_START_PATTERNS = [
    r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t',   # 2 大点（大写中文数字）
    r'^([一二三四五六七八九十百千万亿]+)\t',             # 3 中点（小写中文数字）
    r'^(\d+)\t',                                         # 4 小点（阿拉伯数字）
    r'^([a-z])\t',                                       # 5 a点（字母）
    r'^\(([一二三四五六七八九十百千万亿]+)\)\t',          # 6 （一）
    r'^\((\d+)\)\t',                                     # 7 （1）
    r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]\t',            # 7 带圈数字
    r'^\(([a-z])\)\t',                                   # 8 （a）
]
OUTLINE_START_RE = re.compile(r'(?:' + '|'.join(OUTLINE_START_PATTERNS) + r')')

# 若你还没定义它，请保留这个“句尾标点”正则（已定义就不要重复定义）
try:
    PUNCT_AT_END_RE
except NameError:
    PUNCT_AT_END_RE = re.compile(r'[。！？；.,!?;,，]$')

# 句尾可替换为冒号的标点（中英标点）
PUNCT_AT_END_RE = re.compile(r'[。！？；.,!?;,，]$')

def process_transcript_additions_underline(doc):
    """
    处理复合纲目中的【听抄稿添加开始】和【听抄稿添加结束】标记
    将标记之间的内容设置下划线，并删除标记本身。支持简繁体。
    """
    start_markers = ("【听抄稿添加开始】", "【聽抄稿添加開始】")
    end_markers = ("【听抄稿添加结束】", "【聽抄稿添加結束】")
    in_section = False
    processed_count = 0
    paragraphs_to_delete = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        start_found = any(m in text for m in start_markers)
        end_found = any(m in text for m in end_markers)
        if start_found:
            in_section = True
            new_text = text
            for m in start_markers:
                new_text = new_text.replace(m, "")
            new_text = new_text.replace("。", "").strip()
            if not new_text:
                paragraphs_to_delete.append(i)
            else:
                para.text = new_text
            processed_count += 1
        elif end_found:
            new_text = text
            for m in end_markers:
                new_text = new_text.replace(m, "")
            new_text = new_text.replace("。", "").strip()
            if not new_text:
                paragraphs_to_delete.append(i)
            else:
                para.text = new_text
            in_section = False
            processed_count += 1
        elif in_section and text.strip():
            for run in para.runs:
                run.font.underline = True
            processed_count += 1
    for i in reversed(paragraphs_to_delete):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    return processed_count


def process_compound_outline_additions(doc):
    """
    处理复合纲目中的【添加开始】和【添加结束】标记（晨兴添加）
    将标记之间的内容设置为斜体，并删除标记本身。
    【听抄稿添加开始】～【听抄稿添加结束】由 process_transcript_additions_underline 处理（下划线）。
    """
    in_addition_section = False
    start_marker = "【添加开始】"
    end_marker = "【添加结束】"
    processed_count = 0
    paragraphs_to_delete = []  # 记录需要删除的段落
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 检查是否包含开始标记
        if start_marker in text:
            in_addition_section = True
            # 删除开始标记
            new_text = text.replace(start_marker, "").replace("。", "").strip()
            
            if not new_text:  # 如果删除后段落为空，标记为待删除
                paragraphs_to_delete.append(i)
            else:
                para.text = new_text
            processed_count += 1
        
        # 检查是否包含结束标记
        elif end_marker in text:
            # 删除结束标记
            new_text = text.replace(end_marker, "").replace("。", "").strip()
            
            if not new_text:  # 如果删除后段落为空，标记为待删除
                paragraphs_to_delete.append(i)
            else:
                para.text = new_text
            in_addition_section = False
            processed_count += 1
        
        # 如果在标记区间内，设置斜体
        elif in_addition_section and text.strip():
            for run in para.runs:
                run.font.italic = True
            processed_count += 1
    
    # 删除空段落（从后往前删除，避免索引问题）
    for i in reversed(paragraphs_to_delete):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    return processed_count


# Web 下载建议文件名：格式「【类型】序号 内容.docx」，不含讲者
OUTLINE_TYPE_TO_PREFIX = {
    'original': '【纲目的原文】',
    'with_scripture': '【带经文的纲目】',
    'morning_revival': '【晨兴信息选读的纲目】',
    'transcript': '【听抄稿的纲目】',
    'composite': '【复合的纲目】',
}
# 第三段末尾可能带的类型后缀，用于从正文中剥掉
_TYPE_SUFFIX_PATTERNS = ['（纲目的原文）', '（带经文的纲目）', '（晨兴信息选读的纲目）', '（听抄稿的纲目）', '（复合的纲目）']


def suggest_feast_outline_filename(third_para_text, outline_type):
    """
    Web 用：根据第三段文本和纲目类型生成建议下载文件名，格式「【类型】序号 内容.docx」。
    不含讲者。若解析失败返回 None，调用方用默认名。
    """
    if not third_para_text or not isinstance(third_para_text, str):
        return None
    prefix = OUTLINE_TYPE_TO_PREFIX.get(outline_type)
    if not prefix:
        return None
    try:
        third_para = clean_text_for_filename(third_para_text)
        for suf in _TYPE_SUFFIX_PATTERNS:
            if third_para.endswith(suf):
                third_para = third_para[:-len(suf)].strip()
                break
        if not third_para:
            return None
        # 智能分割：全角空格或空格分出序号与内容
        if '　' in third_para:
            parts = third_para.split('　', 1)
        elif ' ' in third_para:
            match = re.match(r'^(第\s*[^篇章课]*[篇章课])\s+(.+)$', third_para)
            if match:
                parts = [match.group(1).strip(), match.group(2).strip()]
            else:
                parts = third_para.split(' ', 1)
        else:
            parts = [third_para]
        if len(parts) >= 2:
            serial_part = fullwidth_to_halfwidth(parts[0].strip())
            content_part = parts[1].strip()
            pattern = r'^第\s*([0-9]+|[一二三四五六七八九十百千万亿壹贰叁肆伍陆柒捌玖拾佰仟萬億]+)\s*([章篇课])$'
            match = re.match(pattern, serial_part)
            if match:
                cn_num = match.group(1)
                try:
                    if cn2an:
                        arabic_num = cn2an.cn2an(cn_num, 'smart')
                        serial_str = f"msg. {arabic_num}"
                    else:
                        serial_str = serial_part
                except (ValueError, TypeError):
                    serial_str = serial_part
            else:
                serial_str = serial_part
            invalid_chars = r'[\/:*?"<>|]'
            content_part = re.sub(invalid_chars, '', content_part)
            if not content_part:
                return None
            return f"{prefix}{serial_str} {content_part}.docx"
        else:
            invalid_chars = r'[\/:*?"<>|]'
            content_part = re.sub(invalid_chars, '', third_para)
            if not content_part:
                return None
            return f"{prefix}{content_part}.docx"
    except Exception:
        return None


# ===== Web 入口：按类型刷格式（无双出处） =====
FEAST_OUTLINE_TYPES = ('original', 'with_scripture', 'morning_revival', 'transcript', 'composite')


def format_feast_outline_docx(docx_path, outline_type):
    """
    Web 用：对已生成的节期纲目 DOCX 按类型刷格式（不重命名文件）。
    outline_type: 'original' | 'with_scripture' | 'morning_revival' | 'transcript' | 'composite'
    """
    if outline_type not in FEAST_OUTLINE_TYPES:
        raise ValueError(f"不支持的节期纲目类型: {outline_type}")
    doc = Document(docx_path)
    delete_empty_paragraphs(doc)
    is_scripture_outline = (outline_type == 'with_scripture')
    is_compound_outline = (outline_type == 'composite')
    in_dictation_mode = (outline_type == 'transcript')
    if is_scripture_outline:
        create_scripture_styles_fixed(doc)
    process_one_doc(doc, is_scripture_outline, is_compound_outline, in_dictation_mode, is_dual_source_outline=False)
    doc.save(docx_path)


def process_one_doc(doc, is_scripture_outline, is_compound_outline, in_dictation_mode, is_dual_source_outline=False):
    """
    对单个文档执行全部刷格式步骤（不包含加载/保存/重命名）。
    双出处时 preserve_asterisk=True；Web 调用时传 is_dual_source_outline=False。
    """
    # 在开始处理前，先记录所有段落的背景状态
    paragraph_formatting = []
    for para in doc.paragraphs:
        formatting_info = store_paragraph_formatting(para)
        paragraph_formatting.append(formatting_info)

    # 处理1：正则表达式替换
    for para in doc.paragraphs:
        original_text = para.text
        new_text = re.sub(r'^([A-Za-z0-9]+)\.(　|\t)', r'\1\2', original_text)
        if new_text != original_text:
            para.text = new_text

    convert_outline_spaces_to_tabs(doc)

    # 处理2：标点符号处理
    for para in doc.paragraphs:
        original_text = para.text
        text = fullwidth_to_halfwidth(original_text)
        text = replace_english_punctuation(text, preserve_asterisk=is_dual_source_outline)
        if text != original_text:
            para.text = text

    # 处理2.5：句中句号改分号
    for i, para in enumerate(doc.paragraphs):
        if is_in_protected_content_area(doc, i):
            continue
        if is_outline_paragraph(para):
            original_text = para.text
            new_text = convert_mid_sentence_periods_to_semicolons(original_text)
            if new_text != original_text:
                para.text = new_text

    # 处理3：分号处理
    for para in doc.paragraphs:
        original_text = para.text
        if original_text.rstrip().endswith('；'):
            new_text = original_text.rstrip()[:-1] + '。'
            para.text = new_text

    # 处理4：空格和制表符处理
    after_marker = False
    for para in doc.paragraphs:
        text = para.text
        if any(marker in text for marker in after_marker_list):
            after_marker = True
        original_text = text
        text = text.lstrip('　\t')
        if not after_marker:
            text = text.replace('　', '\t')
        if text != original_text:
            para.text = text

    # 处理5：字符替换处理
    for para in doc.paragraphs:
        original_text = para.text
        text = original_text.replace('篇\t', '篇　').replace('章\t', '章　').replace('课\t', '课　')
        text = text.replace('貮\t', '贰\t').replace('参\t', '叁\t').replace('貳\t', '贰\t')
        text = text.replace('彀', '够').replace('\t\t', '\t').replace('～', '~')
        if text != original_text:
            para.text = text

    # 处理6：句号添加处理
    if not is_scripture_outline:
        after_marker = False
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.rstrip()
            if any(marker in text for marker in after_marker_list):
                after_marker = True
            has_red_font = False
            if idx < len(paragraph_formatting):
                has_red_font = has_red_font_in_stored_formatting(paragraph_formatting[idx])
            if ('【添加开始】' in text or '【添加结束】' in text or
                '【听抄稿添加开始】' in text or '【听抄稿添加结束】' in text or
                '【聽抄稿添加開始】' in text or '【聽抄稿添加結束】' in text):
                continue
            if not after_marker and idx >= 4 and not has_red_font:
                if not text.endswith(('。', '！', '？', '…', '”', '\'', '：', '』')):
                    para.text = text + '。'
                elif text.endswith('：'):
                    para.text = text[:-1] + '。'

        # 处理7：冒号处理
        n = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            this_text = (para.text or '').strip()
            if ('【添加开始】' in this_text or '【添加结束】' in this_text or
                '【听抄稿添加开始】' in this_text or '【听抄稿添加结束】' in this_text or
                '【聽抄稿添加開始】' in this_text or '【聽抄稿添加結束】' in this_text):
                continue
            this_level = detect_outline_level(this_text)
            if this_level is None:
                continue
            j = i + 1
            child_found = False
            while j < n:
                next_text = (doc.paragraphs[j].text or '').strip()
                if not next_text:
                    j += 1
                    continue
                next_level = detect_outline_level(next_text)
                if next_level is None:
                    j += 1
                    continue
                if next_level > this_level:
                    child_found = True
                break
            if child_found:
                has_red = False
                if 0 <= i < len(paragraph_formatting):
                    has_red = has_red_font_in_stored_formatting(paragraph_formatting[i])
                if not has_red:
                    new_text = (para.text or '').rstrip()
                    if not new_text.endswith('：'):
                        if PUNCT_AT_END_RE.search(new_text):
                            new_text = PUNCT_AT_END_RE.sub('：', new_text)
                        else:
                            new_text += '：'
                        new_text = re.sub(r'：{2,}$', '：', new_text)
                        para.text = new_text

    # 处理8：李常受文集斜体处理
    if len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        text = last_para.text
        matches = re.findall(r'（(.*?)）', text)
        if matches:
            inner_text = matches[-1]
            sub_match = re.search(r'(李常受文集.*?册)', inner_text)
            if sub_match:
                target_text = sub_match.group(1)
                total_text = ''.join(run.text for run in last_para.runs)
                start_index = total_text.find(target_text)
                end_index = start_index + len(target_text)
                if start_index != -1:
                    current_index = 0
                    new_runs = []
                    for run in last_para.runs:
                        run_text = run.text
                        run_length = len(run_text)
                        run_start = current_index
                        for idx in range(run_length):
                            char_pos = run_start + idx
                            char = run_text[idx]
                            if start_index <= char_pos < end_index:
                                new_runs.append((char, True))
                            else:
                                new_runs.append((char, run.font.italic))
                        current_index += run_length
                    last_para.clear()
                    for text_char, is_italic in new_runs:
                        run = last_para.add_run(text_char)
                        run.font.italic = is_italic

    # 处理9：样式应用
    after_marker = False
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if any(marker in text for marker in after_marker_list):
            after_marker = True
        elif after_marker:
            if not in_dictation_mode:
                if not text.endswith(('。', '！', '？', '…', '"', '\'', '）', '：', '』')):
                    para.style = "81级标题"

    # 处理10：职事信息摘录
    for para in doc.paragraphs:
        if any(marker in para.text for marker in after_marker_list):
            para.style = "9职事信息摘录"
            break

    # 处理11：读经样式
    for para in doc.paragraphs:
        if '读经：' in para.text:
            para.style = "11读经"
            original_text = para.text
            new_text = re.sub(r'[；;]', '，', original_text)
            if new_text != original_text:
                para.text = new_text

    # 处理12：标题样式应用
    if is_scripture_outline:
        pattern_styles = [
            (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)　', '81级标题'),
            (r'^([一二三四五六七八九十百千万亿]+)　', '82级标题'),
            (r'^(\d+)　', '83级标题'),
            (r'^([a-z])　', '84级标题'),
            (r'^（([一二三四五六七八九十百千万亿]+)）　', '84级标题'),
            (r'^（(\d+)）　', '84级标题'),
            (r'^序言[ \t　]+', '2大点'),
            (r'^添言[ \t　]+', '2大点'),
            (r'^前言[ \t　]+', '3中点'),
            (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t', '2大点'),
            (r'^([一二三四五六七八九十百千万亿]+)\t', '3中点'),
            (r'^(\d+)\t', '4小点'),
            (r'^([a-z])\t', '5a点'),
            (r'^\(([一二三四五六七八九十百千万亿]+)\)\t', '6（一）'),
            (r'^\((\d+)\)\t', '7（1）'),
            (r'^\(([a-z])\)\t', '8（a）'),
        ]
        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern, style in pattern_styles:
                if re.match(pattern, text):
                    para.style = style
                    break
        apply_scripture_outline_styles_fixed(doc)
    elif in_dictation_mode:
        pattern_styles = [
            (r'^序言[ \t　]+', '2大点'),
            (r'^添言[ \t　]+', '2大点'),
            (r'^前言[ \t　]+', '3中点'),
            (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t', '2大点'),
            (r'^([一二三四五六七八九十百千万亿]+)\t', '3中点'),
            (r'^(\d+)\t', '4小点'),
            (r'^([a-z])\t', '5a点'),
            (r'^\(([一二三四五六七八九十百千万亿]+)\)\t', '6（一）'),
            (r'^\((\d+)\)\t', '7（1）'),
            (r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]\t', '7（1）'),
            (r'^\(([a-z])\)\t', '8（a）'),
        ]
        after_dictation_marker = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if '听抄信息：' in text:
                after_dictation_marker = True
                continue
            if after_dictation_marker and in_dictation_mode:
                if '。' not in text:
                    apply_custom_92_style(para)
                    continue
            for pattern, style in pattern_styles:
                if re.match(pattern, text):
                    para.style = style
                    break
    else:
        pattern_styles = [
            (r'^序言[ \t　]+', '2大点'),
            (r'^添言[ \t　]+', '2大点'),
            (r'^前言[ \t　]+', '3中点'),
            (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)\t', '2大点'),
            (r'^([一二三四五六七八九十百千万亿]+)\t', '3中点'),
            (r'^(\d+)\t', '4小点'),
            (r'^([a-z])\t', '5a点'),
            (r'^\(([一二三四五六七八九十百千万亿]+)\)\t', '6（一）'),
            (r'^\((\d+)\)\t', '7（1）'),
            (r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]\t', '7（1）'),
            (r'^\(([a-z])\)\t', '8（a）'),
        ]
        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern, style in pattern_styles:
                if re.match(pattern, text):
                    para.style = style
                    break

    # 设置前三段的固定样式
    style_names = ["0系列", "11111西列", "00篇题"]
    for i in range(min(3, len(doc.paragraphs))):
        doc.paragraphs[i].style = style_names[i]

    # 处理13：修正数字间的逗号为顿号
    for para in doc.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        modified_text = re.sub(r'(?<=\d)，(?=\d)', '、', full_text)
        if modified_text != full_text:
            para.clear()
            para.add_run(modified_text)

    # 格式恢复
    for i, para in enumerate(doc.paragraphs):
        if i < len(paragraph_formatting):
            formatting_info = paragraph_formatting[i]
            if formatting_info and formatting_info.get('char_formatting'):
                restore_paragraph_formatting(para, formatting_info)

    if is_compound_outline:
        process_transcript_additions_underline(doc)  # 听抄稿添加：下划线
        process_compound_outline_additions(doc)       # 晨兴添加：斜体

    for para in doc.paragraphs:
        if any(marker in para.text for marker in after_marker_list):
            page_break_para = para.insert_paragraph_before()
            run = page_break_para.add_run()
            run.add_break(WD_BREAK.PAGE)
            break

    if in_dictation_mode:
        pattern_styles_custom92 = [
            (r'^([壹贰貳叁參肆伍陆陸柒捌玖拾佰仟萬万億亿]+)　', 'custom_92'),
            (r'^([一二三四五六七八九十百千万亿]+)　', 'custom_92'),
            (r'^(\d+)　', 'custom_92'),
            (r'^([a-z])　', 'custom_92'),
            (r'^（([一二三四五六七八九十百千万亿]+)）　', 'custom_92'),
            (r'^（(\d+)）　', 'custom_92'),
        ]
        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern, style in pattern_styles_custom92:
                if re.match(pattern, text):
                    apply_custom_92_style(para)
                    break
        after_dictation_marker = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if '听抄信息：' in text:
                after_dictation_marker = True
                continue
            if after_dictation_marker and '。' not in text:
                apply_custom_92_style(para)

    after_marker = False
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if any(marker in text for marker in after_marker_list):
            after_marker = True
        elif after_marker:
            if not in_dictation_mode:
                if not text.endswith(('。', '！', '？', '…', '"', '\'', '）', '：', '』')):
                    para.style = "81级标题"

    if is_scripture_outline:
        force_scripture_font_final(doc)


# ===== 修改后的主处理逻辑（桌面版：含双出处与文件重命名） =====
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    excluded_files = {'template.docx'}
    word_files = [f for f in os.listdir(script_dir) if f.endswith('.docx') and not f.startswith('~$') and f not in excluded_files]

    processed_files = []
    skipped_files = []
    error_files = []

    for filename in word_files:
        is_template_processed = False  # 标记是否使用了模板处理
        
        try:
            filepath = os.path.join(script_dir, filename)
            
            # 先检查是否为双出处文档
            temp_doc = Document(filepath)
            if detect_dual_source_outline(temp_doc):
                # 使用模板处理
                template_processed_path = process_dual_source_with_template_no_highlight(filename, script_dir)
                if template_processed_path:
                    filepath = template_processed_path
                    is_template_processed = True
                    print(f"使用template.docx处理双出处文档：{filename}")
            
            # 加载文档（可能是模板处理后的）
            doc = Document(filepath)

            delete_empty_paragraphs(doc)

            # 检查文档类型
            is_scripture_outline = False
            is_dual_source_outline = False  # 新增双出处标志
            is_compound_outline = False

            if len(doc.paragraphs) >= 3:
                third_para_text = doc.paragraphs[2].text
                if '（带经文的纲目）' in third_para_text:
                    is_scripture_outline = True
                    # 创建自定义经节样式
                    create_scripture_styles_fixed(doc)
                elif '（复合的纲目）' in third_para_text:  # 添加这个条件
                    is_compound_outline = True

            # 新增：检测双出处纲目
            is_dual_source_outline = detect_dual_source_outline(doc)
            if is_dual_source_outline:
                add_dual_source_marker_to_title(doc)

            # 提前检查是否在听抄信息模式
            in_dictation_mode = False
            for para in doc.paragraphs:
                if '听抄信息：' in para.text:
                    in_dictation_mode = True
                    break

            process_one_doc(doc, is_scripture_outline, is_compound_outline, in_dictation_mode, is_dual_source_outline)


            # 文件重命名逻辑 - 修复版本（包含讲者检查）
            new_filename = filename  # 默认保持原文件名
            if len(doc.paragraphs) >= 3:
                # 获取第三段文本并清理不可见字符
                raw_text = doc.paragraphs[2].text
                third_para = clean_text_for_filename(raw_text)

                # 使用类型提取函数（现在包含讲者检查）
                try:
                    category_prefix, third_para, speaker_info = extract_category_info_with_dual_source(third_para, filename, is_dual_source_outline)
                    
                    # 智能分割处理
                    if '　' in third_para:
                        parts = third_para.split('　')
                    elif ' ' in third_para:
                        match = re.match(r'^(第\s*[^篇章课]*[篇章课])\s+(.+)$', third_para)
                        if match:
                            parts = [match.group(1).strip(), match.group(2).strip()]
                        else:
                            parts = third_para.split(' ', 1)
                    else:
                        parts = [third_para]
                    
                    if len(parts) >= 2:
                        serial_part = parts[0].strip()
                        content_part = parts[1].strip()
                        
                        # 将序号部分的全角字符转换为半角
                        serial_part = fullwidth_to_halfwidth(serial_part)
                        
                        # 提取序号并转换为阿拉伯数字
                        pattern = r'^第\s*([0-9]+|[一二三四五六七八九十百千万亿壹贰叁肆伍陆柒捌玖拾佰仟萬億]+)\s*([章篇课])$'
                        match = re.match(pattern, serial_part)
                        
                        if match:
                            cn_num = match.group(1)
                            try:
                                if cn2an:
                                    arabic_num = cn2an.cn2an(cn_num, 'smart')
                                    serial_str = f"msg. {arabic_num}"
                                else:
                                    serial_str = serial_part
                            except (ValueError, TypeError):
                                serial_str = serial_part
                        else:
                            serial_str = serial_part

                        # 移除文件名中不允许的字符
                        invalid_chars = r'[\/:*?"<>|]'
                        content_part = re.sub(invalid_chars, '', content_part)

                        # 组合新的文件名，包含讲者信息
                        if speaker_info and "听抄稿" not in category_prefix:  # 听抄稿不包含讲者信息
                            new_filename = f"{category_prefix}{serial_str} {content_part}{speaker_info}.docx"
                        else:
                            new_filename = f"{category_prefix}{serial_str} {content_part}.docx"
                    else:
                        # 组合新的文件名，包含讲者信息
                        if speaker_info and "听抄稿" not in category_prefix:  # 听抄稿不包含讲者信息
                            new_filename = f"{category_prefix}{third_para}{speaker_info}.docx"
                        else:
                            new_filename = f"{category_prefix}{third_para}.docx"
                            
                except FileSkipException as e:
                    skipped_files.append((filename, str(e)))
                    continue

            # 保存文件
            if filename == 'Comparison.docx':
                original_path = os.path.join(script_dir, filename)
                new_path = os.path.join(script_dir, new_filename)
                doc.save(new_path)
            else:
                new_path = os.path.join(script_dir, new_filename)
                doc.save(new_path)
            # 清理临时文件
            if is_template_processed:
                try:
                    if os.path.exists(filepath) and 'temp_' in os.path.basename(filepath):
                        os.remove(filepath)
                        print(f"已清理临时文件：{os.path.basename(filepath)}")
                except Exception as e:
                    print(f"清理临时文件失败：{e}")
            processed_files.append((filename, new_filename))
            
        except FileSkipException as e:
            skipped_files.append((filename, str(e)))
            continue
        except Exception as e:
            error_files.append((filename, str(e)))
            continue


    processed_skipped_files = []
    for filename, reason in skipped_files:
        clean_reason = remove_duplicate_chinese_punctuation(reason)
        processed_skipped_files.append((filename, clean_reason))

    # 处理 error_files 中的重复标点  
    processed_error_files = []
    for filename, error in error_files:
        clean_error = remove_duplicate_chinese_punctuation(error)
        processed_error_files.append((filename, clean_error))

    # 统一弹出错误总结窗口
    show_final_error_summary(processed_skipped_files, processed_error_files)


    # 统一弹出错误总结窗口
    show_final_error_summary(skipped_files, error_files)
    
    # 5.5版本新增：显示性能统计信息
    total_files = len(word_files)
    if DEBUG_MODE or total_files > 5:  # 处理较多文件时显示统计
        cache_stats = optimizer.get_cache_stats()
        print(f"\n===== 5.5版本性能统计 =====")
        print(f"缓存命中率: {cache_stats['hit_rate']}")
        print(f"缓存请求总数: {cache_stats['total_requests']}")
        print(f"缓存大小: {cache_stats['cache_size']}/{Config.MAX_CACHE_SIZE}")
        print("===============================")

# 主程序入口
if __name__ == "__main__":
    main()