import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import winsound
import cn2an  # 确保已安装该模块
from docx.enum.text import WD_BREAK  # 导入分页符常量
import tkinter as tk
from tkinter import messagebox
import threading
import time

# 在文件顶部添加这个函数
def show_error_dialog(title, message):
    """显示错误对话框，确保在最顶层并居中"""
    def create_dialog():
        # 创建主窗口但不显示
        root = tk.Tk()
        root.withdraw()  # 隐藏主窗口
        
        # 设置对话框属性
        root.attributes('-topmost', True)  # 确保在最顶层
        root.lift()  # 提升窗口层级
        root.focus_force()  # 强制获取焦点
        
        # 显示消息框
        messagebox.showerror(title, message)
        
        # 销毁主窗口
        root.destroy()
    
    # 在新线程中创建对话框，避免阻塞主程序
    dialog_thread = threading.Thread(target=create_dialog)
    dialog_thread.daemon = True
    dialog_thread.start()
    dialog_thread.join()  # 等待对话框关闭


def get_header_count(doc):
    """动态检测文档开头的标题段落数（1-3段）。
    优先查找"读经："段落的索引；若不存在，查找第一个以"壹"开头的段落；默认3。
    """
    for idx, para in enumerate(doc.paragraphs):
        if '读经：' in para.text or '讀經：' in para.text:
            return idx
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('壹'):
            return idx
    return 3


# 定义一个函数，将全角数字和字母转换为半角
def fullwidth_to_halfwidth(text):
    result = ''
    for char in text:
        code = ord(char)
        # 转换全角数字
        if code >= 0xFF10 and code <= 0xFF19:
            char = chr(code - 0xFF10 + ord('0'))
        # 转换全角大写字母
        elif code >= 0xFF21 and code <= 0xFF3A:
            char = chr(code - 0xFF21 + ord('A'))
        # 转换全角小写字母
        elif code >= 0xFF41 and code <= 0xFF5A:
            char = chr(code - 0xFF41 + ord('a'))
        result += char
    return result

# 修改后的函数
def replace_english_punctuation(text):
    punctuation_mapping = {
        ',': '，',
        ';': '；',
        ':': '：',
        '?': '？',
        '!': '！',
        '(': '（',
        ')': '）',
        '[': '【',
        ']': '】',
        '{': '｛',
        '}': '｝',
    }

    # 替换除 '.' 和 '-' 之外的英文标点为中文标点
    for en_punc, zh_punc in punctuation_mapping.items():
        text = text.replace(en_punc, zh_punc)

    # 特殊处理 '.'，仅当它不在英文字母之间时替换为全角句号 '。'
    text = re.sub(r'(?<![A-Za-z])\.(?![A-Za-z])', '。', text)

    # 将所有直破折号替换为长破折号
    text = text.replace('-', '—')

    # 处理单引号和双引号的替换，注意区分前后引号
    def replace_quotes(text, quote_char, left_quote, right_quote):
        parts = text.split(quote_char)
        result = ''
        for i, part in enumerate(parts):
            result += part
            if i < len(parts) - 1:
                if i % 2 == 0:
                    result += left_quote
                else:
                    result += right_quote
        return result

    text = replace_quotes(text, '"', '“', '”')
    text = replace_quotes(text, "'", '‘', '’')
    return text

# 定义一个函数，删除空白段落
def delete_empty_paragraphs(doc):
    empty_paras = []
    for para in doc.paragraphs:
        if not para.text.strip():
            empty_paras.append(para)

    for para in empty_paras:
        p_element = para._element
        p_element.getparent().remove(p_element)

    # 定义一个函数，将中文单引号改为中文双引号
def replace_chinese_single_quotes(text):
    text = text.replace('‘', '“')
    text = text.replace('’', '”')
    return text

# 以下保持您原有的代码逻辑
# 获取脚本所在的目录
script_dir = os.path.dirname(os.path.abspath(__file__))

# 获取所有的Word文档（.docx），并排除以~$开头的临时文件
word_files = [f for f in os.listdir(script_dir) if f.endswith('.docx') and not f.startswith('~$')]

# 遍历所有的Word文档
for filename in word_files:
    filepath = os.path.join(script_dir, filename)
    doc = Document(filepath)
    header_count = get_header_count(doc)

    # 步骤3：删除空白段落
    delete_empty_paragraphs(doc)


    # 调用函数
    delete_empty_paragraphs(doc)

    # 步骤1和2：替换全角数字字母、英文标点为中文标点
    for para in doc.paragraphs:
        text = para.text
        text = fullwidth_to_halfwidth(text)  # 转换全角数字和字母为半角
        text = replace_english_punctuation(text)  # 替换英文标点为中文标点
        para.text = text  # 更新段落文本

    # 步骤4：将中文单引号改为中文双引号
    for para in doc.paragraphs:
        text = para.text
        text = replace_chinese_single_quotes(text)  # 替换中文单引号为双引号
        para.text = text

    # 步骤5：将段落结尾的分号改为句号
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if text.endswith('；'):
            text = text[:-1] + '。'  # 将分号替换为句号
            para.text = text
  
    # 步骤6：处理全角空格和半角空格
    after_marker = False  # 标记是否已经到达"职事信息摘录："
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True  # 遇到标记，之后的段落不再替换空格

        # 删除段首的全角空格和半角空格
        text = text.lstrip('　\t ')

        # 在标记之前，将全角空格和半角空格替换为Tab键，但跳过前三段
        if not after_marker and idx >= header_count:
            text = text.replace('　', '\t')  # 替换全角空格
            text = text.replace(' ', '\t')   # 替换半角空格
        para.text = text

    # 步骤7：替换篇、章、课后的Tab为全角空格
    for para in doc.paragraphs:
        text = para.text
        text = text.replace('篇\t', '篇　')
        text = text.replace('章\t', '章　')
        text = text.replace('课\t', '课　')
        text = text.replace('貮\t', '贰\t')
        text = text.replace('参\t', '叁\t')
        text = text.replace('貳\t', '贰\t')
        text = text.replace('参　', '叁　')   # 参+全角空格
        text = text.replace('貳　', '贰　')   # 繁体贰+全角空格
        text = text.replace('貮　', '贰　')   # 异体贰+全角空格
        text = text.replace('陸\t', '陆\t')   # 繁体陆+Tab
        text = text.replace('陸　', '陆　')   # 繁体陆+全角空格
        text = text.replace('億\t', '亿\t')   # 繁体亿+Tab
        text = text.replace('億　', '亿　')   # 繁体亿+全角空格
        text = text.replace('彀', '够')
        # 替换连续的两个Tab键为一个
        text = text.replace('\t\t', '\t')
        # 将全角的'～'替换为半角的'~'
        text = text.replace('～', '~')        
        para.text = text



    # 步骤8：在段落结尾添加句号（在标记之前，且不包括前三段）
    after_marker = False
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.rstrip()
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True  # 遇到标记
        if not after_marker and idx >= header_count + 1:
            # 如果段落结尾没有标点符号，添加句号
            if not text.endswith(('。', '！', '？', '…', '”', '’', '：', '』')):
                text += '。'
                para.text = text
        # 如果段落结尾是分号（'；'），替换为句号
            elif text.endswith('：'):
                text = text[:-1] + '。'  # 将分号替换为句号
                para.text = text         

    # 步骤9：如果段落以一、1、a加Tab键开头，替换上一段结尾的标点为冒号
    prev_para = None
    for para in doc.paragraphs:
        text = para.text
        if text.startswith('一\t') or text.startswith('1\t') or text.startswith('a\t'):
            if prev_para is not None:
                prev_text = prev_para.text.rstrip()
                # 检查上一段是否以标点符号结尾，且不是冒号
                if prev_text.endswith(('。', '！', '？', '；', '.', '!', '?', ';', ',')) and not prev_text.endswith('：'):
                    # 将结尾的标点符号替换为冒号
                    prev_text = re.sub(r'[。！？；.,!?;]$', '：', prev_text)
                # 检查是否有连续的冒号，并将它们替换为一个
                prev_text = re.sub(r'：{2,}$', '：', prev_text)
                prev_para.text = prev_text
            elif prev_para and not prev_para.text.rstrip().endswith('：'):
                # 如果上一段没有标点符号，且结尾不是冒号，直接添加冒号
                prev_para.text = prev_para.text.rstrip() + '：'
                # 检查是否有连续的冒号，并将它们替换为一个
                prev_para.text = re.sub(r'：{2,}$', '：', prev_para.text)
        prev_para = para  # 更新上一段

    # 步骤10：处理最后一段的出处，只将最后一个括号内的“李常受文集.*?册”部分设为斜体
    if not doc.paragraphs:
        print(f"Warning: Document {filename} has no paragraphs. Skipping further processing for this document.")
        continue # Skip to the next file in the loop
    last_para = doc.paragraphs[-1]
    text = last_para.text

    # 匹配所有括号内的内容
    matches = re.findall(r'（(.*?)）', text)
    if matches:
        # 取最后一个括号内的内容
        inner_text = matches[-1]
        # 在最后一个括号内查找"李常受文集.*?册"部分
        sub_match = re.search(r'(李常受文集.*?册)', inner_text)
        if sub_match:
            target_text = sub_match.group(1)
            
            # 新增逻辑：检查前三行内容与括号内容的匹配
            if len(doc.paragraphs) >= header_count:
                # 获取前三行内容并用逗号拼接
                first_three_lines = []
                for i in range(header_count):
                    line_text = doc.paragraphs[i].text.strip()
                    if line_text:  # 只添加非空行
                        first_three_lines.append(line_text)
                
                # 用逗号拼接前三行内容
                title_content = '，'.join(first_three_lines)
                
                # 与最后一个括号内的内容进行完全匹配比较
                # if title_content != inner_text:
                    # 不匹配时的处理 - 已删除提示，继续刷格式
                    # pass
            
            # 将段落中的所有运行（Run）的文本拼接起来
            total_text = ''.join(run.text for run in last_para.runs)
            # 查找目标文本在总文本中的起始和结束位置
            start_index = total_text.find(target_text)
            end_index = start_index + len(target_text)
            if start_index != -1:
                current_index = 0
                new_runs = []
                for run in last_para.runs:
                    run_text = run.text
                    run_length = len(run_text)
                    run_start = current_index
                    run_end = current_index + run_length

                    # 判断运行文本是否与目标文本部分或全部重叠
                    if run_end <= start_index or run_start >= end_index:
                        # 运行文本在目标文本之外，直接添加
                        new_runs.append((run_text, run.font.italic))
                    else:
                        # 运行文本与目标文本有重叠，需要拆分
                        for i in range(run_length):
                            char_pos = run_start + i
                            char = run_text[i]
                            if start_index <= char_pos < end_index:
                                # 字符在目标文本范围内，设为斜体
                                new_runs.append((char, True))
                            else:
                                # 字符不在目标文本范围内，保持原格式
                                new_runs.append((char, run.font.italic))
                    current_index += run_length

                # 清空原有的运行（Run）
                last_para.clear()
                # 重新添加运行
                for text, is_italic in new_runs:
                    new_run = last_para.add_run(text)
                    new_run.font.italic = is_italic
            # else:
                # 未找到目标文本的处理 - 已删除提示，继续刷格式
                # pass
        # else:
            # 未在最后一个括号内找到"李常受文集...册"的处理 - 已删除提示，继续刷格式
            # pass
    # else:
        # 未找到任何括号的处理 - 已删除提示，继续刷格式

    # 步骤11：应用前三段的样式
    # 注意：需要确保这些样式在您的Word模板中存在
    style_names = ["0系列", "11111西列", "00篇题"]
    for i in range(min(header_count, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        para.style = style_names[i]

    # 步骤12：在标记之后，如果段落结尾没有标点符号，则应用特定样式
    after_marker = False
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True  # 遇到标记
        elif after_marker:
            # 如果段落结尾没有标点符号，应用样式：81级标题；否则应用样式：0000模板
            if not text.endswith(('。', '！', '？', '…', '\u201d', '\u201c', '\u2019', '\u2018', '）', '』', '\u0022', '\u0027')):
                para.style = "81级标题"
            else:
                para.style = "0000模板"

    # 步骤13：将"职事信息摘录："、"研读问题："、"出处与参读："、"参考与参读信息："应用样式：9职事信息摘录
    keywords = ['职事信息摘录：', '研读问题：', '出处与参读：', '参考与参读信息：', '参考与参读资料：',
                '職事信息摘錄：', '研讀問題：', '出處與參讀：', '參考與參讀信息：', '參考與參讀資料：']
    for para in doc.paragraphs:
        text = para.text
        for keyword in keywords:
            if keyword in text:
                para.style = "9职事信息摘录"
                break  # 找到匹配后不再检查其他关键词

    # 步骤14：将“读经：”应用样式：11读经，并将该段落中的所有分号替换为逗号
    for para in doc.paragraphs:
        text = para.text
        if '读经：' in text or '讀經：' in text:
            # 应用样式
            para.style = "11读经"
            # 替换段落中的所有分号为逗号（处理全角和半角分号）
            text = re.sub(r'[；;]', '，', text)
            para.clear()
            para.add_run(text)

    # 步骤15：根据标识应用相应的段落样式
    # 定义模式和对应的样式名称
    # 定义模式和对应的样式名称
    pattern_styles = [
        # 一级标题
        (r'^([壹贰貳叁参肆伍陆陸柒捌玖拾佰仟萬亿億]+)　', '81级标题'),
        # 二级标题
        (r'^([一二三四五六七八九十百千万萬亿億]+)　', '82级标题'),
        # 三级标题
        (r'^(\d+)　', '83级标题'),
        # 四级标题
        (r'^([a-z])　', '84级标题'),
        # 子级标题（中文数字）
        (r'^（([一二三四五六七八九十百千万萬亿億]+)）　', '84级标题'),
        # 子级标题（阿拉伯数字）
        (r'^（(\d+)）　', '84级标题'),
        # 大点
        (r'^([壹贰貳叁参肆伍陆陸柒捌玖拾佰仟萬亿億]+)\t', '2大点'),
        # 中点
        (r'^([一二三四五六七八九十百千万萬亿億]+)\t', '3中点'),
        # 小点
        (r'^(\d+)\t', '4小点'),
        # a 点
        (r'^([a-z])\t', '5a点'),
        # （一）
        (r'^\(([一二三四五六七八九十百千万萬亿億]+)\)\t', '6（一）'),
        # （1）
        (r'^\((\d+)\)\t', '7（1）'),
        # （a）
        (r'^\(([a-z])\)\t', '8（a）'),
    ]



    # 遍历段落，匹配模式并应用相应的样式
    for para in doc.paragraphs:
        text = para.text.strip()
        for pattern, style_name in pattern_styles:
            if re.match(pattern, text):
                para.style = style_name
                break  # 匹配成功后不再检查其他模式



    # 步骤16：如果逗号前后都是阿拉伯数字，将逗号替换为顿号
    for para in doc.paragraphs:
        # 拼接段落中所有的运行文本为一个完整的字符串
        full_text = ''.join(run.text for run in para.runs)
        
        # 使用正则表达式查找所有阿拉伯数字之间的逗号，并将其替换为顿号
        modified_text = re.sub(r'(?<=\d)，(?=\d)', '、', full_text)
        
        # 清空原来的运行（Run），重新赋值修改后的文本
        if modified_text != full_text:  # 如果有修改才进行更新
            para.clear()  # 清空段落中的运行
            para.add_run(modified_text)  # 将修改后的文本作为新的运行

    # 最后一步2：在"职事信息摘录："或"参考与参读信息："之前添加分页符
    page_break_keywords = ['职事信息摘录：', '参考与参读信息：', '参考与参读资料：',
                           '職事信息摘錄：', '參考與參讀信息：', '參考與參讀資料：']
    page_break_paras = [para for para in doc.paragraphs
                        if any(kw in para.text for kw in page_break_keywords)]
    for para in page_break_paras:
        # 在该段落之前插入一个新的段落
        page_break_para = para.insert_paragraph_before()
        # 在新段落中添加分页符
        run = page_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)


    # 保存修改后的文档，文件名根据最后一段标题内容生成
    if len(doc.paragraphs) >= header_count:
        third_para = doc.paragraphs[header_count - 1].text.strip()
        # 用全角空格分割为两部分
        parts = third_para.split('　')
        if len(parts) >= 2:
            serial_part = parts[0].strip()
            content_part = parts[1].strip()

            # 将序号部分的全角字符转换为半角
            serial_part = fullwidth_to_halfwidth(serial_part)
            print(f"第三段序号部分：'{serial_part}'")  # 调试信息

            # 提取序号并转换为阿拉伯数字
            # 可能的格式：“第3章”、“第十一篇”、“第20课”等
            pattern = r'^第\s*([0-9]+|[一二三四五六七八九十百千万亿]+)\s*([章篇课])$'
            match = re.match(pattern, serial_part)
            if match:
                cn_num = match.group(1)
                unit = match.group(2)
                print(f"提取的数字：'{cn_num}'，单位：'{unit}'")  # 调试信息
                try:
                    # 使用 cn2an 进行转换，使用 'smart' 模式以兼容中文数字和阿拉伯数字
                    arabic_num = cn2an.cn2an(cn_num, 'smart')
                    serial_str = f"msg. {arabic_num}"
                    print(f"生成的序号字符串：'{serial_str}'")  # 调试信息
                except ValueError as e:
                    # 转换失败，使用原始序号
                    print(f"数字转换错误：{e}")  # 调试信息
                    serial_str = serial_part
            else:
                # 不匹配预期格式，使用原始序号
                print("正则表达式未匹配到序号部分")  # 调试信息
                serial_str = serial_part

            # 移除文件名中不允许的字符
            invalid_chars = r'[\/:*?"<>|]'
            content_part = re.sub(invalid_chars, '', content_part)


            # 组合新的文件名
            # 组合新的文件名
            base_filename = f"{serial_str} {content_part}.docx"
            
            # 检查文件是否已存在，如果存在则递增版本号
            new_filename = base_filename
            new_filepath = os.path.join(script_dir, new_filename)
            version_num = 2
            while os.path.exists(new_filepath):
                new_filename = f"R{version_num}_{serial_str} {content_part}.docx"
                new_filepath = os.path.join(script_dir, new_filename)
                version_num += 1
                
        else:
            # 如果第三段没有全角空格，使用整个第三段作为文件名
            base_filename = f"{third_para}.docx"
            
            # 检查文件是否已存在，如果存在则递增版本号
            new_filename = base_filename
            new_filepath = os.path.join(script_dir, new_filename)
            version_num = 2
            while os.path.exists(new_filepath):
                new_filename = f"R{version_num}_{third_para}.docx"
                new_filepath = os.path.join(script_dir, new_filename)
                version_num += 1
    else:
        # 如果文档少于三段，使用原文件名
        new_filename = filename

    # 步骤100：将"研读问题："和"出处与参读："段落中的空格替换为Tab
    target_prefixes = ('研读问题：', '出处与参读：', '研讀問題：', '出處與參讀：')
    for para in doc.paragraphs:
        if para.text.strip().startswith(target_prefixes):
            for run in para.runs:
                run.text = run.text.replace(' ', '\t')

    # 保存修改后的文档，使用新的文件名
    new_filepath = os.path.join(script_dir, new_filename)
    doc.save(new_filepath)
