"""
鸟瞰纲目刷格式
专用于词典-鸟瞰纲目的 DOCX 生成与格式化。
篇头固定四段：
  第1段：主恢复真理的词典（0系列样式）
  第2段：关键词/篇题（11111西列样式）
  第3段：第三段内容（00篇题样式）
  第4段：鸟瞰类型行，如「3b　节期纲目的鸟瞰」（00篇题样式）
第5段起：读经、纲目正文（走中文纲目刷格式逻辑）
"""
from __future__ import annotations
import logging
import shutil
import tempfile
import os
from pathlib import Path
from docx import Document
from format_chinese_outline import format_chinese_outline_docx

logger = logging.getLogger(__name__)

# 篇头样式映射（与节期纲目模板一致）
_HEADER_STYLES = ["0系列", "11111西列", "00篇题", "00篇题"]

# 鸟瞰类型第四段文字
BIRD_VIEW_LINE4 = {
    "ministry": "3a\t职事信息的鸟瞰",
    "feast":    "3b\t节期纲目的鸟瞰",
}


def format_bird_view_docx(
    outline_text: str,
    keyword: str,
    bird_type: str = "feast",
) -> bytes:
    """
    将鸟瞰纲目正文生成格式化 DOCX，返回字节流。

    Args:
        outline_text: LLM 生成的纲目正文（纯文本，含「读经：」行）
        keyword:      关键词，作为第二段篇题
        bird_type:    "ministry" 或 "feast"

    Returns:
        docx 字节流
    """
    backend_dir = Path(__file__).resolve().parent

    # 1. 找模板（优先节期纲目模板）
    template_path = None
    for name in ("节期纲目模板.docx", "中文纲目模板.docx"):
        p = backend_dir / name
        if p.exists():
            template_path = p
            break
    if template_path is None:
        raise FileNotFoundError("找不到纲目模板文件（节期纲目模板.docx 或 中文纲目模板.docx）")

    # 2. 复制模板到临时文件
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    tmp_path = tmp.name
    try:
        shutil.copy2(str(template_path), tmp_path)

        # 3. 清空模板段落，写入篇头 + 正文
        doc = Document(tmp_path)
        for para in list(doc.paragraphs):
            para._element.getparent().remove(para._element)

        # 篇头四段
        line4 = BIRD_VIEW_LINE4.get(bird_type, BIRD_VIEW_LINE4["feast"])
        headers = [
            "主恢复真理的词典",
            keyword.strip() or "（篇题）",
            "第三段\t鸟瞰的纲目",
            line4,
        ]
        for text in headers:
            doc.add_paragraph(text)

        # 正文按行写入（跳过第一行如果是篇题）
        outline_text = outline_text.replace("\r\n", "\n").replace("\r", "\n")
        lines = outline_text.split("\n")
        # 跳过开头与关键词相同的行
        start = 0
        for i, line in enumerate(lines):
            if line.strip() == keyword.strip():
                start = i + 1
                break
            elif line.strip():
                break
        for line in lines[start:]:
            if line.strip():
                doc.add_paragraph(line)
            elif len(doc.paragraphs) > 4:
                doc.add_paragraph("")

        doc.save(tmp_path)

        # 4. 套用篇头样式（前四段）
        doc2 = Document(tmp_path)
        for i, style_name in enumerate(_HEADER_STYLES):
            if i < len(doc2.paragraphs):
                try:
                    doc2.paragraphs[i].style = doc2.styles[style_name]
                except KeyError:
                    logger.warning("样式不存在：%s，跳过", style_name)
        doc2.save(tmp_path)

        # 5. 调用中文纲目格式刷处理第5段起的纲目内容
        format_chinese_outline_docx(tmp_path, traditional_quotes=False)

        # 6. 读取字节流返回
        with open(tmp_path, "rb") as f:
            return f.read()

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
