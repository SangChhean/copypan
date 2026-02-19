"""
中文纲目格式刷程序 - 函数式版本
用于后端 API 调用，处理单个 DOCX 文件
"""
import os
import re
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import cn2an

logger = logging.getLogger(__name__)


def get_header_count(doc):
    """动态检测文档开头的标题段落数（1-3段）。"""
    for idx, para in enumerate(doc.paragraphs):
        if '读经：' in para.text or '讀經：' in para.text:
            return idx
    for idx, para in enumerate(doc.paragraphs):
        # 检测简体或繁体大点开头（壹、貳、參、叄等）
        text_stripped = para.text.strip()
        if text_stripped.startswith(('壹', '貳', '參', '叄', '肆', '伍', '陸', '柒', '捌', '玖', '拾')):
            return idx
    return 3


def fullwidth_to_halfwidth(text):
    """将全角数字和字母转换为半角"""
    result = ''
    for char in text:
        code = ord(char)
        if code >= 0xFF10 and code <= 0xFF19:
            char = chr(code - 0xFF10 + ord('0'))
        elif code >= 0xFF21 and code <= 0xFF3A:
            char = chr(code - 0xFF21 + ord('A'))
        elif code >= 0xFF41 and code <= 0xFF5A:
            char = chr(code - 0xFF41 + ord('a'))
        result += char
    return result


def replace_english_punctuation(text):
    """替换英文标点为中文标点"""
    punctuation_mapping = {
        ',': '，', ';': '；', ':': '：', '?': '？', '!': '！',
        '(': '（', ')': '）', '[': '【', ']': '】', '{': '｛', '}': '｝',
    }
    for en_punc, zh_punc in punctuation_mapping.items():
        text = text.replace(en_punc, zh_punc)
    text = re.sub(r'(?<![A-Za-z])\.(?![A-Za-z])', '。', text)
    text = text.replace('-', '—')
    
    def replace_quotes(text_inner, quote_char, left_quote, right_quote):
        """内部函数：替换引号"""
        parts = text_inner.split(quote_char)
        result = ''
        for i, part in enumerate(parts):
            result += part
            if i < len(parts) - 1:
                result += left_quote if i % 2 == 0 else right_quote
        return result
    
    # 替换双引号
    text = replace_quotes(text, '"', '"', '"')
    # 替换单引号
    text = replace_quotes(text, "'", '\u2018', '\u2019')  # 使用Unicode转义避免编码问题
    return text


def delete_empty_paragraphs(doc):
    """删除空白段落"""
    empty_paras = []
    for para in doc.paragraphs:
        if not para.text.strip():
            empty_paras.append(para)
    for para in empty_paras:
        p_element = para._element
        p_element.getparent().remove(p_element)


def replace_chinese_single_quotes(text):
    """将中文单引号改为中文双引号"""
    return text.replace('\u2018', '"').replace('\u2019', '"')


def _apply_style_if_exists(doc, para, style_name: str) -> None:
    """仅当模板中存在该样式时应用，避免 KeyError 导致整次格式刷失败。"""
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        logger.debug("样式不存在，跳过: %s", style_name)


def format_chinese_outline_docx(docx_path: str) -> None:
    """
    格式化中文纲目 DOCX 文件（原地修改）。
    
    Args:
        docx_path: DOCX 文件路径
    """
    logger.info(f"开始格式化中文纲目: {docx_path}")
    doc = Document(docx_path)
    header_count = get_header_count(doc)
    logger.info(f"检测到标题段落数: {header_count}, 总段落数: {len(doc.paragraphs)}")

    # 删除空白段落
    delete_empty_paragraphs(doc)
    delete_empty_paragraphs(doc)  # 重复一次确保清理干净

    # 步骤1和2：替换全角数字字母、英文标点为中文标点
    for para in doc.paragraphs:
        text = para.text
        text = fullwidth_to_halfwidth(text)
        text = replace_english_punctuation(text)
        para.text = text

    # 步骤4：将中文单引号改为中文双引号
    for para in doc.paragraphs:
        text = para.text
        text = replace_chinese_single_quotes(text)
        para.text = text

    # 步骤5：将段落结尾的分号改为句号
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if text.endswith('；'):
            text = text[:-1] + '。'
            para.text = text

    # 步骤6：处理全角空格和半角空格
    after_marker = False
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True
        text = text.lstrip('　\t ')
        if not after_marker and idx >= header_count:
            text = text.replace('　', '\t')
            text = text.replace(' ', '\t')
        para.text = text

    # 步骤7：替换篇、章、课后的Tab为全角空格
    for para in doc.paragraphs:
        text = para.text
        text = text.replace('篇\t', '篇　')
        text = text.replace('章\t', '章　')
        text = text.replace('课\t', '课　')
        text = text.replace('貮\t', '贰\t')
        text = text.replace('参\t', '叁\t')
        text = text.replace('參\t', '叁\t')  # 繁体参→简体叁
        text = text.replace('叄\t', '叁\t')  # 繁体叄→简体叁
        text = text.replace('貳\t', '贰\t')
        text = text.replace('参　', '叁　')
        text = text.replace('參　', '叁　')  # 繁体参→简体叁
        text = text.replace('叄　', '叁　')  # 繁体叄→简体叁
        text = text.replace('貳　', '贰　')
        text = text.replace('貮　', '贰　')
        text = text.replace('陸\t', '陆\t')
        text = text.replace('陸　', '陆　')
        text = text.replace('億\t', '亿\t')
        text = text.replace('億　', '亿　')
        text = text.replace('彀', '够')
        text = text.replace('\t\t', '\t')
        text = text.replace('～', '~')
        para.text = text

    # 步骤8：在段落结尾添加句号
    after_marker = False
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.rstrip()
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True
        if not after_marker and idx >= header_count + 1:
            if not text.endswith(('。', '！', '？', '…', '"', '\u2019', '：', '』')):
                text += '。'
                para.text = text
            elif text.endswith('：'):
                text = text[:-1] + '。'
                para.text = text

    # 步骤9：如果段落以一、1、a加Tab键开头，替换上一段结尾的标点为冒号
    prev_para = None
    for para in doc.paragraphs:
        text = para.text
        if text.startswith('一\t') or text.startswith('1\t') or text.startswith('a\t'):
            if prev_para is not None:
                prev_text = prev_para.text.rstrip()
                if prev_text.endswith(('。', '！', '？', '；', '.', '!', '?', ';', ',')) and not prev_text.endswith('：'):
                    prev_text = re.sub(r'[。！？；.,!?;]$', '：', prev_text)
                prev_text = re.sub(r'：{2,}$', '：', prev_text)
                prev_para.text = prev_text
            elif prev_para and not prev_para.text.rstrip().endswith('：'):
                prev_para.text = prev_para.text.rstrip() + '：'
                prev_para.text = re.sub(r'：{2,}$', '：', prev_para.text)
        prev_para = para

    # 步骤10：处理最后一段的出处，将"李常受文集.*?册"部分设为斜体
    # 注意：此步骤会清除段落样式，需要在步骤11-16中重新应用样式
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        text = last_para.text
        matches = re.findall(r'（(.*?)）', text)
        if matches:
            inner_text = matches[-1]
            sub_match = re.search(r'(李常受文集.*?册)', inner_text)
            if sub_match:
                target_text = sub_match.group(1)
                total_text = ''.join(run.text for run in last_para.runs)
                start_index = total_text.find(target_text)
                end_index = start_index + len(target_text)
                if start_index != -1:
                    current_index = 0
                    new_runs = []
                    for run in last_para.runs:
                        run_text = run.text
                        run_length = len(run_text)
                        run_start = current_index
                        run_end = current_index + run_length
                        if run_end <= start_index or run_start >= end_index:
                            new_runs.append((run_text, run.font.italic))
                        else:
                            for i in range(run_length):
                                char_pos = run_start + i
                                char = run_text[i]
                                if start_index <= char_pos < end_index:
                                    new_runs.append((char, True))
                                else:
                                    new_runs.append((char, run.font.italic))
                        current_index += run_length
                    # 保存当前样式，以便后续重新应用
                    saved_style = last_para.style.name if hasattr(last_para.style, 'name') else None
                    last_para.clear()
                    for text, is_italic in new_runs:
                        new_run = last_para.add_run(text)
                        new_run.font.italic = is_italic
                    # 注意：样式会在步骤11-16中重新应用

    # 步骤11：应用读经前标题的样式（第1段=0系列，最后一段=00篇题，中间段=11111西列）
    style_names = ["0系列", "11111西列", "00篇题"]
    n_header = min(header_count, len(doc.paragraphs))
    for i in range(n_header):
        para = doc.paragraphs[i]
        if i == 0:
            style_name = style_names[0]  # 0系列
        elif i == n_header - 1:
            style_name = style_names[2]  # 00篇题（读经前最后一段为篇题）
        else:
            style_name = style_names[1]  # 11111西列
        _apply_style_if_exists(doc, para, style_name)

    # 步骤12：在「职事信息摘录：」之后，根据段落结尾是否有标点应用样式（无标点=小标题 81级标题，有标点=正文 0000模板）
    after_marker = False
    # 使用 Unicode 转义避免引号字符导致的语法错误
    _excerpts_end_punctuation = ('。', '！', '？', '…', '"', '\u201d', '\u201c', '\u2019', '\u2018', '）', '』', '\u0022', '\u0027')
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_marker = True
        elif after_marker:
            if not text.endswith(_excerpts_end_punctuation):
                _apply_style_if_exists(doc, para, "81级标题")
            else:
                _apply_style_if_exists(doc, para, "0000模板")

    # 步骤13：将特定关键词应用样式：9职事信息摘录
    keywords = ['职事信息摘录：', '研读问题：', '出处与参读：', '参考与参读信息：', '参考与参读资料：',
                '職事信息摘錄：', '研讀問題：', '出處與參讀：', '參考與參讀信息：', '參考與參讀資料：']
    for para in doc.paragraphs:
        text = para.text
        for keyword in keywords:
            if keyword in text:
                _apply_style_if_exists(doc, para, "9职事信息摘录")
                break

    # 步骤14：将"读经："应用样式：11读经，并将该段落中的所有分号替换为逗号
    # 注意：先修改文本，再应用样式，避免样式被清除
    for para in doc.paragraphs:
        text = para.text
        if '读经：' in text or '讀經：' in text:
            text = re.sub(r'[；;]', '，', text)
            para.clear()
            para.add_run(text)
            # 在修改文本后重新应用样式
            _apply_style_if_exists(doc, para, "11读经")

    # 步骤15：如果逗号前后都是阿拉伯数字，将逗号替换为顿号（在应用样式之前，避免清除样式）
    for para in doc.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        modified_text = re.sub(r'(?<=\d)，(?=\d)', '、', full_text)
        if modified_text != full_text:
            # 只修改文本，不清除段落（避免清除样式）
            para.clear()
            para.add_run(modified_text)

    # 步骤16：根据标识应用相应的段落样式（仅对「读经」后、「职事信息摘录」前的纲目正文应用，不覆盖读经前标题和职事信息摘录段）
    pattern_styles = [
        # 大点：包含简体（壹贰叁参肆伍陆柒捌玖拾）和繁体（貳參叄陸等）
        (r'^([壹貳贰參叄叁参肆伍陸陆柒捌玖拾佰仟萬亿億]+)　', '81级标题'),
        (r'^([一二三四五六七八九十百千万萬亿億]+)　', '82级标题'),
        (r'^(\d+)　', '83级标题'),
        (r'^([a-z])　', '84级标题'),
        (r'^（([一二三四五六七八九十百千万萬亿億]+)）　', '84级标题'),
        (r'^（(\d+)）　', '84级标题'),
        (r'^([壹貳贰參叄叁参肆伍陸陆柒捌玖拾佰仟萬亿億]+)\t', '2大点'),
        (r'^([一二三四五六七八九十百千万萬亿億]+)\t', '3中点'),
        (r'^(\d+)\t', '4小点'),
        (r'^([a-z])\t', '5a点'),
        (r'^\(([一二三四五六七八九十百千万萬亿億]+)\)\t', '6（一）'),
        (r'^\((\d+)\)\t', '7（1）'),
        (r'^\(([a-z])\)\t', '8（a）'),
    ]
    after_excerpts_marker = False
    for idx, para in enumerate(doc.paragraphs):
        # 跳过读经前的标题段（由步骤11已应用 0系列/11111西列/00篇题）
        if idx < header_count:
            continue
        text = para.text.rstrip()
        # 进入职事信息摘录段后，不再用 pattern_styles 覆盖（由步骤12已应用 81级标题/0000模板）
        if '职事信息摘录：' in text or '職事信息摘錄：' in text:
            after_excerpts_marker = True
        if after_excerpts_marker:
            continue
        text_stripped = text.strip()
        for pattern, style_name in pattern_styles:
            if re.match(pattern, text_stripped):
                _apply_style_if_exists(doc, para, style_name)
                break

    # 最后一步：在特定关键词之前添加分页符
    page_break_keywords = ['职事信息摘录：', '参考与参读信息：', '参考与参读资料：',
                           '職事信息摘錄：', '參考與參讀信息：', '參考與參讀資料：']
    page_break_paras = [para for para in doc.paragraphs
                        if any(kw in para.text for kw in page_break_keywords)]
    for para in page_break_paras:
        page_break_para = para.insert_paragraph_before()
        run = page_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)

    # 步骤100：将"研读问题："和"出处与参读："段落中的空格替换为Tab
    target_prefixes = ('研读问题：', '出处与参读：', '研讀問題：', '出處與參讀：')
    for para in doc.paragraphs:
        if para.text.strip().startswith(target_prefixes):
            for run in para.runs:
                run.text = run.text.replace(' ', '\t')

    # 保存修改后的文档
    logger.info(f"保存格式化后的文档: {docx_path}")
    doc.save(docx_path)
    logger.info("中文纲目格式化完成")
