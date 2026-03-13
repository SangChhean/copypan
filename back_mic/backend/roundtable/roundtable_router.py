"""
AI 圆桌会议 API 路由
"""
import asyncio
import json
import logging
import os
import re
import tempfile
from typing import Dict, List
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field

from user.token import test_token, test_token_optional

from .roundtable_service import RoundTableService
from .roundtable_db import get_all_records, get_record_by_id, toggle_pin, delete_record

logger = logging.getLogger(__name__)

roundtable_router = APIRouter(prefix="/api")


class RoundtableStartBody(BaseModel):
    scene_type: str = Field(..., description="scene_one | scene_two | scene_three | scene_four")
    topic: str = Field(..., min_length=1)
    participants: List[str] = Field(..., min_length=1, max_length=6)
    ai_roles: Dict[str, str] = Field(default_factory=dict)


@roundtable_router.post("/roundtable/start", dependencies=[Depends(test_token)])
async def start_roundtable(body: RoundtableStartBody):
    """创建 Session，支持 scene_one（十二支派）、scene_two（神学辩论）、scene_three（重大讨论）、scene_four（顶级模型思考）。"""
    if body.scene_type not in ("scene_one", "scene_two", "scene_three", "scene_four"):
        raise HTTPException(status_code=400, detail="scene_type 须为 scene_one、scene_two、scene_three 或 scene_four")
    if body.scene_type == "scene_four":
        if len(body.participants) != 1:
            raise HTTPException(status_code=400, detail="场景④ 仅支持选择 1 个 AI")
    elif len(body.participants) < 2 or len(body.participants) > 6:
        raise HTTPException(status_code=400, detail="participants 数量须在 2～6 之间")
    svc = RoundTableService()
    session_id = await svc.create_session(
        body.scene_type,
        body.topic,
        body.participants,
        body.ai_roles,
    )
    return {"session_id": session_id}


async def _sse_stream(session_id: str):
    """根据 session 的 scene_type 选择对应场景生成器，产出 SSE。结论轮期间每 20 秒发心跳保活。"""
    svc = RoundTableService()
    try:
        session = await svc.get_session(session_id)
        st = session.get("scene_type")
        if st == "scene_one":
            generator = svc.run_scene_one(session_id)
        elif st == "scene_three":
            generator = svc.run_scene_three(session_id)
        elif st == "scene_four":
            generator = svc.run_scene_four(session_id)
        else:
            generator = svc.run_scene_two(session_id)
    except ValueError as e:
        yield f"data: {json.dumps({'type': 'error', 'reason': str(e)}, ensure_ascii=False)}\n\n"
        return

    queue = asyncio.Queue()

    async def feed_generator():
        async for event in generator:
            await queue.put(event)
        await queue.put(StopIteration)

    async def heartbeat(writer_queue):
        while True:
            await asyncio.sleep(20)
            await writer_queue.put(None)

    feed_task = asyncio.create_task(feed_generator())
    hb_task = asyncio.create_task(heartbeat(queue))

    try:
        while True:
            item = await queue.get()
            if item is StopIteration:
                break
            if item is None:
                logger.info("[RoundTable] SSE heartbeat")
                yield ": heartbeat\n\n"
                continue
            yield f"data: {json.dumps(item, ensure_ascii=False)}\n\n"
    finally:
        hb_task.cancel()
        feed_task.cancel()
        await asyncio.gather(hb_task, feed_task, return_exceptions=True)


@roundtable_router.get("/roundtable/stream/{session_id}", dependencies=[Depends(test_token_optional)])
async def stream_roundtable(session_id: str):
    """SSE 流：运行场景②，推送 speech_* / round_complete / conclusion_* 等事件。"""
    return StreamingResponse(
        _sse_stream(session_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@roundtable_router.get("/roundtable/history", dependencies=[Depends(test_token)])
async def get_history():
    """历史记录列表，仅返回摘要字段。"""
    records = get_all_records()
    return [
        {
            "record_id": r.get("record_id"),
            "scene_type": r.get("scene_type"),
            "topic": r.get("topic"),
            "participants": r.get("participants"),
            "created_at": r.get("created_at"),
            "is_pinned": r.get("is_pinned"),
            "total_cost": r.get("total_cost"),
        }
        for r in records
    ]


def _roundtable_export_text(record: dict) -> str:
    """将圆桌记录拼成导出用纯文本。"""
    topic = record.get("topic") or ""
    participants = record.get("participants") or []
    participants_str = "、".join(participants)
    rounds = record.get("rounds") or []
    conclusion = record.get("conclusion") or ""
    scene_type = record.get("scene_type") or ""
    if scene_type == "scene_two":
        scene_label = "神学辩论"
        round_titles = ["第1轮 · 亮明立场", "第2轮 · 正面交锋", "第3轮 · 总结陈词"]
    elif scene_type == "scene_three":
        scene_label = "重大讨论"
        round_titles = ["第1轮 · 作答", "第2轮 · 互相指出", "第3轮 · 最终评价"]
    elif scene_type == "scene_four":
        scene_label = "顶级模型思考"
        round_titles = ["深度思考"]
    else:
        scene_label = "十二支派"
        round_titles = ["第1步 · 各AI独立研究"]
    lines = [
        f"题目：{topic}",
        f"场景：{scene_label}",
        f"参与AI：{participants_str}",
        "",
    ]
    for idx, round_data in enumerate(rounds):
        title = round_titles[idx] if idx < len(round_titles) else f"第{idx + 1}轮"
        lines.append(f"===== {title} =====")
        lines.append("")
        if isinstance(round_data, dict):
            for ai_name, content in round_data.items():
                lines.append(f"【{ai_name}】")
                lines.append(content if content else "")
                lines.append("")
        lines.append("")
    lines.append("===== 圆桌结论 =====")
    lines.append(conclusion)
    return "\n".join(lines)


def _build_roundtable_docx(docx_path: str, record: dict) -> None:
    from docx import Document
    from docx.shared import Pt
    import os

    template_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        '中文纲目模板.docx'
    )
    doc = Document(template_path)

    # 清空模板原有段落内容
    for p in doc.paragraphs:
        for r in list(p.runs):
            r.text = ''

    topic = record.get('topic') or ''
    scene_type = record.get('scene_type') or ''
    ai_roles = record.get('ai_roles') or {}
    participants = record.get('participants') or []
    conclusion = record.get('conclusion') or ''

    # 第0段：题目，应用「0系列」样式
    p_title = doc.add_paragraph(style='0系列')
    p_title.add_run(topic)

    # 第1段：参与角色（仅场景②），应用「0000模板」样式；场景③、④无角色
    if scene_type == 'scene_two' and ai_roles:
        stances = '、'.join(
            ai_roles[ai] for ai in participants
            if ai in ai_roles and ai_roles[ai]
        )
        if stances:
            p_roles = doc.add_paragraph(style='0000模板')
            p_roles.add_run(f'参与角色：{stances}')

    # 场景④：单轮思考内容
    if scene_type == 'scene_four':
        rounds_data = record.get('rounds') or []
        if rounds_data and isinstance(rounds_data[0], dict):
            for _ai_name, content in rounds_data[0].items():
                for line in (content or '').split('\n'):
                    p = doc.add_paragraph(style='0000模板')
                    p.add_run(line if line.strip() else ' ')

    # 结论正文：每行一段，均应用「0000模板」样式
    for line in conclusion.split('\n'):
        p = doc.add_paragraph(style='0000模板')
        p.add_run(line if line.strip() else ' ')

    doc.save(docx_path)


@roundtable_router.get("/roundtable/{record_id}/export/docx", dependencies=[Depends(test_token)])
async def export_roundtable_docx(record_id: str):
    """导出圆桌记录为 DOCX 文件。"""
    record = get_record_by_id(record_id)
    if record is None:
        raise HTTPException(status_code=404, detail="记录不存在")
    topic = (record.get("topic") or "圆桌")[:50]
    safe_name = re.sub(r'[\/:*?"<>|]', "_", topic) or "圆桌"
    filename_docx = f"圆桌_{safe_name}.docx"
    encoded_filename_docx = quote(filename_docx, encoding="utf-8")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        _build_roundtable_docx(tmp.name, record)
        try:
            with open(tmp.name, "rb") as f:
                docx_bytes = f.read()
        finally:
            try:
                os.unlink(tmp.name)
            except Exception:
                pass
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename_docx}"},
    )


@roundtable_router.get("/roundtable/{record_id}/export/pdf", dependencies=[Depends(test_token)])
async def export_roundtable_pdf(record_id: str):
    """导出圆桌记录为 PDF 文件（先生成 DOCX 再转换）。"""
    record = get_record_by_id(record_id)
    if record is None:
        raise HTTPException(status_code=404, detail="记录不存在")
    topic = (record.get("topic") or "圆桌")[:50]
    safe_name = re.sub(r'[\/:*?"<>|]', "_", topic) or "圆桌"
    docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_path = tmp.name
        _build_roundtable_docx(docx_path, record)
        from ai_search.ai_service import AISearchService
        svc = AISearchService()
        pdf_bytes = svc._convert_docx_to_pdf(docx_path)
        if not pdf_bytes:
            raise HTTPException(status_code=500, detail="PDF 转换失败")
        filename_pdf = f"圆桌_{safe_name}.pdf"
        encoded_filename_pdf = quote(filename_pdf, encoding="utf-8")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename_pdf}"},
        )
    finally:
        if docx_path and os.path.exists(docx_path):
            try:
                os.unlink(docx_path)
            except Exception:
                pass


@roundtable_router.get("/roundtable/{record_id}", dependencies=[Depends(test_token)])
async def get_record(record_id: str):
    """单条记录详情（含 rounds、conclusion）。"""
    r = get_record_by_id(record_id)
    if r is None:
        raise HTTPException(status_code=404, detail="记录不存在")
    return r


@roundtable_router.delete("/roundtable/{record_id}", dependencies=[Depends(test_token)])
async def delete_roundtable_record(record_id: str):
    """删除圆桌记录。"""
    ok = delete_record(record_id)
    if not ok:
        raise HTTPException(status_code=404, detail="记录不存在")
    return {"ok": True}


@roundtable_router.post("/roundtable/{record_id}/pin", dependencies=[Depends(test_token)])
async def pin_record(record_id: str):
    """置顶/取消置顶，返回操作后的 is_pinned。"""
    try:
        new_pin = toggle_pin(record_id)
        return {"is_pinned": new_pin}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
