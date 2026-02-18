# -*- coding: utf-8 -*-
"""
英文纲目格式刷程序
遍历同级文件夹中的所有 .docx 文件，根据预定义规则自动应用格式。
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn


# 全局字体设置
FONT_NAME = "Century Schoolbook"
FONT_SIZE = Pt(10.5)

# 纲目编号样式映射
ROMAN_STYLE_MAP = {
    "I": "3、大点I",
    "II": "3、大点II-",
    "V": "3、大点II-",
    "X": "3、大点II-",
    "III": "3、大点VII",
    "IV": "3、大点VII",
    "VI": "3、大点VII",
    "IX": "3、大点VII",
    "XI": "3、大点VII",
    "XV": "3、大点VII",
    "XX": "3、大点VII",
    "VII": "3、大点VIII",
    "VIII": "3、大点VIII",
    "XII": "3、大点VIII",
    "XIII": "3、大点VIII",
    "XIV": "3、大点VIII",
    "XVI": "3、大点VIII",
    "XVII": "3、大点VIII",
    "XVIII": "3、大点VIII",
    "XIX": "3、大点VIII",
    "XXVI": "3、大点VIII",
}

# 正则表达式
RE_ROMAN = re.compile(r'^([IVXL]+)\.\s')
RE_UPPER_LETTER = re.compile(r'^([A-H])\.\s')
RE_NUMBER = re.compile(r'^(\d{1,2})\.\s')
RE_LOWER_LETTER = re.compile(r'^([a-z])\.\s')

# 用于替换空格为 Tab 的正则表达式
RE_SPACE_PATTERNS = [
    (re.compile(r'^([IVXL]+)\. '), r'\1.\t'),      # 罗马数字
    (re.compile(r'^([A-H])\. '), r'\1.\t'),        # 大写字母
    (re.compile(r'^(\d{1,2})\. '), r'\1.\t'),      # 阿拉伯数字
    (re.compile(r'^([a-z])\. '), r'\1.\t'),        # 小写字母
]

# 行尾标点
END_PUNCTUATION = '.;:?!)）'

# Markdown 格式字符（刷格式前统一删除）
RE_MARKDOWN = re.compile(r'[*#_`]')


def clean_markdown_chars(paragraph):
    """删除段落中所有 Markdown 格式字符（* # _ `）"""
    for run in paragraph.runs:
        if RE_MARKDOWN.search(run.text):
            run.text = RE_MARKDOWN.sub('', run.text)


def set_run_font(run):
    """设置 run 的字体为 Century Schoolbook, 10.5pt"""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    # 设置中文字体（确保西文字体生效）
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def set_paragraph_font(paragraph):
    """设置段落中所有 run 的字体"""
    for run in paragraph.runs:
        set_run_font(run)


def apply_style_with_font(paragraph, style_name, doc):
    """应用样式并设置字体"""
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        print(f"  警告: 样式 '{style_name}' 不存在")
    set_paragraph_font(paragraph)


def set_center_bold(paragraph):
    """设置段落居中加粗"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = True
        set_run_font(run)


def set_center(paragraph):
    """设置段落居中（不加粗）"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(paragraph)


def set_left_bold(paragraph):
    """设置段落左对齐加粗"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.bold = True
        set_run_font(run)


def get_outline_style(text):
    """根据编号获取纲目样式名称"""
    text = text.strip()

    # 检查罗马数字
    match = RE_ROMAN.match(text)
    if match:
        roman = match.group(1)
        return ROMAN_STYLE_MAP.get(roman, "3、大点VIII")

    # 检查大写字母 A-H
    match = RE_UPPER_LETTER.match(text)
    if match:
        return "4、中点A"

    # 检查数字 1-20
    match = RE_NUMBER.match(text)
    if match:
        num = int(match.group(1))
        if 1 <= num <= 20:
            return "5、小点1"

    # 检查小写字母 a-z
    match = RE_LOWER_LETTER.match(text)
    if match:
        return "6、小小点a"

    return None


def is_title_line(text):
    """检查是否为信息摘录中的标题行（末尾无标点的单独行）"""
    text = text.strip()
    if not text:
        return False
    return text[-1] not in END_PUNCTUATION


def ensure_tab_after_numbering(paragraph):
    """确保编号后面是 Tab 字符而非空格"""
    if not paragraph.runs:
        return

    first_run = paragraph.runs[0]
    text = first_run.text

    for pattern, replacement in RE_SPACE_PATTERNS:
        if pattern.match(text):
            first_run.text = pattern.sub(replacement, text, count=1)
            break


def process_document(doc_path):
    """处理单个文档"""
    print(f"处理文件: {doc_path.name}")

    doc = Document(doc_path)

    # 状态: title, outline, excerpts, references
    state = "title"
    last_title = None  # 记录题目区最后一行（居中加粗段落）

    for para in doc.paragraphs:
        # 刷格式前先清除 Markdown 格式字符
        clean_markdown_chars(para)

        text = para.text.strip()

        if not text:
            # 空段落，保持字体设置
            set_paragraph_font(para)
            continue

        # 状态转换检测
        if text.lower().startswith("excerpts from the ministry:"):
            state = "excerpts"
            set_left_bold(para)
            para.paragraph_format.space_after = Pt(12)
            continue

        if text.lower().startswith("references and further reading:"):
            state = "references"
            set_left_bold(para)
            para.paragraph_format.space_after = Pt(12)
            continue

        # 根据状态处理
        if state == "title":
            if text.lower().startswith("scripture reading:"):
                set_center(para)
                para.paragraph_format.space_after = Pt(12)
            elif RE_ROMAN.match(text):
                state = "outline"
                ensure_tab_after_numbering(para)
                apply_style_with_font(para, get_outline_style(text), doc)
            else:
                set_center_bold(para)
                last_title = text

        elif state == "outline":
            style_name = get_outline_style(text)
            if style_name:
                ensure_tab_after_numbering(para)
                apply_style_with_font(para, style_name, doc)
            else:
                set_paragraph_font(para)

        elif state == "excerpts":
            if is_title_line(text):
                set_center_bold(para)
                para.paragraph_format.space_before = Pt(12)
            else:
                apply_style_with_font(para, "paragraph", doc)

        elif state == "references":
            apply_style_with_font(para, "paragraph", doc)

    # 在 "Excerpts from the Ministry:" 和 "References and Further Reading:" 之前插入分页符
    page_break_keywords = ['excerpts from the ministry:', 'references and further reading:']
    page_break_paras = [para for para in doc.paragraphs
                        if any(kw in para.text.lower() for kw in page_break_keywords)]
    for para in page_break_paras:
        page_break_para = para.insert_paragraph_before()
        run = page_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)

    # 保存文档
    doc.save(doc_path)
    print(f"  已保存: {doc_path.name}")

    # 按题目区最后一行重命名文件
    if last_title:
        safe_name = re.sub(r'[\\/:*?"<>|]', '', last_title).strip()
        if safe_name:
            new_path = doc_path.parent / f"{safe_name}.docx"
            if new_path != doc_path:
                doc_path.rename(new_path)
                print(f"  已重命名为: {new_path.name}")


def main():
    """主函数"""
    # 获取程序所在目录
    script_dir = Path(__file__).parent

    # 遍历所有 .docx 文件
    docx_files = list(script_dir.glob("*.docx"))

    if not docx_files:
        print("未找到 .docx 文件")
        return

    print(f"找到 {len(docx_files)} 个 .docx 文件")
    print("-" * 40)

    for doc_path in docx_files:
        # 跳过临时文件
        if doc_path.name.startswith("~$"):
            continue
        try:
            process_document(doc_path)
        except Exception as e:
            print(f"  错误: {e}")

    print("-" * 40)
    print("处理完成")


if __name__ == "__main__":
    main()
