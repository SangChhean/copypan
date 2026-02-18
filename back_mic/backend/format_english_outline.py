"""
英文纲目格式刷程序 - 函数式版本
用于后端 API 调用，处理单个 DOCX 文件
"""
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 全局字体设置
FONT_NAME = "Century Schoolbook"
FONT_SIZE = Pt(10.5)

# 纲目编号样式映射
ROMAN_STYLE_MAP = {
    "I": "3、大点I",
    "II": "3、大点II-",
    "V": "3、大点II-",
    "X": "3、大点II-",
    "III": "3、大点VII",
    "IV": "3、大点VII",
    "VI": "3、大点VII",
    "IX": "3、大点VII",
    "XI": "3、大点VII",
    "XV": "3、大点VII",
    "XX": "3、大点VII",
    "VII": "3、大点VIII",
    "VIII": "3、大点VIII",
    "XII": "3、大点VIII",
    "XIII": "3、大点VIII",
    "XIV": "3、大点VIII",
    "XVI": "3、大点VIII",
    "XVII": "3、大点VIII",
    "XVIII": "3、大点VIII",
    "XIX": "3、大点VIII",
    "XXVI": "3、大点VIII",
}

# 正则表达式
RE_ROMAN = re.compile(r'^([IVXL]+)\.\s')
RE_UPPER_LETTER = re.compile(r'^([A-H])\.\s')
RE_NUMBER = re.compile(r'^(\d{1,2})\.\s')
RE_LOWER_LETTER = re.compile(r'^([a-z])\.\s')

# 用于替换空格为 Tab 的正则表达式
RE_SPACE_PATTERNS = [
    (re.compile(r'^([IVXL]+)\. '), r'\1.\t'),
    (re.compile(r'^([A-H])\. '), r'\1.\t'),
    (re.compile(r'^(\d{1,2})\. '), r'\1.\t'),
    (re.compile(r'^([a-z])\. '), r'\1.\t'),
]

# 行尾标点
END_PUNCTUATION = '.;:?!)）'

# Markdown 格式字符（刷格式前统一删除）
RE_MARKDOWN = re.compile(r'[*#_`]')


def clean_markdown_chars(paragraph):
    """删除段落中所有 Markdown 格式字符（* # _ `）"""
    for run in paragraph.runs:
        if RE_MARKDOWN.search(run.text):
            run.text = RE_MARKDOWN.sub('', run.text)


def set_run_font(run):
    """设置 run 的字体为 Century Schoolbook, 10.5pt"""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    try:
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except (AttributeError, Exception):
        pass


def set_paragraph_font(paragraph):
    """设置段落中所有 run 的字体"""
    for run in paragraph.runs:
        set_run_font(run)


def apply_style_with_font(paragraph, style_name, doc):
    """应用样式并设置字体"""
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        logger.warning(f"样式 '{style_name}' 不存在")
    set_paragraph_font(paragraph)


def set_center_bold(paragraph):
    """设置段落居中加粗"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = True
        set_run_font(run)


def set_center(paragraph):
    """设置段落居中（不加粗）"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(paragraph)


def set_left_bold(paragraph):
    """设置段落左对齐加粗"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.bold = True
        set_run_font(run)


def get_outline_style(text):
    """根据编号获取纲目样式名称"""
    text = text.strip()
    match = RE_ROMAN.match(text)
    if match:
        roman = match.group(1)
        return ROMAN_STYLE_MAP.get(roman, "3、大点VIII")
    match = RE_UPPER_LETTER.match(text)
    if match:
        return "4、中点A"
    match = RE_NUMBER.match(text)
    if match:
        num = int(match.group(1))
        if 1 <= num <= 20:
            return "5、小点1"
    match = RE_LOWER_LETTER.match(text)
    if match:
        return "6、小小点a"
    return None


def is_title_line(text):
    """检查是否为信息摘录中的标题行（末尾无标点的单独行）"""
    text = text.strip()
    if not text:
        return False
    return text[-1] not in END_PUNCTUATION


def ensure_tab_after_numbering(paragraph):
    """确保编号后面是 Tab 字符而非空格"""
    if not paragraph.runs:
        return
    first_run = paragraph.runs[0]
    text = first_run.text
    for pattern, replacement in RE_SPACE_PATTERNS:
        if pattern.match(text):
            first_run.text = pattern.sub(replacement, text, count=1)
            break


def format_english_outline_docx(docx_path: str) -> None:
    """
    格式化英文纲目 DOCX 文件（原地修改）。
    
    Args:
        docx_path: DOCX 文件路径
    """
    doc = Document(docx_path)
    state = "title"
    last_title = None

    for para in doc.paragraphs:
        clean_markdown_chars(para)
        text = para.text.strip()

        if not text:
            set_paragraph_font(para)
            continue

        # 状态转换检测
        if text.lower().startswith("excerpts from the ministry:"):
            state = "excerpts"
            set_left_bold(para)
            para.paragraph_format.space_after = Pt(12)
            continue

        if text.lower().startswith("references and further reading:"):
            state = "references"
            set_left_bold(para)
            para.paragraph_format.space_after = Pt(12)
            continue

        # 根据状态处理
        if state == "title":
            if text.lower().startswith("scripture reading:"):
                set_center(para)
                para.paragraph_format.space_after = Pt(12)
            elif RE_ROMAN.match(text):
                state = "outline"
                ensure_tab_after_numbering(para)
                apply_style_with_font(para, get_outline_style(text), doc)
            else:
                set_center_bold(para)
                last_title = text

        elif state == "outline":
            style_name = get_outline_style(text)
            if style_name:
                ensure_tab_after_numbering(para)
                apply_style_with_font(para, style_name, doc)
            else:
                set_paragraph_font(para)

        elif state == "excerpts":
            if is_title_line(text):
                set_center_bold(para)
                para.paragraph_format.space_before = Pt(12)
            else:
                apply_style_with_font(para, "paragraph", doc)

        elif state == "references":
            apply_style_with_font(para, "paragraph", doc)

    # 在特定关键词之前插入分页符
    page_break_keywords = ['excerpts from the ministry:', 'references and further reading:']
    page_break_paras = [para for para in doc.paragraphs
                        if any(kw in para.text.lower() for kw in page_break_keywords)]
    for para in page_break_paras:
        page_break_para = para.insert_paragraph_before()
        run = page_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)

    # 保存文档
    doc.save(docx_path)
