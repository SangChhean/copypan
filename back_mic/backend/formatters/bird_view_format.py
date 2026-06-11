"""
鸟瞰纲目刷格式
专用于词典-鸟瞰纲目的 DOCX 生成与格式化。
篇头固定四段：
  第1段：主恢复真理的词典（0系列样式）
  第2段：关键词/篇题（11111西列样式）
  第3段：第三段内容（00篇题样式）
  第4段：鸟瞰类型行，如「3b　节期纲目的鸟瞰」（00篇题样式）
第5段起：读经、纲目正文（走中文纲目刷格式逻辑）
"""
from __future__ import annotations
import logging
import re
import shutil
import tempfile
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from formatters.format_chinese_outline import format_chinese_outline_docx

logger = logging.getLogger(__name__)

# 篇头样式映射（与节期纲目模板一致）
_HEADER_STYLES = ["0系列", "11111西列", "00篇题", "00篇题"]

# 鸟瞰类型第四段文字
BIRD_VIEW_LINE4 = {
    "ministry": "3a\t职事信息的鸟瞰",
    "feast":    "3b\t节期纲目的鸟瞰",
}


def format_bird_view_docx(
    outline_text: str,
    keyword: str,
    bird_type: str = "feast",
    with_source: bool = False,
) -> bytes:
    """
    将鸟瞰纲目正文生成格式化 DOCX，返回字节流。

    Args:
        outline_text: LLM 生成的纲目正文（纯文本，含「读经：」行）
        keyword:      关键词，作为第二段篇题
        bird_type:    "ministry" 或 "feast"
        with_source:  是否为带出处模式（触发标红/标绿着色）

    Returns:
        docx 字节流
    """
    backend_dir = Path(__file__).resolve().parent.parent

    # 1. 找模板（优先节期纲目模板）
    template_path = None
    for name in ("节期纲目模板.docx", "中文纲目模板.docx"):
        p = backend_dir / name
        if p.exists():
            template_path = p
            break
    if template_path is None:
        raise FileNotFoundError("找不到纲目模板文件（节期纲目模板.docx 或 中文纲目模板.docx）")

    # 2. 复制模板到临时文件
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    tmp_path = tmp.name
    try:
        shutil.copy2(str(template_path), tmp_path)

        # 3. 清空模板段落，写入篇头 + 正文
        doc = Document(tmp_path)
        for para in list(doc.paragraphs):
            para._element.getparent().remove(para._element)

        # 篇头四段
        line4 = BIRD_VIEW_LINE4.get(bird_type, BIRD_VIEW_LINE4["feast"])
        headers = [
            "主恢复真理的词典",
            keyword.strip() or "（篇题）",
            "第三段\t鸟瞰的纲目",
            line4,
        ]
        for text in headers:
            doc.add_paragraph(text)

        # 正文按行写入（跳过第一行如果是篇题）
        outline_text = outline_text.replace("\r\n", "\n").replace("\r", "\n")
        lines = outline_text.split("\n")
        # 跳过开头与关键词相同的行
        start = 0
        for i, line in enumerate(lines):
            if line.strip() == keyword.strip():
                start = i + 1
                break
            elif line.strip():
                break
        for line in lines[start:]:
            if line.strip():
                doc.add_paragraph(line)
            elif len(doc.paragraphs) > 4:
                doc.add_paragraph("")

        doc.save(tmp_path)

        # 4. 套用篇头样式（前四段）
        doc2 = Document(tmp_path)
        for i, style_name in enumerate(_HEADER_STYLES):
            if i < len(doc2.paragraphs):
                try:
                    doc2.paragraphs[i].style = doc2.styles[style_name]
                except KeyError:
                    logger.warning("样式不存在：%s，跳过", style_name)
        doc2.save(tmp_path)

        # 5. 调用中文纲目格式刷处理第5段起的纲目内容
        format_chinese_outline_docx(tmp_path, traditional_quotes=False)

        # Step 10/11：出处标红 + 无出处行标绿 + 关键词标红（仅 with_source 模式）
        if with_source:
            doc_color = Document(tmp_path)
            _colorize_bird_view_sources(doc_color, keyword)
            doc_color.save(tmp_path)

        # 6. 读取字节流返回
        with open(tmp_path, "rb") as f:
            return f.read()

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _is_scripture_reference(text: str) -> bool:
    """判断括号内容是否为圣经经节出处（如：弗一10、创一26~28）"""
    pattern = r'^[\u4e00-\u9fa5]{1,3}[一二三四五六七八九十百千0-9]+[上下~\-,，、0-9一二三四五六七八九十]*$'
    return bool(re.match(pattern, text.strip()))


def _colorize_bird_view_sources(doc, keyword: str) -> None:
    """
    鸟瞰纲目带出处模式的着色逻辑：
    - 以 ）结尾且括号内非圣经经节 → 括号部分（含括号）标红
    - 不以 ）结尾且非读经行 → 整行标绿（找不到出处）
    - keyword 在全文中出现处标红（优先级高于标绿）
    跳过前4段篇头。
    """
    has_bracket_ending = False
    for i, para in enumerate(doc.paragraphs):
        if i < 4:
            continue
        text = para.text.strip()
        if text.endswith('）'):
            left = text.rfind('（')
            if left != -1:
                inner = text[left + 1:-1]
                if not _is_scripture_reference(inner):
                    has_bracket_ending = True
                    break

    if not has_bracket_ending:
        return

    keywords = [kw.strip() for kw in keyword.split('、') if kw.strip()] if keyword else []

    for i, para in enumerate(doc.paragraphs):
        if i < 4:
            # 第2段（index=1）是关键词段，只做关键词标红
            if i == 1 and keywords:
                text = para.text.strip()
                if any(kw in text for kw in keywords):
                    para.clear()
                    _add_run_with_keyword_red(para, text, keywords)
            continue
        text = para.text.strip()
        if not text:
            continue
        is_reading = text.startswith('读经：')
        ends_with_bracket = text.endswith('）')

        if ends_with_bracket and not is_reading:
            nested = 0
            left_index = -1
            for j, ch in enumerate(text):
                if ch == '（':
                    if nested == 0:
                        left_index = j
                    nested += 1
                elif ch == '）':
                    nested -= 1
            if left_index != -1:
                inner = text[left_index + 1:text.rfind('）')]
                if not _is_scripture_reference(inner):
                    para.clear()
                    before = text[:left_index]
                    bracket = text[left_index:]
                    if before:
                        _add_run_with_keyword_red(para, before, keywords)
                    run_bracket = para.add_run(bracket)
                    run_bracket.font.color.rgb = RGBColor(0xFF, 0, 0)
                    continue

        if not ends_with_bracket and not is_reading:
            para.clear()
            run_green = para.add_run(text)
            run_green.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if keywords and any(kw in text for kw in keywords):
                para.clear()
                _add_run_with_keyword_red(para, text, keywords, apply_highlight=True)


def _add_run_with_keyword_red(para, text: str, keywords: list, apply_highlight: bool = False) -> None:
    """将文本写入段落，关键词标红，其余部分按 apply_highlight 决定是否标绿。"""
    if not keywords:
        run = para.add_run(text)
        if apply_highlight:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        return
    pattern = '(' + '|'.join(re.escape(kw) for kw in keywords) + ')'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part)
        if part in keywords:
            run.font.color.rgb = RGBColor(0xFF, 0, 0)
        elif apply_highlight:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
