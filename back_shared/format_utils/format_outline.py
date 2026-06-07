import os
import shutil
import base64
import tempfile
import asyncio
import re
import logging
import subprocess
import sys
from pathlib import Path
from docx import Document
from typing import Optional

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
ZH_TEMPLATE = TEMPLATES_DIR / "中文纲目模板.docx"
EN_TEMPLATE = TEMPLATES_DIR / "英文纲目模板.docx"

_ZH_OUTLINE_PATTERNS = [
    re.compile(r"^[壹贰叁肆伍陆柒捌玖拾佰仟萬亿億貳參叄叁参肆伍陸陆柒捌玖拾][\t　]"),
    re.compile(r"^[一二三四五六七八九十百千]+[\t　]"),
    re.compile(r"^\d+\t"),
    re.compile(r"^[a-z]\t"),
]
_EN_OUTLINE_PATTERNS = [
    re.compile(r"^(I{1,3}|IV|VI{0,3}|IX|X{1,3}|XI{0,3}|XIV|XV|XVI{0,3}|XIX|XX)\.\s", re.I),
    re.compile(r"^[A-H]\.\s"),
    re.compile(r"^[1-9]\.\s"),
    re.compile(r"^[a-k]\.\s"),
]

_ZH_PUNCT_END = set("。，；：？！、")
_EN_PUNCT_END = set(".,;:!?")


def _apply_style_if_exists(doc: Document, para, style_name: str) -> None:
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        logger.debug("样式不存在，跳过: %s", style_name)


def detect_doc_type(text: str) -> str:
    """扫描文本前30行，判断 outline_zh / outline_en / prose。"""
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")[:30]
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        for pat in _ZH_OUTLINE_PATTERNS:
            if pat.match(stripped):
                return "outline_zh"
        for pat in _EN_OUTLINE_PATTERNS:
            if pat.match(stripped):
                return "outline_en"
    return "prose"


def _is_zh_direction(direction: str) -> bool:
    return direction in ("zh", "zh_tw")


def _ends_with_punct(text: str, direction: str) -> bool:
    if not text:
        return False
    punct = _ZH_PUNCT_END if _is_zh_direction(direction) else _EN_PUNCT_END
    return text[-1] in punct


def apply_prose_format(doc: Document, direction: str) -> None:
    """正文模式：篇题 / 小标题 / 正文 / 职事信息摘录段。"""
    is_zh = _is_zh_direction(direction)
    title_style = "00篇题" if is_zh else "11篇题"
    subtitle_style = "81级标题" if is_zh else "B2摘录标题1级"
    body_style = "0000模板" if is_zh else "00正文"
    ministry_marker_styles = {
        "职事信息摘录：": "9职事信息摘录",
        "職事信息摘錄：": "9职事信息摘录",
        "Excerpts from the Ministry:": "9职事信息摘录",
        "excerpts from the ministry:": "9职事信息摘录",
    }

    first_non_empty_done = False
    in_ministry = False

    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()
        if not stripped:
            continue

        lower = stripped.lower()
        marker_hit = None
        for marker, style in ministry_marker_styles.items():
            if marker in stripped or (marker.lower() in lower and not is_zh):
                marker_hit = style
                in_ministry = True
                break

        if marker_hit:
            _apply_style_if_exists(doc, para, marker_hit)
            continue

        if in_ministry:
            if is_zh:
                if len(stripped) < 30 and not _ends_with_punct(stripped, direction):
                    _apply_style_if_exists(doc, para, subtitle_style)
                else:
                    _apply_style_if_exists(doc, para, body_style)
            else:
                if len(stripped) < 30 and not _ends_with_punct(stripped, direction):
                    _apply_style_if_exists(doc, para, subtitle_style)
                else:
                    _apply_style_if_exists(doc, para, body_style)
            continue

        if not first_non_empty_done:
            _apply_style_if_exists(doc, para, title_style)
            first_non_empty_done = True
            continue

        if len(stripped) < 30 and not _ends_with_punct(stripped, direction):
            _apply_style_if_exists(doc, para, subtitle_style)
        else:
            _apply_style_if_exists(doc, para, body_style)


def convert_to_pdf(docx_path: str) -> Optional[bytes]:
    """将 DOCX 转为 PDF bytes；LibreOffice 优先，docx2pdf 回退。"""
    if not os.path.exists(docx_path):
        logger.error("DOCX 文件不存在: %s", docx_path)
        return None

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        pdf_path = tmp_pdf.name

    logger.info("开始转换 DOCX 到 PDF: %s -> %s", docx_path, pdf_path)

    try:
        result = subprocess.run(
            ["which", "libreoffice"],
            capture_output=True,
            text=True,
            timeout=5,
        )
        if result.returncode == 0:
            output_dir = os.path.dirname(pdf_path)
            temp_output_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            pdf_export_opts = 'pdf:writer_pdf_Export:{"SelectPdfVersion":{"type":"long","value":"2"}}'
            convert_result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    pdf_export_opts,
                    "--outdir",
                    output_dir,
                    docx_path,
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if convert_result.returncode == 0:
                generated_pdf = os.path.join(output_dir, temp_output_name)
                if os.path.exists(generated_pdf):
                    if generated_pdf != pdf_path:
                        os.rename(generated_pdf, pdf_path)
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        with open(pdf_path, "rb") as f:
                            pdf_bytes = f.read()
                        logger.info("PDF 转换成功（LibreOffice）: 大小 %s bytes", len(pdf_bytes))
                        try:
                            os.unlink(pdf_path)
                        except Exception:
                            pass
                        return pdf_bytes
    except FileNotFoundError:
        logger.debug("LibreOffice 未找到，尝试其他方法")
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice 转换超时")
    except Exception as e:
        logger.warning("LibreOffice 转换失败: %s", e)

    if sys.platform == "win32":
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except Exception:
            pass
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            logger.info("PDF 转换成功（docx2pdf）: 大小 %s bytes", len(pdf_bytes))
            try:
                os.unlink(pdf_path)
            except Exception:
                pass
            return pdf_bytes
    except NotImplementedError as e:
        logger.error("DOCX 转 PDF 失败: %s", e)
    except Exception as e:
        logger.error("DOCX 转 PDF 失败: %s", e, exc_info=True)

    try:
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
    except Exception:
        pass
    return None


def _pick_template(direction: str) -> Path:
    if direction in ("zh", "zh_tw"):
        return ZH_TEMPLATE
    if direction in ("en", "es"):
        return EN_TEMPLATE
    raise ValueError(f"不支持的方向: {direction}")


def _default_filename(direction: str, output_format: str) -> str:
    base = {
        "zh": "outline_zh",
        "zh_tw": "outline_zh_tw",
        "en": "outline_en",
        "es": "outline_es",
    }.get(direction, "document")
    ext = "pdf" if output_format == "pdf" else "docx"
    return f"{base}.{ext}"


def _build_filename(text: str, direction: str, output_format: str, filename: str) -> str:
    ext = "pdf" if output_format == "pdf" else "docx"
    lines = [ln.strip() for ln in (text or "").split("\n") if ln.strip()]
    title_line = lines[0] if lines else ""
    if title_line:
        safe = re.sub(r'[\/:*?"<>|]', "_", title_line[:50])
        if safe:
            return f"{safe}.{ext}"
    if filename:
        name = filename if filename.endswith(f".{ext}") else f"{filename}.{ext}"
        return name
    return _default_filename(direction, output_format)


def format_and_download(
    text: str,
    direction: str,
    output_format: str,
    filename: str = "",
) -> dict:
    """统一刷格式并返回 base64 编码文件。"""
    error = None
    temp_docx_path = None
    try:
        template = _pick_template(direction)
        if not template.exists():
            return {
                "content_base64": "",
                "filename": filename or _default_filename(direction, output_format),
                "format": output_format,
                "error": f"模板文件不存在: {template.name}",
            }

        with tempfile.TemporaryDirectory() as tmpdir:
            temp_docx_path = os.path.join(tmpdir, "formatted.docx")
            shutil.copy2(template, temp_docx_path)

            doc = Document(temp_docx_path)
            for para in list(doc.paragraphs):
                p_element = para._element
                p_element.getparent().remove(p_element)

            lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
            for line in lines:
                if line.strip() or len(doc.paragraphs) == 0:
                    doc.add_paragraph(line)
            doc.save(temp_docx_path)

            doc_type = detect_doc_type(text)

            if doc_type == "outline_zh" and direction in ("zh", "zh_tw"):
                from .format_chinese import format_chinese_outline_docx
                format_chinese_outline_docx(
                    temp_docx_path,
                    traditional_quotes=(direction == "zh_tw"),
                )
            elif doc_type == "outline_en" and direction in ("en", "es"):
                from .format_english import format_english_outline_docx
                format_english_outline_docx(temp_docx_path)
            else:
                doc = Document(temp_docx_path)
                apply_prose_format(doc, direction)
                doc.save(temp_docx_path)

            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()

            out_filename = _build_filename(text, direction, output_format, filename)
            result_format = "docx"
            content_bytes = docx_bytes

            if output_format == "pdf":
                pdf_bytes = convert_to_pdf(temp_docx_path)
                if pdf_bytes:
                    content_bytes = pdf_bytes
                    result_format = "pdf"
                else:
                    error = "PDF 转换失败，已返回 DOCX 文件"
                    out_filename = out_filename.replace(".pdf", ".docx")
                    result_format = "docx"
            else:
                result_format = "docx"

            return {
                "content_base64": base64.b64encode(content_bytes).decode("ascii"),
                "filename": out_filename,
                "format": result_format,
                "error": error,
            }
    except Exception as e:
        logger.error("format_and_download 失败: %s", e, exc_info=True)
        return {
            "content_base64": "",
            "filename": filename or _default_filename(direction, output_format),
            "format": output_format if output_format in ("docx", "pdf") else "docx",
            "error": str(e),
        }
